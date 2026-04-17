from __future__ import annotations

import base64
import io
import json
import os
import re
import zipfile
from datetime import datetime, timezone
from typing import Any

import pandas as pd
import requests
import streamlit as st

APP_VERSION = "superadmin-autonomous-2026-04-17-v1"
RESPONSE_PATH_ROOT = "data/validation_doc"
TEST_EMAILS = {"kl@od.sd", "in@bc.sd", "de@re.bh", "gh@fg.jh"}

st.set_page_config(page_title="Dashboard Superadmin - validation du document", layout="wide")


def require_password() -> None:
    expected = ""
    try:
        expected = st.secrets.get("SUPERADMIN_PASSWORD", "")
    except Exception:
        expected = ""
    if not expected:
        st.warning("Aucun mot de passe Superadmin n’est défini dans les secrets Streamlit. Le tableau de bord est ouvert.")
        return
    if st.session_state.get("superadmin_ok"):
        return
    st.title("Dashboard Superadmin")
    pwd = st.text_input("Mot de passe Superadmin", type="password")
    if st.button("Ouvrir"):
        if pwd == expected:
            st.session_state.superadmin_ok = True
            st.rerun()
        else:
            st.error("Mot de passe incorrect.")
    st.stop()


def _sanitize_secret(value: Any) -> str:
    if value is None:
        return ""
    value = str(value).strip().strip('"').strip("'")
    value = value.replace("\r", "").replace("\n", "").strip()
    if value.lower().startswith("bearer "):
        value = value[7:].strip()
    return value


def get_github_config_from_streamlit() -> dict[str, str]:
    owner = os.getenv("GITHUB_OWNER", "mniangj-png")
    repo = os.getenv("GITHUB_REPO", "consultation-statafric_niang")
    branch = os.getenv("GITHUB_BRANCH", "main")
    token = os.getenv("GITHUB_TOKEN", "") or os.getenv("GH_TOKEN", "")
    try:
        gh = st.secrets.get("github", {})
        owner = _sanitize_secret(gh.get("owner", owner)) or owner
        repo = _sanitize_secret(gh.get("repo", repo)) or repo
        branch = _sanitize_secret(gh.get("branch", branch)) or branch
        token = _sanitize_secret(gh.get("token", token)) or token
    except Exception:
        pass
    try:
        token = _sanitize_secret(st.secrets.get("GITHUB_TOKEN", token)) or token
    except Exception:
        pass
    return {"owner": owner, "repo": repo, "branch": branch, "token": token}


def github_headers(cfg: dict[str, str]) -> dict[str, str]:
    headers = {"Accept": "application/vnd.github+json"}
    token = _sanitize_secret(cfg.get("token", ""))
    if token:
        headers["Authorization"] = f"Bearer {token}"
    return headers


def github_api_get(cfg: dict[str, str], url: str, **kwargs) -> requests.Response:
    r = requests.get(url, headers=github_headers(cfg), timeout=30, **kwargs)
    r.raise_for_status()
    return r


def list_json_paths(cfg: dict[str, str], source_kind: str) -> list[str]:
    root = f"{RESPONSE_PATH_ROOT}/{source_kind}"
    base = f"https://api.github.com/repos/{cfg['owner']}/{cfg['repo']}/contents/{root}"
    paths: list[str] = []
    stack = [(base, root)]
    while stack:
        url, path = stack.pop()
        r = github_api_get(cfg, url, params={"ref": cfg["branch"]})
        payload = r.json()
        if isinstance(payload, dict) and payload.get("type") == "file":
            if str(payload.get("path", "")).lower().endswith(".json"):
                paths.append(payload["path"])
            continue
        if not isinstance(payload, list):
            continue
        for item in payload:
            item_type = item.get("type")
            item_path = item.get("path", "")
            item_url = item.get("url")
            if item_type == "dir" and item_url:
                stack.append((item_url, item_path))
            elif item_type == "file" and str(item_path).lower().endswith(".json"):
                paths.append(item_path)
    return sorted(paths)


def load_json_record(cfg: dict[str, str], path: str) -> dict[str, Any] | None:
    url = f"https://api.github.com/repos/{cfg['owner']}/{cfg['repo']}/contents/{path}"
    try:
        r = github_api_get(cfg, url, params={"ref": cfg["branch"]})
        item = r.json()
        content = base64.b64decode(item["content"]).decode("utf-8")
        data = json.loads(content)
        if isinstance(data, dict):
            data["_source_path"] = path
            return data
    except Exception:
        return None
    return None


@st.cache_data(show_spinner=False)
def load_submissions() -> tuple[str, list[dict[str, Any]], list[str]]:
    cfg = get_github_config_from_streamlit()
    paths = list_json_paths(cfg, "submissions")
    records = []
    for path in paths:
        rec = load_json_record(cfg, path)
        if rec is not None:
            records.append(rec)
    return cfg["branch"], records, paths


def records_to_dataframe(records: list[dict[str, Any]]) -> pd.DataFrame:
    if not records:
        return pd.DataFrame()
    df = pd.DataFrame(records)
    for col in ["submitted_at", "saved_at"]:
        if col in df.columns:
            df[col] = pd.to_datetime(df[col], errors="coerce", utc=True)
    return df


def _norm_text(value: Any) -> str:
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    return re.sub(r"\s+", " ", str(value).strip().lower())


def _respondent_key(row: pd.Series) -> str:
    email = _norm_text(row.get("email"))
    if email and email not in TEST_EMAILS:
        return f"email:{email}"
    parts = [
        _norm_text(row.get("institution_type")),
        _norm_text(row.get("country_or_rec")),
        _norm_text(row.get("institution_acronym")),
        _norm_text(row.get("respondent_title")),
    ]
    return "fallback:" + "|".join(parts)


def clean_records_df(df: pd.DataFrame) -> tuple[pd.DataFrame, dict[str, int], pd.DataFrame]:
    stats = {"source_rows": int(len(df)), "tests_removed": 0, "duplicates_removed": 0, "final_rows": 0}
    if df.empty:
        return df.copy(), stats, pd.DataFrame()

    work = df.copy()
    if "email" not in work.columns:
        work["email"] = ""
    work["_email_norm"] = work["email"].map(_norm_text)
    work["_respondent_key"] = work.apply(_respondent_key, axis=1)

    if "submitted_at" in work.columns:
        work["_sort_time"] = pd.to_datetime(work["submitted_at"], errors="coerce", utc=True)
    elif "saved_at" in work.columns:
        work["_sort_time"] = pd.to_datetime(work["saved_at"], errors="coerce", utc=True)
    else:
        work["_sort_time"] = pd.NaT

    journal_rows = []

    test_mask = work["_email_norm"].isin(TEST_EMAILS)
    if test_mask.any():
        tests = work.loc[test_mask].copy()
        tests["_removal_reason"] = "test_email"
        journal_rows.append(tests)
    stats["tests_removed"] = int(test_mask.sum())
    work = work.loc[~test_mask].copy()

    ranked = work.sort_values(
        by=["_respondent_key", "_sort_time", "_source_path"],
        ascending=[True, False, False],
        na_position="last",
    ).copy()
    ranked["_rank_within_respondent"] = ranked.groupby("_respondent_key").cumcount() + 1
    dup_mask = ranked["_rank_within_respondent"] > 1
    if dup_mask.any():
        dups = ranked.loc[dup_mask].copy()
        dups["_removal_reason"] = "older_duplicate"
        journal_rows.append(dups)
    stats["duplicates_removed"] = int(dup_mask.sum())

    cleaned = ranked.loc[~dup_mask].copy()
    stats["final_rows"] = int(len(cleaned))

    journal = pd.concat(journal_rows, ignore_index=True) if journal_rows else pd.DataFrame()

    drop_cols = ["_email_norm", "_respondent_key", "_sort_time", "_rank_within_respondent"]
    cleaned.drop(columns=drop_cols, inplace=True, errors="ignore")
    if not journal.empty:
        journal.drop(columns=drop_cols, inplace=True, errors="ignore")
    return cleaned, stats, journal


def _safe_col(df: pd.DataFrame, col: str) -> pd.Series:
    if col in df.columns:
        return df[col]
    return pd.Series(dtype="object")


def build_summary_sheet(df: pd.DataFrame, stats: dict[str, int]) -> pd.DataFrame:
    rows = [
        {"indicateur": "Lignes source", "valeur": stats["source_rows"]},
        {"indicateur": "Tests supprimés", "valeur": stats["tests_removed"]},
        {"indicateur": "Doublons supprimés", "valeur": stats["duplicates_removed"]},
        {"indicateur": "Lignes finales conservées", "valeur": stats["final_rows"]},
        {"indicateur": "Institutions distinctes", "valeur": _safe_col(df, "institution_acronym").replace("", pd.NA).dropna().nunique()},
        {"indicateur": "Pays / CER distincts", "valeur": _safe_col(df, "country_or_rec").replace("", pd.NA).dropna().nunique()},
        {"indicateur": "Langues distinctes", "valeur": _safe_col(df, "language").replace("", pd.NA).dropna().nunique()},
    ]
    return pd.DataFrame(rows)


def _value_counts_sheet(df: pd.DataFrame, col: str, sheet_label: str) -> pd.DataFrame:
    if col not in df.columns:
        return pd.DataFrame(columns=[sheet_label, "effectif"])
    vc = df[col].fillna("NA").astype(str).value_counts(dropna=False).reset_index()
    vc.columns = [sheet_label, "effectif"]
    return vc


def build_traceability_sheet(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    source_time_col = "submitted_at" if "submitted_at" in df.columns else ("saved_at" if "saved_at" in df.columns else None)

    def add_rows(section: str, item_map: dict[str, str], reason_suffix: str = "_why"):
        for code, label in item_map.items():
            if code not in df.columns:
                continue
            cols = [c for c in ["institution_acronym", "country_or_rec", "email", "submission_id", source_time_col, code, f"{code}{reason_suffix}"] if c and c in df.columns]
            subset = df[cols].copy()
            subset = subset.loc[subset[code].isin(["go_with_reservations", "no_go"])]
            if subset.empty:
                continue
            subset["section"] = section
            subset["item_code"] = code
            subset["item_label"] = label
            subset["position"] = subset[code]
            why_col = f"{code}{reason_suffix}"
            subset["commentaire"] = subset[why_col] if why_col in subset.columns else ""
            keep_cols = ["section", "item_code", "item_label", "position", "commentaire", "institution_acronym", "country_or_rec", "email", "submission_id"]
            if source_time_col:
                keep_cols.append(source_time_col)
            rows.append(subset[keep_cols])

    strategic_map = {
        "strategic_prioritization_criteria": "Critères de priorisation retenus",
        "strategic_scoring_logic": "Logique de notation multicritère",
        "strategic_core_extensions": "Distinction noyau / extensions",
        "strategic_gender_integration": "Intégration transversale du genre",
        "strategic_min_disaggregations": "Désagrégations minimales proposées",
        "strategic_data_sources": "Sources de données et dispositifs de production",
        "strategic_governance_roles": "Gouvernance et répartition des rôles",
        "strategic_roadmap_update": "Feuille de route de mise en œuvre et mécanisme de mise à jour",
    }
    domain_map = {
        "domain_d01": "D01 Croissance économique, transformation structurelle et commerce",
        "domain_d02": "D02 Emploi, travail décent et protection sociale",
        "domain_d03": "D03 Agriculture durable, sécurité alimentaire et nutrition",
        "domain_d04": "D04 Infrastructures, industrialisation et innovation",
        "domain_d05": "D05 Inclusion, pauvreté et inégalités",
        "domain_d06": "D06 Éducation, compétences et capital humain",
        "domain_d07": "D07 Santé, bien-être et accès universel",
        "domain_d08": "D08 Égalité des genres et autonomisation",
        "domain_d09": "D09 Environnement, résilience climatique et villes durables",
        "domain_d10": "D10 Gouvernance, paix et institutions",
        "domain_d11": "D11 Économie bleue et gestion des océans",
        "domain_d12": "D12 Partenariats et financement du développement",
    }
    add_rows("Choix stratégiques", strategic_map)
    add_rows("Domaines", domain_map)

    for col, label in [
        ("overall_validation_why", "Validation globale"),
        ("operational_usability_why", "Caractère opérationnel"),
        ("final_institutional_position_why", "Position finale"),
        ("strategic_comments", "Commentaires stratégiques"),
        ("domain_comments", "Commentaires domaines"),
        ("top_3_revisions", "Révisions prioritaires"),
    ]:
        if col in df.columns:
            cols = [c for c in ["institution_acronym", "country_or_rec", "email", "submission_id", source_time_col, col] if c and c in df.columns]
            subset = df[cols].copy()
            subset = subset.loc[subset[col].fillna("").astype(str).str.strip() != ""]
            if not subset.empty:
                subset["section"] = "Commentaires transversaux"
                subset["item_code"] = col
                subset["item_label"] = label
                subset["position"] = ""
                subset["commentaire"] = subset[col]
                keep_cols = ["section", "item_code", "item_label", "position", "commentaire", "institution_acronym", "country_or_rec", "email", "submission_id"]
                if source_time_col:
                    keep_cols.append(source_time_col)
                rows.append(subset[keep_cols])

    if not rows:
        return pd.DataFrame(columns=["section", "item_code", "item_label", "position", "commentaire", "institution_acronym", "country_or_rec", "email", "submission_id"])
    return pd.concat(rows, ignore_index=True)


def build_arbitration_sheet(trace_df: pd.DataFrame) -> pd.DataFrame:
    if trace_df.empty:
        return pd.DataFrame(columns=["section", "item_code", "item_label", "reserves_ou_rejets", "commentaires_non_vides", "orientation_recommandee"])
    grouped = (
        trace_df.assign(commentaire_non_vide=trace_df["commentaire"].fillna("").astype(str).str.strip().ne(""))
        .groupby(["section", "item_code", "item_label"], dropna=False)
        .agg(reserves_ou_rejets=("item_code", "size"), commentaires_non_vides=("commentaire_non_vide", "sum"))
        .reset_index()
    )
    grouped["orientation_recommandee"] = grouped["reserves_ou_rejets"].apply(lambda n: "Arbitrage explicite recommandé" if n >= 2 else "A examiner / consolider")
    return grouped


def build_consolidation_support(df: pd.DataFrame, trace_df: pd.DataFrame) -> pd.DataFrame:
    if trace_df.empty:
        return pd.DataFrame(columns=["reference", "section", "item_label", "commentaire", "institution_acronym", "country_or_rec", "email", "arbitrage", "statut_traitement"])
    comments = trace_df[["section", "item_label", "commentaire", "institution_acronym", "country_or_rec", "email"]].copy()
    comments["reference"] = comments.index + 1
    comments["arbitrage"] = ""
    comments["statut_traitement"] = ""
    return comments[["reference", "section", "item_label", "commentaire", "institution_acronym", "country_or_rec", "email", "arbitrage", "statut_traitement"]]


def build_matrix_sheet(df: pd.DataFrame, row_cols: list[str], index_col: str, title_col_name: str) -> pd.DataFrame:
    cols_present = [c for c in row_cols if c in df.columns]
    if not cols_present or index_col not in df.columns:
        return pd.DataFrame()
    matrix = df[[index_col] + cols_present].copy()
    matrix[index_col] = matrix[index_col].fillna("").replace("", "NA")
    grouped = []
    for col in cols_present:
        tmp = matrix.groupby([index_col, col], dropna=False).size().reset_index(name="effectif")
        tmp["item"] = col
        grouped.append(tmp)
    out = pd.concat(grouped, ignore_index=True) if grouped else pd.DataFrame()
    out.rename(columns={index_col: title_col_name}, inplace=True)
    return out


def build_note_support(df: pd.DataFrame, stats: dict[str, int]) -> pd.DataFrame:
    rows = [
        {"rubrique": "Base", "message": f"{stats['final_rows']} réponses finales retenues après nettoyage."},
        {"rubrique": "Nettoyage", "message": f"{stats['tests_removed']} données de test supprimées."},
        {"rubrique": "Nettoyage", "message": f"{stats['duplicates_removed']} doublons supprimés en conservant la dernière réponse par répondant."},
    ]
    if "overall_validation" in df.columns:
        vc = df["overall_validation"].fillna("NA").astype(str).value_counts()
        for k, v in vc.items():
            rows.append({"rubrique": "Validation globale", "message": f"{k} : {v}"})
    if "final_institutional_position" in df.columns:
        vc = df["final_institutional_position"].fillna("NA").astype(str).value_counts()
        for k, v in vc.items():
            rows.append({"rubrique": "Position finale", "message": f"{k} : {v}"})
    return pd.DataFrame(rows)


def build_analysis_sheets(df: pd.DataFrame, stats: dict[str, int], journal: pd.DataFrame) -> dict[str, pd.DataFrame]:
    trace_df = build_traceability_sheet(df)
    return {
        "synthese_kpis": build_summary_sheet(df, stats),
        "base_nettoyee": df.copy(),
        "journal_nettoyage": journal.copy() if not journal.empty else pd.DataFrame(columns=["_removal_reason"]),
        "validations_globales": _value_counts_sheet(df, "overall_validation", "overall_validation"),
        "positions_finales": _value_counts_sheet(df, "final_institutional_position", "final_institutional_position"),
        "traceabilite_commentaires": trace_df,
        "arbitrages_recommandes": build_arbitration_sheet(trace_df),
        "version_consolidee_support": build_consolidation_support(df, trace_df),
        "matrice_strategique": build_matrix_sheet(
            df,
            [
                "strategic_prioritization_criteria",
                "strategic_scoring_logic",
                "strategic_core_extensions",
                "strategic_gender_integration",
                "strategic_min_disaggregations",
                "strategic_data_sources",
                "strategic_governance_roles",
                "strategic_roadmap_update",
            ],
            "institution_acronym",
            "institution",
        ),
        "matrice_domaines": build_matrix_sheet(
            df,
            [
                "domain_d01", "domain_d02", "domain_d03", "domain_d04", "domain_d05", "domain_d06",
                "domain_d07", "domain_d08", "domain_d09", "domain_d10", "domain_d11", "domain_d12",
            ],
            "institution_acronym",
            "institution",
        ),
        "note_synthese_support": build_note_support(df, stats),
    }


def _excel_safe_scalar(value: Any) -> Any:
    if isinstance(value, (datetime, pd.Timestamp)):
        try:
            if pd.isna(value):
                return ""
        except Exception:
            pass
        if getattr(value, "tzinfo", None) is not None:
            value = value.tz_convert(None) if hasattr(value, "tz_convert") else value.astimezone(timezone.utc).replace(tzinfo=None)
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, (list, dict)):
        return json.dumps(value, ensure_ascii=False)
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    return value


def dataframe_to_xlsx_bytes(sheets: dict[str, pd.DataFrame]) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for sheet_name, df in sheets.items():
            safe_name = re.sub(r"[:\\/*?\[\]]", "_", sheet_name)[:31] or "sheet"
            safe_df = df.copy().astype(object)
            for col in safe_df.columns:
                safe_df[col] = safe_df[col].map(_excel_safe_scalar)
            safe_df.to_excel(writer, index=False, sheet_name=safe_name)
    output.seek(0)
    return output.getvalue()


def build_report_docx_bytes(df: pd.DataFrame, stats: dict[str, int], trace_df: pd.DataFrame) -> bytes:
    try:
        from docx import Document
    except Exception:
        return b""
    doc = Document()
    doc.add_heading("Note de synthèse des consultations / validation", level=1)
    doc.add_paragraph(f"Version du tableau de bord : {APP_VERSION}")
    doc.add_paragraph(f"Date de génération : {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    doc.add_heading("1. Base utilisée", level=2)
    for line in [
        f"Lignes source : {stats['source_rows']}",
        f"Tests supprimés : {stats['tests_removed']}",
        f"Doublons supprimés : {stats['duplicates_removed']}",
        f"Lignes finales conservées : {stats['final_rows']}",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("2. Validation globale", level=2)
    if "overall_validation" in df.columns and not df.empty:
        vc = df["overall_validation"].fillna("NA").astype(str).value_counts()
        for k, v in vc.items():
            doc.add_paragraph(f"{k} : {v}", style="List Bullet")
    else:
        doc.add_paragraph("Aucune donnée disponible.")
    doc.add_heading("3. Position finale des institutions", level=2)
    if "final_institutional_position" in df.columns and not df.empty:
        vc = df["final_institutional_position"].fillna("NA").astype(str).value_counts()
        for k, v in vc.items():
            doc.add_paragraph(f"{k} : {v}", style="List Bullet")
    else:
        doc.add_paragraph("Aucune donnée disponible.")
    doc.add_heading("4. Commentaires et arbitrages", level=2)
    if not trace_df.empty:
        excerpt = trace_df.head(20)
        for _, row in excerpt.iterrows():
            doc.add_paragraph(
                f"{row.get('section', '')} | {row.get('item_label', '')} | {row.get('institution_acronym', '')} | {row.get('commentaire', '')}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("Aucun commentaire traçable disponible.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def make_zip(files: dict[str, bytes]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for name, content in files.items():
            zf.writestr(name, content)
    buf.seek(0)
    return buf.getvalue()


def main() -> None:
    require_password()
    st.title("Dashboard Superadmin")
    st.caption(f"Génération d’un classeur d’analyse et d’un rapport à partir des soumissions finales | version {APP_VERSION}")

    with st.spinner("Chargement des soumissions finales..."):
        branch, records, paths = load_submissions()

    raw_df = records_to_dataframe(records)
    cleaned_df, stats, journal = clean_records_df(raw_df)

    st.success(f"Soumissions chargées depuis GitHub / branche : {branch}")

    c0, c1, c2, c3 = st.columns(4)
    c0.metric("Lignes source", stats["source_rows"])
    c1.metric("Tests supprimés", stats["tests_removed"])
    c2.metric("Doublons supprimés", stats["duplicates_removed"])
    c3.metric("Lignes finales conservées", stats["final_rows"])

    if cleaned_df.empty:
        st.warning("Aucune soumission finale n’est disponible après nettoyage.")
        return

    col1, col2, col3 = st.columns(3)
    col1.metric("Soumissions finales retenues", len(cleaned_df))
    col2.metric("Institutions distinctes", cleaned_df.get("institution_acronym", pd.Series(dtype=str)).replace("", pd.NA).dropna().nunique())
    col3.metric("Pays / CER distincts", cleaned_df.get("country_or_rec", pd.Series(dtype=str)).replace("", pd.NA).dropna().nunique())

    with st.expander("Aperçu de la base nettoyée", expanded=False):
        st.dataframe(cleaned_df, use_container_width=True, hide_index=True)

    with st.expander("Journal de nettoyage", expanded=False):
        if journal.empty:
            st.info("Aucune ligne retirée.")
        else:
            st.dataframe(journal, use_container_width=True, hide_index=True)

    sheets = build_analysis_sheets(cleaned_df, stats, journal)
    trace_df = sheets["traceabilite_commentaires"]
    workbook_bytes = dataframe_to_xlsx_bytes(sheets)
    report_bytes = build_report_docx_bytes(cleaned_df, stats, trace_df)
    now = datetime.utcnow().strftime("%Y%m%d_%H%M%S")

    st.subheader("Fichiers générés")
    c1, c2, c3 = st.columns(3)
    c1.download_button(
        "Télécharger le classeur d’analyse (.xlsx)",
        data=workbook_bytes,
        file_name=f"classeur_analyse_validation_{now}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    c2.download_button(
        "Télécharger le rapport (.docx)",
        data=report_bytes,
        file_name=f"rapport_validation_{now}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        disabled=(report_bytes == b""),
    )
    bundle_bytes = make_zip(
        {
            f"classeur_analyse_validation_{now}.xlsx": workbook_bytes,
            f"rapport_validation_{now}.docx": report_bytes if report_bytes else b"",
        }
    )
    c3.download_button(
        "Télécharger le paquet complet (.zip)",
        data=bundle_bytes,
        file_name=f"package_validation_{now}.zip",
        mime="application/zip",
    )

    tab1, tab2, tab3, tab4 = st.tabs(["Classeur - aperçu", "Traçabilité", "Rapport - contenu synthétique", "Paramètres"])
    with tab1:
        sheet_name = st.selectbox("Feuille du classeur", list(sheets.keys()))
        st.dataframe(sheets[sheet_name], use_container_width=True, hide_index=True)
    with tab2:
        st.dataframe(trace_df, use_container_width=True, hide_index=True)
    with tab3:
        st.markdown("Le rapport DOCX généré comprend une vue d’ensemble de la base retenue, la validation globale, la position finale des institutions, ainsi qu’un extrait traçable des commentaires.")
    with tab4:
        cfg = get_github_config_from_streamlit()
        st.code("\n".join([
            f"GitHub owner : {cfg['owner']}",
            f"GitHub repo  : {cfg['repo']}",
            f"GitHub branch: {cfg['branch']}",
            f"Fichiers JSON détectés : {len(paths)}",
        ]))
        with st.expander("Chemins GitHub détectés"):
            st.write(paths)
        st.info("Le tableau de bord génère les fichiers à la demande. Il ne pousse pas automatiquement le classeur ni le rapport dans GitHub.")


if __name__ == "__main__":
    main()
