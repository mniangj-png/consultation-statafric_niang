
import io
import json
import os
import re
import sqlite3
import zipfile
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, List, Tuple, Optional

import numpy as np
import pandas as pd
import streamlit as st

# Optional deps (Google Sheets / Dropbox)
try:
    import gspread  # type: ignore
    from google.oauth2.service_account import Credentials  # type: ignore
except Exception:
    gspread = None
    Credentials = None

try:
    import dropbox  # type: ignore
except Exception:
    dropbox = None


# =========================
# Configuration
# =========================

APP_TITLE_FR = "Questionnaire de consultation"
APP_TITLE_EN = "Consultation questionnaire"

DB_PATH = "responses.db"
LONG_LIST_CSV = os.path.join("data", "indicator_longlist.csv")
LONG_LIST_XLSX = os.path.join("data", "longlist.xlsx")

UK_FR = "NSP (Ne sais pas)"
UK_EN = "NSP (Don’t know)"


# Scores affichés (notation multicritères)
SCORE_LABELS_FR = {0: "NSP", 1: "Faible", 2: "Moyen", 3: "Élevé"}
SCORE_LABELS_EN = {0: "NSP", 1: "Low", 2: "Medium", 3: "High"}

def score_format(lang: str):
    """Formatter for score selectboxes.
    We include a None option (placeholder) so we don't prefill answers.
    """
    placeholder_fr = "— Sélectionner —"
    placeholder_en = "— Select —"
    def _fmt(v):
        if v is None or v == "":
            return placeholder_fr if lang == "fr" else placeholder_en
        try:
            iv = int(v)
        except Exception:
            return str(v)
        return SCORE_LABELS_FR.get(iv, str(v)) if lang == "fr" else SCORE_LABELS_EN.get(iv, str(v))
    return _fmt


ROLE_OPTIONS_FR = [
    "DG/DGA/SG",
    "Directeur",
    "Conseiller",
    "Chef de division",
    "Chef de bureau",
    "Autre",
]
ROLE_OPTIONS_EN = [
    "DG/DGA/SG",
    "Director",
    "Advisor",
    "Head of division",
    "Head of office",
    "Other",
]


# =========================
# Helpers : i18n and state
# =========================

def t(lang: str, fr: str, en: str) -> str:
    return fr if lang == "fr" else en


def now_utc_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def get_query_params() -> Dict[str, List[str]]:
    """Compatibility across Streamlit versions."""
    try:
        # Streamlit >= 1.30
        qp = st.query_params  # type: ignore
        return {k: list(v) if isinstance(v, (list, tuple)) else [str(v)] for k, v in qp.items()}
    except Exception:
        try:
            return st.experimental_get_query_params()
        except Exception:
            return {}


def set_query_params(params: Dict[str, Any]) -> None:
    try:
        st.query_params.update(params)  # type: ignore
    except Exception:
        try:
            st.experimental_set_query_params(**params)
        except Exception:
            pass


def init_session() -> None:
    if "lang" not in st.session_state:
        st.session_state.lang = "fr"
    if "nav_idx" not in st.session_state:
        st.session_state.nav_idx = 0
    if "responses" not in st.session_state:
        st.session_state.responses = {}
    if "submission_id" not in st.session_state:
        st.session_state.submission_id = None
    if "admin_authed" not in st.session_state:
        st.session_state.admin_authed = False


def resp_get(key: str, default=None):
    return st.session_state.responses.get(key, default)


def resp_set(key: str, value) -> None:
    st.session_state.responses[key] = value


# =========================
# Data : longlist loader
# =========================

@st.cache_data(show_spinner=False)
def load_longlist() -> pd.DataFrame:
    """
    Charge la longlist (statistiques par domaine) depuis :
    - CSV : data/indicator_longlist.csv (prioritaire)
    - XLSX : data/longlist.xlsx (fallback)
    Tolère aussi les fichiers placés à la racine du dépôt.

    Si aucun fichier n'est trouvé, l'application démarre quand même,
    mais les listes déroulantes de la Rubrique 4/5 seront vides.
    """
    csv_candidates = [
        LONG_LIST_CSV,
        "indicator_longlist.csv",
        os.path.join(".", "indicator_longlist.csv"),
        os.path.join(".", "data", "indicator_longlist.csv"),
    ]
    xlsx_candidates = [
        LONG_LIST_XLSX,
        "longlist.xlsx",
        os.path.join(".", "longlist.xlsx"),
        os.path.join(".", "data", "longlist.xlsx"),
    ]

    # 1) CSV (prioritaire si la traduction EN est suffisamment complète)
    df_csv = None
    df_csv_path = None
    for p in csv_candidates:
        if os.path.exists(p):
            df_csv = pd.read_csv(p, dtype=str).fillna("")
            df_csv_path = p
            break

    if df_csv is not None:
        # Sanity check : si beaucoup de libellés EN sont vides, on préfère l'XLSX (souvent plus à jour)
        if "stat_label_en" in df_csv.columns:
            miss_ratio = (df_csv["stat_label_en"].astype(str).str.strip() == "").mean()
        else:
            miss_ratio = 1.0

        if miss_ratio <= 0.20:
            df_csv.attrs["source_path"] = df_csv_path
            return df_csv

    # 2) XLSX (format utilisateur) (format utilisateur)
    for p in xlsx_candidates:
        if os.path.exists(p):
            df = pd.read_excel(p, dtype=str).fillna("")
            df.attrs["source_path"] = p

            # Colonnes attendues (minimum) : Domain_code, Domain_label_fr, Stat_label_fr
            if set(["Domain_code", "Domain_label_fr", "Stat_label_fr"]).issubset(df.columns):
                out = pd.DataFrame()
                out["domain_code"] = df["Domain_code"].astype(str).str.strip()

                # Labels FR (on retire le préfixe code "D01|...")
                out["domain_label_fr"] = df["Domain_label_fr"].astype(str).str.split("|", n=1).str[-1].str.strip()
                out["stat_code"] = df["Stat_label_fr"].astype(str).str.split("|", n=1).str[0].str.strip()
                out["stat_label_fr"] = df["Stat_label_fr"].astype(str).str.split("|", n=1).str[-1].str.strip()

                # Labels EN si disponibles, sinon fallback FR
                if "Domain_label_en" in df.columns:
                    out["domain_label_en"] = df["Domain_label_en"].astype(str).str.split("|", n=1).str[-1].str.strip()
                else:
                    out["domain_label_en"] = out["domain_label_fr"]

                if "Stat_label_en" in df.columns:
                    out["stat_label_en"] = df["Stat_label_en"].astype(str).str.split("|", n=1).str[-1].str.strip()
                else:
                    out["stat_label_en"] = out["stat_label_fr"]

                out.attrs["source_path"] = p
                return out[[
                    "domain_code",
                    "domain_label_fr",
                    "domain_label_en",
                    "stat_code",
                    "stat_label_fr",
                    "stat_label_en",
                ]]


    # Fallback final : si un CSV a été trouvé (même avec traduction EN incomplète), on le renvoie
    if df_csv is not None:
        df_csv.attrs["source_path"] = df_csv_path or ""
        return df_csv
# Aucun fichier trouvé : dataframe vide
    empty = pd.DataFrame(columns=[
        "domain_code",
        "domain_label_fr",
        "domain_label_en",
        "stat_code",
        "stat_label_fr",
        "stat_label_en",
    ])
    empty.attrs["source_path"] = ""
    return empty


def domains_from_longlist(df_long: pd.DataFrame, lang: str) -> List[Tuple[str, str]]:
    if df_long.empty:
        return []
    col = "domain_label_fr" if lang == "fr" else "domain_label_en"
    tmp = df_long[["domain_code", col]].drop_duplicates().sort_values(["domain_code", col])
    return [(r["domain_code"], r[col]) for _, r in tmp.iterrows()]


def stats_for_domain(df_long: pd.DataFrame, domain_code: str, lang: str) -> List[Tuple[str, str]]:
    if df_long.empty or not domain_code:
        return []
    col = "stat_label_fr" if lang == "fr" else "stat_label_en"
    tmp = df_long[df_long["domain_code"] == domain_code][["stat_code", col]].drop_duplicates()
    tmp = tmp.sort_values(["stat_code", col])
    return [(r["stat_code"], (str(r[col]).strip() if str(r[col]).strip() else r["stat_code"])) for _, r in tmp.iterrows()]


# =========================
# Storage : SQLite + optional Google Sheets + Dropbox
# =========================


def domain_label_map(df_long: pd.DataFrame, lang: str) -> Dict[str, str]:
    """Map domain_code -> label in selected language."""
    if df_long is None or df_long.empty:
        return {}
    col = "domain_label_fr" if lang == "fr" else "domain_label_en"
    m = {}
    for _, r in df_long.drop_duplicates("domain_code").iterrows():
        code = str(r["domain_code"])
        lbl = str(r.get(col, "")).strip()
        if not lbl:
            lbl = str(r.get("domain_label_fr", "")).strip() or code
        m[code] = lbl
    return m

def stat_label_map(df_long: pd.DataFrame, lang: str) -> Dict[str, str]:
    """Map stat_code -> label in selected language."""
    if df_long is None or df_long.empty:
        return {}
    col = "stat_label_fr" if lang == "fr" else "stat_label_en"
    m = {}
    for _, r in df_long.drop_duplicates("stat_code").iterrows():
        code = str(r["stat_code"])
        lbl = str(r.get(col, "")).strip()
        if not lbl:
            lbl = str(r.get("stat_label_fr", "")).strip() or code
        m[code] = lbl
    return m

def build_publication_report_docx(lang: str, filtered_payloads: pd.DataFrame, by_domain: pd.DataFrame, by_stat: pd.DataFrame, scored_rows: pd.DataFrame) -> bytes:
    """
    Génère un rapport Word 'publication' enrichi :
    - profil des répondants
    - domaines TOP 5 (fréquences)
    - tableau agrégé des statistiques et scores moyens
    - graphiques (bar charts)
    - annexes
    """
    from docx import Document
    from docx.shared import Inches
    import matplotlib.pyplot as plt

    doc = Document()
    title = t(lang, "Rapport de synthèse – Consultation sur les statistiques prioritaires", "Summary report – Consultation on priority statistics")
    doc.add_heading(title, level=0)
    doc.add_paragraph(t(lang, f"Date : {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", f"Date: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"))
    doc.add_paragraph("")

    # Sample profile
    doc.add_heading(t(lang, "Profil des répondants", "Respondent profile"), level=1)
    n = len(filtered_payloads)
    doc.add_paragraph(t(lang, f"Nombre de réponses analysées : {n}", f"Number of responses analyzed: {n}"))

    # Countries
    if "pays" in filtered_payloads.columns:
        vc = filtered_payloads["pays"].fillna("").astype(str)
        vc = vc[vc.str.strip() != ""].value_counts().head(10)
        if len(vc) > 0:
            doc.add_paragraph(t(lang, "Top pays (10 premiers) :", "Top countries (top 10):"))
            for k, v in vc.items():
                doc.add_paragraph(f"- {k} : {v}", style=None)

    # Actor types
    if "type_acteur" in filtered_payloads.columns:
        vc = filtered_payloads["type_acteur"].fillna("").astype(str)
        vc = vc[vc.str.strip() != ""].value_counts()
        if len(vc) > 0:
            doc.add_paragraph(t(lang, "Répartition par type d’acteur :", "Distribution by stakeholder type:"))
            for k, v in vc.items():
                doc.add_paragraph(f"- {k} : {v}", style=None)

    # Domain aggregation
    doc.add_heading(t(lang, "Domaines prioritaires (scores moyens)", "Priority domains (mean scores)"), level=1)
    top_dom = by_domain.head(15).copy()
    # Table
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = t(lang, "Domaine", "Domain")
    hdr[1].text = t(lang, "Nb. soumissions", "Submissions")
    hdr[2].text = t(lang, "Nb. stats notées", "Scored indicators")
    hdr[3].text = t(lang, "Score moyen", "Mean score")
    for _, r in top_dom.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["domain_label"])
        row[1].text = str(int(r["n_submissions"]))
        row[2].text = str(int(r["n_stats"]))
        row[3].text = f"{float(r['mean_overall']):.2f}"

    # Chart domain
    try:
        fig = plt.figure()
        plt.bar(top_dom["domain_label"], top_dom["mean_overall"])
        plt.xticks(rotation=75, ha="right")
        plt.ylabel(t(lang, "Score moyen", "Mean score"))
        plt.tight_layout()
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format="png", dpi=150)
        plt.close(fig)
        img_stream.seek(0)
        doc.add_paragraph(t(lang, "Graphique : score moyen par domaine (top 15).", "Chart: mean score by domain (top 15)."))
        doc.add_picture(img_stream, width=Inches(6.5))
    except Exception:
        pass

    # Statistic aggregation
    doc.add_heading(t(lang, "Statistiques prioritaires (scores moyens)", "Priority indicators (mean scores)"), level=1)
    top_stat = by_stat.sort_values(["mean_overall", "n"], ascending=[False, False]).head(30).copy()
    table2 = doc.add_table(rows=1, cols=6)
    h = table2.rows[0].cells
    h[0].text = t(lang, "Domaine", "Domain")
    h[1].text = t(lang, "Statistique", "Indicator")
    h[2].text = t(lang, "N", "N")
    h[3].text = t(lang, "Écart", "Gap")
    h[4].text = t(lang, "Demande", "Demand")
    h[5].text = t(lang, "Faisabilité", "Feasibility")
    for _, r in top_stat.iterrows():
        row = table2.add_row().cells
        row[0].text = str(r["domain_label"])
        row[1].text = str(r["stat_label"])
        row[2].text = str(int(r["n"]))
        row[3].text = f"{float(r['mean_gap']):.2f}"
        row[4].text = f"{float(r['mean_demand']):.2f}"
        row[5].text = f"{float(r['mean_feasibility']):.2f}"

    # Chart top overall indicators
    try:
        fig = plt.figure()
        plt.barh(top_stat["stat_label"].iloc[::-1], top_stat["mean_overall"].iloc[::-1])
        plt.xlabel(t(lang, "Score moyen", "Mean score"))
        plt.tight_layout()
        img2 = io.BytesIO()
        plt.savefig(img2, format="png", dpi=150)
        plt.close(fig)
        img2.seek(0)
        doc.add_paragraph(t(lang, "Graphique : score moyen par statistique (top 30).", "Chart: mean score by indicator (top 30)."))
        doc.add_picture(img2, width=Inches(6.5))
    except Exception:
        pass

    # Interpretation auto
    doc.add_heading(t(lang, "Interprétations automatiques", "Automatic interpretations"), level=1)
    # Simple rules
    best_dom = top_dom.iloc[0]
    doc.add_paragraph(
        t(
            lang,
            f"Le domaine le mieux noté est « {best_dom['domain_label']} » avec un score moyen de {best_dom['mean_overall']:.2f} (sur 3).",
            f"The highest-rated domain is “{best_dom['domain_label']}” with a mean score of {best_dom['mean_overall']:.2f} (out of 3)."
        )
    )
    best_stat = top_stat.iloc[0]
    doc.add_paragraph(
        t(
            lang,
            f"La statistique la mieux notée est « {best_stat['stat_label']} » (domaine : {best_stat['domain_label']}) avec un score moyen de {best_stat['mean_overall']:.2f}.",
            f"The highest-rated indicator is “{best_stat['stat_label']}” (domain: {best_stat['domain_label']}) with a mean score of {best_stat['mean_overall']:.2f}."
        )
    )

    # Annexes
    doc.add_heading(t(lang, "Annexes", "Annexes"), level=1)
    doc.add_paragraph(t(lang, "A1. Tableau détaillé (statistiques agrégées) – extrait", "A1. Detailed table (aggregated indicators) – excerpt"))
    annex = by_stat.head(50).copy()
    tab3 = doc.add_table(rows=1, cols=5)
    hh = tab3.rows[0].cells
    hh[0].text = t(lang, "Domaine", "Domain")
    hh[1].text = t(lang, "Statistique", "Indicator")
    hh[2].text = t(lang, "N", "N")
    hh[3].text = t(lang, "Score moyen", "Mean score")
    hh[4].text = t(lang, "Détail", "Detail")
    for _, r in annex.iterrows():
        rr = tab3.add_row().cells
        rr[0].text = str(r["domain_label"])
        rr[1].text = str(r["stat_label"])
        rr[2].text = str(int(r["n"]))
        rr[3].text = f"{float(r['mean_overall']):.2f}"
        rr[4].text = f"Gap={float(r['mean_gap']):.2f}, Demand={float(r['mean_demand']):.2f}, Feas={float(r['mean_feasibility']):.2f}"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()




def db_init() -> None:
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS submissions(
            submission_id TEXT PRIMARY KEY,
            submitted_at_utc TEXT,
            lang TEXT,
            email TEXT,
            payload_json TEXT
        )
    """)

    # Backward compatibility : add email column if existing DB was created with older schema
    try:
        cur.execute("PRAGMA table_info(submissions)")
        cols = [r[1] for r in cur.fetchall()]
        if "email" not in cols:
            cur.execute("ALTER TABLE submissions ADD COLUMN email TEXT")
    except Exception:
        pass

    # Helpful index (non-unique) for email lookups
    try:
        cur.execute("CREATE INDEX IF NOT EXISTS idx_submissions_email ON submissions(email)")
    except Exception:
        pass

    con.commit()
    con.close()


def db_email_exists(email: str) -> bool:
    email = (email or "").strip().lower()
    if not email or not os.path.exists(DB_PATH):
        return False
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT 1 FROM submissions WHERE lower(email)=? LIMIT 1", (email,))
    row = cur.fetchone()
    con.close()
    return row is not None


def db_save_submission(submission_id: str, lang: str, email: str, payload: Dict[str, Any]) -> None:
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
        INSERT OR REPLACE INTO submissions(submission_id, submitted_at_utc, lang, email, payload_json)
        VALUES(?, ?, ?, ?, ?)
    """, (submission_id, now_utc_iso(), lang, (email or "").strip().lower(), json.dumps(payload, ensure_ascii=False)))
    con.commit()
    con.close()

def db_read_submissions(limit: int = 2000) -> pd.DataFrame:
    if not os.path.exists(DB_PATH):
        return pd.DataFrame(columns=["submission_id", "submitted_at_utc", "lang", "email", "payload_json"])
    con = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query(
        "SELECT submission_id, submitted_at_utc, lang, email, payload_json FROM submissions ORDER BY submitted_at_utc DESC LIMIT ?",
        con,
        params=(limit,),
    )
    con.close()
    return df


def db_dump_csv_bytes(limit: int = 2000000) -> bytes:
    """Export the SQLite submissions table to CSV bytes."""
    df = db_read_submissions(limit=limit)
    return df.to_csv(index=False).encode("utf-8-sig")


def flatten_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Create a 'flat' row for exports / Google Sheets (comprehensive).
    - Keeps keys stable across FR/EN by mapping table items to canonical ids.
    - Serializes list/dict fields into '; ' / JSON strings as needed.
    """
    def _join_list(v: Any) -> str:
        if isinstance(v, list):
            return "; ".join([str(x) for x in v if x is not None and str(x).strip() != ""])
        return ""

    def _json(v: Any) -> str:
        try:
            return json.dumps(v, ensure_ascii=False)
        except Exception:
            return ""

    # Canonical mappings for table questions (FR/EN)
    GENDER_ITEM_MAP = {
        "Sexe": "sex",
        "Sex": "sex",
        "Âge": "age",
        "Age": "age",
        "Milieu urbain/rural": "urban_rural",
        "Urban/rural residence": "urban_rural",
        "Handicap": "disability",
        "Disability": "disability",
        "Quintile de richesse": "wealth_quintile",
        "Wealth quintile": "wealth_quintile",
    }
    CAPACITY_ITEM_MAP = {
        "Compétences (RH)": "skills_hr",
        "Human resources skills": "skills_hr",
        "Accès aux données administratives": "access_admin_data",
        "Access to administrative data": "access_admin_data",
        "Financement": "funding",
        "Funding": "funding",
        "Outils numériques": "digital_tools",
        "Digital tools": "digital_tools",
        "Cadre juridique": "legal_framework",
        "Legal framework": "legal_framework",
        "Coordination institutionnelle": "institutional_coordination",
        "Institutional coordination": "institutional_coordination",
    }

    def _extract_table(table_obj: Any, mapping: Dict[str, str], prefix: str) -> Dict[str, Any]:
        out_tbl: Dict[str, Any] = {}
        # Ensure stable columns even when a respondent skips the section
        canons = sorted(set(mapping.values()))
        for canon in canons:
            out_tbl[f"{prefix}_{canon}"] = ""
            out_tbl[f"{prefix}_{canon}_spec"] = ""
        if not isinstance(table_obj, dict):
            return out_tbl
        for label, canon in mapping.items():
            cell = table_obj.get(label, None)
            if isinstance(cell, dict):
                out_tbl[f"{prefix}_{canon}"] = cell.get("code", "")
                out_tbl[f"{prefix}_{canon}_spec"] = cell.get("spec", "")
        return out_tbl

    out: Dict[str, Any] = {}

    # Identification (Rubrique 2)
    out["organisation"] = payload.get("organisation", "")
    out["pays"] = payload.get("pays", "")
    out["type_acteur"] = payload.get("type_acteur", "")
    out["fonction"] = payload.get("fonction", "")
    out["email"] = payload.get("email", "")
    out["lang"] = payload.get("lang", "")

    # Rubrique 3 : portée
    out["scope"] = payload.get("scope", "")
    out["scope_other"] = payload.get("scope_other", "")

    # Rubrique 4 : domaines
    pre = payload.get("preselection_domains", [])
    out["preselection_domains"] = _join_list(pre)
    out["nb_preselection_domains"] = len(pre) if isinstance(pre, list) else 0

    top5 = payload.get("top5_domains", [])
    for i in range(5):
        out[f"top_domain_{i+1}"] = top5[i] if i < len(top5) else ""

    # Rubrique 5 : stats et notation
    selected_stats = payload.get("selected_stats", [])
    out["nb_stats"] = len(selected_stats) if isinstance(selected_stats, list) else 0
    out["stats_list"] = _join_list(selected_stats)
    out["selected_by_domain_json"] = _json(payload.get("selected_by_domain", {}))
    out["scoring_json"] = _json(payload.get("scoring", {}))

    # Rubrique 6 : perspective de genre (table)
    out.update(_extract_table(payload.get("gender_table", {}), GENDER_ITEM_MAP, "gender"))

    # Rubrique 8 : capacité & faisabilité (table)
    out.update(_extract_table(payload.get("capacity_table", {}), CAPACITY_ITEM_MAP, "capacity"))

    # Rubrique 9 : harmonisation & qualité
    out["quality_expectations"] = _join_list(payload.get("quality_expectations", []))
    out["quality_other"] = payload.get("quality_other", "")

    # Rubrique 10 : diffusion
    out["dissemination_channels"] = _join_list(payload.get("dissemination_channels", []))
    out["dissemination_other"] = payload.get("dissemination_other", "")

    # Rubrique 12 : questions ouvertes
    out["comment_1"] = payload.get("open_q1", "")
    out["missing_indicators"] = payload.get("open_q2", "")
    out["support_needs"] = payload.get("open_q3", "")

    return out


def google_sheets_append(payload: Dict[str, Any]) -> Tuple[bool, str]:
    """
    Append a row into a Google Sheet if configured.
    Requires secrets:
      GOOGLE_SHEET_ID
      GOOGLE_SERVICE_ACCOUNT (dict)
    """
    if gspread is None or Credentials is None:
        return False, "Bibliothèques Google Sheets non disponibles (gspread/google-auth)."
    try:
        sheet_id = st.secrets.get("GOOGLE_SHEET_ID", None)
        sa_info = st.secrets.get("GOOGLE_SERVICE_ACCOUNT", None)
        if not sheet_id or not sa_info:
            return False, "Google Sheets non configuré (secrets manquants)."
        scopes = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
        creds = Credentials.from_service_account_info(sa_info, scopes=scopes)
        gc = gspread.authorize(creds)
        sh = gc.open_by_key(sheet_id)
        try:
            ws = sh.worksheet("responses")
        except Exception:
            ws = sh.add_worksheet(title="responses", rows=2000, cols=80)

        row = flatten_payload(payload)
        # Ensure header
        existing = ws.get_all_values()
        if not existing:
            ws.append_row(list(row.keys()), value_input_option="RAW")
        else:
            header = existing[0]
            # Add any missing columns at end
            for k in row.keys():
                if k not in header:
                    header.append(k)
            # If header grew, update first row
            ws.update("A1", [header])
        # Align order with header
        header = ws.row_values(1)
        values = [row.get(h, "") for h in header]
        ws.append_row(values, value_input_option="RAW")
        return True, "OK"
    except Exception as e:
        return False, f"Erreur Google Sheets : {e}"


def google_sheets_write_df(df: pd.DataFrame, worksheet_title: str, sheet_id: Optional[str] = None) -> Tuple[bool, str]:
    """
    Write a dataframe to a worksheet (overwrite), if configured.
    Uses secrets:
      GOOGLE_SHEET_ID (or provided sheet_id)
      GOOGLE_SERVICE_ACCOUNT (dict)
    """
    if gspread is None or Credentials is None:
        return False, "Bibliothèques Google Sheets non disponibles (gspread/google-auth)."
    try:
        sid = sheet_id or st.secrets.get("GOOGLE_SHEET_ID", None)
        sa_info = st.secrets.get("GOOGLE_SERVICE_ACCOUNT", None)
        if not sid or not sa_info:
            return False, "Google Sheets non configuré (secrets manquants)."

        scopes = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
        creds = Credentials.from_service_account_info(sa_info, scopes=scopes)
        gc = gspread.authorize(creds)
        sh = gc.open_by_key(sid)

        try:
            ws = sh.worksheet(worksheet_title)
        except Exception:
            ws = sh.add_worksheet(title=worksheet_title, rows=max(200, len(df) + 20), cols=max(20, len(df.columns) + 10))

        # Prepare values (header + rows), keep everything as string to avoid type surprises
        values = [list(df.columns)]
        if not df.empty:
            values += df.astype(object).where(pd.notnull(df), "").astype(str).values.tolist()

        ws.clear()
        ws.update("A1", values, value_input_option="RAW")
        return True, "OK"
    except Exception as e:
        return False, f"Erreur Google Sheets : {e}"


def dropbox_upload_bytes(content: bytes, filename: str, subfolder: str = "exports") -> Tuple[bool, str]:
    """
    Upload arbitrary bytes to Dropbox, if configured.
    Requires secret: DROPBOX_ACCESS_TOKEN
    Optional: DROPBOX_FOLDER (default /consultation_stat_niang)
    """
    if dropbox is None:
        return False, "Bibliothèque Dropbox non disponible."
    try:
        token = st.secrets.get("DROPBOX_ACCESS_TOKEN", None)
        if not token:
            return False, "Dropbox non configuré (DROPBOX_ACCESS_TOKEN manquant)."
        folder = st.secrets.get("DROPBOX_FOLDER", "/consultation_stat_niang")
        folder = folder if folder.startswith("/") else "/" + folder
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        path = f"{folder}/{subfolder}/{ts}_{filename}"
        dbx = dropbox.Dropbox(token)
        dbx.files_upload(content, path, mode=dropbox.files.WriteMode.overwrite)
        return True, "OK"
    except Exception as e:
        return False, f"Erreur Dropbox : {e}"

def dropbox_upload_json(submission_id: str, payload: Dict[str, Any]) -> Tuple[bool, str]:
    """
    Upload the JSON submission to Dropbox if configured.
    Requires secret: DROPBOX_ACCESS_TOKEN
    Optional: DROPBOX_FOLDER (default /consultation_stat_niang)
    """
    if dropbox is None:
        return False, "Bibliothèque Dropbox non disponible."
    try:
        token = st.secrets.get("DROPBOX_ACCESS_TOKEN", None)
        if not token:
            return False, "Dropbox non configuré (DROPBOX_ACCESS_TOKEN manquant)."
        folder = st.secrets.get("DROPBOX_FOLDER", "/consultation_stat_niang")
        folder = folder if folder.startswith("/") else "/" + folder
        path = f"{folder}/submissions/{submission_id}.json"
        dbx = dropbox.Dropbox(token)
        content = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
        dbx.files_upload(content, path, mode=dropbox.files.WriteMode.overwrite)
        return True, "OK"
    except Exception as e:
        return False, f"Erreur Dropbox : {e}"


# =========================
# Validation logic (quality controls)
# =========================

EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


def validate_r2(lang: str) -> List[str]:
    errs: List[str] = []
    organisation = resp_get("organisation", "").strip()
    pays = resp_get("pays", "").strip()
    type_acteur = resp_get("type_acteur", "").strip()
    fonction = resp_get("fonction", "").strip()
    fonction_autre = resp_get("fonction_autre", "").strip()
    email = resp_get("email", "").strip()

    if not organisation:
        errs.append(t(lang, "Organisation : champ obligatoire.", "Organization: required field."))
    if not pays:
        errs.append(t(lang, "Pays : champ obligatoire.", "Country: required field."))
    if not type_acteur:
        errs.append(t(lang, "Type d’acteur : champ obligatoire.", "Stakeholder type: required field."))
    if not fonction:
        errs.append(t(lang, "Fonction : champ obligatoire.", "Role/Function: required field."))
    if fonction == "Autre" or fonction == "Other":
        if not fonction_autre:
            errs.append(t(lang, "Fonction (Autre) : précisez.", "Role (Other): please specify."))
    if not email:
        errs.append(t(lang, "Email : champ obligatoire.", "Email: required field."))
    elif not EMAIL_RE.match(email):
        errs.append(t(lang, "Email : format invalide.", "Email: invalid format."))
    return errs


def validate_r4(lang: str) -> List[str]:
    errs: List[str] = []
    pre = resp_get("preselected_domains", [])
    top5 = resp_get("top5_domains", [])
    if not isinstance(pre, list):
        pre = []
    if not isinstance(top5, list):
        top5 = []
    if len(pre) < 5 or len(pre) > 10:
        errs.append(t(lang, "Rubrique 4 : pré-sélectionnez entre 5 et 10 domaines.", "Section 4: pre-select 5 to 10 domains."))
    if len(set(pre)) != len(pre):
        errs.append(t(lang, "Rubrique 4 : la pré-sélection contient des doublons.", "Section 4: duplicates found in pre-selection."))
    if len(top5) != 5:
        errs.append(t(lang, "Rubrique 4 : le TOP 5 doit contenir exactement 5 domaines.", "Section 4: TOP 5 must contain exactly 5 domains."))
    else:
        if len(set(top5)) != 5:
            errs.append(t(lang, "Rubrique 4 : le TOP 5 contient des doublons.", "Section 4: TOP 5 contains duplicates."))
        missing = [d for d in top5 if d not in pre]
        if missing:
            errs.append(t(lang, "Rubrique 4 : chaque domaine du TOP 5 doit provenir de la pré-sélection.", "Section 4: TOP 5 must be selected from pre-selection."))
    return errs


def validate_r5(lang: str) -> List[str]:
    errs: List[str] = []
    top5 = resp_get("top5_domains", [])
    selected_by_domain: Dict[str, List[str]] = resp_get("selected_by_domain", {})
    if not isinstance(selected_by_domain, dict):
        selected_by_domain = {}

    all_stats: List[str] = []
    for d in top5:
        stats = selected_by_domain.get(d, [])
        if not isinstance(stats, list):
            stats = []
        if len(stats) < 1:
            errs.append(t(lang, f"Rubrique 5 : choisissez au moins 1 statistique pour {d}.",
                          f"Section 5: select at least 1 indicator for {d}."))
        if len(stats) > 3:
            errs.append(t(lang, f"Rubrique 5 : maximum 3 statistiques pour {d}.",
                          f"Section 5: maximum 3 indicators for {d}."))
        all_stats.extend(stats)

    if len(all_stats) < 5 or len(all_stats) > 15:
        errs.append(t(lang, "Rubrique 5 : le total des statistiques doit être entre 5 et 15.",
                      "Section 5: total number of indicators must be between 5 and 15."))

    if len(set(all_stats)) != len(all_stats):
        errs.append(t(lang, "Rubrique 5 : une même statistique ne doit pas être sélectionnée plusieurs fois.",
                      "Section 5: the same indicator must not be selected more than once."))

    # scoring
    scoring: Dict[str, Dict[str, Any]] = resp_get("scoring", {})
    if not isinstance(scoring, dict):
        scoring = {}

    for s in all_stats:
        if s not in scoring:
            errs.append(t(lang, f"Rubrique 5 : vous devez noter la statistique {s}.",
                          f"Section 5: you must score indicator {s}."))
            continue
        for k in ["gap", "demand", "feasibility"]:
            v_raw = scoring.get(s, {}).get(k, None)
            if v_raw is None or str(v_raw).strip() == "":
                errs.append(t(lang, f"Rubrique 5 : la note '{k}' manque pour {s}.",
                              f"Section 5: missing score '{k}' for {s}."))
            else:
                try:
                    v = int(scoring[s][k])
                    if v < 0 or v > 3:
                        errs.append(t(lang, f"Rubrique 5 : note invalide pour {s} ({k}).",
                                      f"Section 5: invalid score for {s} ({k})."))
                except Exception:
                    errs.append(t(lang, f"Rubrique 5 : note invalide pour {s} ({k}).",
                                  f"Section 5: invalid score for {s} ({k})."))
    return errs


def validate_r6(lang: str) -> List[str]:
    errs: List[str] = []
    tbl = resp_get("gender_table", {})
    if not isinstance(tbl, dict) or not tbl:
        errs.append(t(lang, "Rubrique 6 : veuillez renseigner le tableau.", "Section 6: please complete the table."))
        return errs
    for k, v in tbl.items():
        if not v:
            errs.append(t(lang, f"Rubrique 6 : ligne non renseignée : {k}.", f"Section 6: missing answer for: {k}."))
    return errs


def validate_r8(lang: str) -> List[str]:
    errs: List[str] = []
    tbl = resp_get("capacity_table", {})
    if not isinstance(tbl, dict) or not tbl:
        errs.append(t(lang, "Rubrique 8 : veuillez renseigner le tableau.", "Section 8: please complete the table."))
        return errs
    for k, v in tbl.items():
        if not v:
            errs.append(t(lang, f"Rubrique 8 : ligne non renseignée : {k}.", f"Section 8: missing answer for: {k}."))
    return errs


def validate_all(lang: str) -> List[str]:
    errs = []
    errs.extend(validate_r2(lang))
    errs.extend(validate_r4(lang))
    errs.extend(validate_r5(lang))
    errs.extend(validate_r6(lang))
    errs.extend(validate_r8(lang))
    # Open questions optional: handled in send step as warnings
    return errs


# =========================
# Navigation
# =========================

def get_steps(lang: str) -> List[Tuple[str, str]]:
    # Rubrics 7 and 11 removed, plus final SEND tab
    return [
        ("R1", t(lang, "Rubrique 1 : Instructions", "Section 1: Instructions")),
        ("R2", t(lang, "Rubrique 2 : Identification du répondant", "Section 2: Respondent identification")),
        ("R3", t(lang, "Rubrique 3 : Portée de la réponse", "Section 3: Scope of response")),
        ("R4", t(lang, "Rubrique 4 : Domaines prioritaires", "Section 4: Priority domains")),
        ("R5", t(lang, "Rubrique 5 : Statistiques prioritaires et notation", "Section 5: Priority indicators and scoring")),
        ("R6", t(lang, "Rubrique 6 : Perspective de genre", "Section 6: Gender perspective")),
        ("R8", t(lang, "Rubrique 8 : Capacité et faisabilité (12–24 mois)", "Section 8: Capacity and feasibility (12–24 months)")),
        ("R9", t(lang, "Rubrique 9 : Harmonisation et qualité", "Section 9: Harmonization and quality")),
        ("R10", t(lang, "Rubrique 10 : Diffusion", "Section 10: Dissemination")),
        ("R12", t(lang, "Rubrique 12 : Questions ouvertes", "Section 12: Open questions")),
        ("SEND", t(lang, "ENVOYER", "SUBMIT")),
    ]


def render_sidebar(lang: str, steps: List[Tuple[str, str]]) -> None:
    st.sidebar.header(t(lang, "Navigation", "Navigation"))
    labels = [s[1] for s in steps]

    # Keep sidebar selection in sync with nav_idx
    st.session_state.nav_radio = int(st.session_state.nav_idx)
    chosen = st.sidebar.radio(
        t(lang, "Aller à", "Go to"),
        options=list(range(len(labels))),
        index=int(st.session_state.nav_idx),
        format_func=lambda i: labels[i],
        key="nav_radio"
    )

    # User clicked in sidebar
    if int(chosen) != int(st.session_state.nav_idx):
        st.session_state.nav_idx = int(chosen)

    st.sidebar.divider()
    st.sidebar.caption(
        t(
            lang,
            "Note : les contrôles qualité peuvent bloquer la progression si une contrainte n’est pas respectée.",
            "Note: quality checks may prevent moving forward when constraints are not met."
        )
    )

    st.sidebar.markdown("---")
    st.sidebar.caption(
        t(
            lang,
            "NSP : Ne sait pas (score 0). Utilisez NSP uniquement si l’information est indisponible.",
            "UK: Unknown (score 0). Use UK only when information is unavailable."
        )
    )


def nav_buttons(lang: str, steps: List[Tuple[str, str]], df_long: pd.DataFrame) -> None:
    """Bottom nav buttons, with blocking based on current step validations."""
    step_key = steps[st.session_state.nav_idx][0]
    errors: List[str] = []

    # Blocking rules per step
    if step_key == "R2":
        errors = validate_r2(lang)
    elif step_key == "R4":
        errors = validate_r4(lang)
    elif step_key == "R5":
        errors = validate_r5(lang)
    elif step_key == "R6":
        errors = validate_r6(lang)
    elif step_key == "R8":
        errors = validate_r8(lang)

    col1, col2, col3 = st.columns([1, 1, 3])
    with col1:
        prev_disabled = st.session_state.nav_idx <= 0
        if st.button(t(lang, "⬅ Précédent", "⬅ Previous"), disabled=prev_disabled):
            st.session_state.nav_idx = max(0, st.session_state.nav_idx - 1)
            st.rerun()
    with col2:
        next_disabled = (st.session_state.nav_idx >= len(steps) - 1) or bool(errors)
        if st.button(t(lang, "Suivant ➡", "Next ➡"), disabled=next_disabled):
            st.session_state.nav_idx = min(len(steps) - 1, st.session_state.nav_idx + 1)
            st.rerun()
    with col3:
        if errors:
            st.error("\n".join(errors))


# =========================
# UI : Rubrics
# =========================

def rubric_1(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 1 : Instructions", "Section 1: Instructions"))
    st.markdown(
        t(
            lang,
            """
### Objectif
Ce questionnaire vise à recueillir votre avis sur **les statistiques socio-économiques prioritaires** à produire et diffuser au niveau continental.

### Comment répondre
1. **Identifiez** votre organisation (Rubrique 2).
2. **Pré-sélectionnez 5 à 10 domaines** et classez un **TOP 5** (Rubrique 4).
3. Pour chaque domaine du TOP 5 : choisissez **1 à 3 statistiques** et attribuez des **notes** (Rubrique 5).
4. Complétez les rubriques transversales : **genre** et **capacité/faisabilité**.

### Barème de notation (Rubrique 5)
- **3** : élevé / très important  
- **2** : moyen  
- **1** : faible  
- **0** : NSP (Ne sais pas)

            """,
            """
### Purpose
This questionnaire collects your views on **priority socio-economic statistics** to be produced and disseminated at continental level.

### How to answer
1. **Identify** your organization (Section 2).
2. **Pre-select 5–10 domains** and rank a **TOP 5** (Section 4).
3. For each TOP 5 domain: select **1–3 indicators** and provide **scores** (Section 5).
4. Complete cross-cutting sections: **gender** and **capacity/feasibility**.

### Scoring scale (Section 5)
- **High** (3)
- **Medium** (2)
- **Low** (1)
- **NSP (Don’t know)** (0)

            """
        )
    )


def rubric_2(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 2 : Identification du répondant", "Section 2: Respondent identification"))
    st.info(
        t(
            lang,
            "Merci de renseigner ces informations. Elles servent uniquement à l’analyse et ne seront pas publiées nominativement.",
            "Please provide these details. They are used for analysis and will not be published in a personally identifiable way."
        )
    )

    resp_set("lang", lang)

    st.text_input(t(lang, "Nom de l'organisation", "Organization Name"), key="org_input", value=resp_get("organisation", ""))
    resp_set("organisation", st.session_state.get("org_input", "").strip())

    col1, col2 = st.columns(2)
    with col1:
        st.text_input(t(lang, "Pays", "Country"), key="country_input", value=resp_get("pays", ""))
        resp_set("pays", st.session_state.get("country_input", "").strip())
    with col2:
        st.text_input(t(lang, "Email", "Email"), key="email_input", value=resp_get("email", ""))
        resp_set("email", st.session_state.get("email_input", "").strip())

    type_options = [
        ("NSO", {"fr": "Institut national de statistique", "en": "National Statistical Office"}),
        ("Ministry", {"fr": "Ministère / Service statistique sectoriel", "en": "Ministry / Sector statistical unit"}),
        ("REC", {"fr": "Communauté économique régionale", "en": "Regional Economic Community"}),
        ("CivilSoc", {"fr": "Société civile", "en": "Civil society"}),
        ("DevPartner", {"fr": "Partenaire technique et financier", "en": "Development partner"}),
        ("Academia", {"fr": "Université / Recherche", "en": "Academia / Research"}),
        ("Other", {"fr": "Autre", "en": "Other"}),
    ]
    type_labels = [t(lang, x[1]["fr"], x[1]["en"]) for x in type_options]
    type_keys = [x[0] for x in type_options]

    # Type d’acteur : pas de pré-remplissage (placeholder)
    type_opts = [""] + type_keys
    prev_type = resp_get("type_acteur", "")
    idx = type_opts.index(prev_type) if prev_type in type_opts else 0

    chosen_type = st.selectbox(
        t(lang, "Type d’acteur", "Stakeholder type"),
        options=type_opts,
        index=idx,
        format_func=lambda k: (t(lang, "— Sélectionner —", "— Select —") if k == "" else type_labels[type_keys.index(k)]),
        help=t(lang, "Choisissez la catégorie correspondant le mieux à votre organisation.", 
               "Choose the category that best matches your organization.")
    )
    resp_set("type_acteur", chosen_type)
# Fonction dropdown : pas de pré-remplissage (placeholder)
    role_opts = ROLE_OPTIONS_FR if lang == "fr" else ROLE_OPTIONS_EN
    role_options = [""] + role_opts
    prev_role = resp_get("fonction", "")
    role_idx = role_options.index(prev_role) if prev_role in role_options else 0

    chosen_role = st.selectbox(
        t(lang, "Fonction", "Role/Function"),
        options=role_options,
        index=role_idx,
        format_func=lambda x: (t(lang, "— Sélectionner —", "— Select —") if x == "" else x),
        help=t(lang, "Indiquez votre fonction principale dans l’organisation.", "Indicate your main role in the organization."),
    )
    resp_set("fonction", chosen_role)
    if chosen_role in ["Autre", "Other"]:
        st.text_input(t(lang, "Préciser (fonction)", "Specify (role)"),
                      key="fonction_autre_input", value=resp_get("fonction_autre", ""))
        resp_set("fonction_autre", st.session_state.get("fonction_autre_input", "").strip())
    else:
        resp_set("fonction_autre", "")

    # Live errors
    errs = validate_r2(lang)
    if errs:
        st.warning(t(lang, "Veuillez corriger les éléments ci-dessous :", "Please fix the following:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_3(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 3 : Portée de la réponse", "Section 3: Scope of response"))
    st.markdown(
        t(
            lang,
            "Indiquez le périmètre principal de votre réponse. Cela aide à interpréter vos priorités.",
            "Indicate the main scope of your response. This helps interpret your priorities."
        )
    )

    scope_opts_raw = [
        ("National", {"fr": "National", "en": "National"}),
        ("Regional", {"fr": "Régional (CER)", "en": "Regional (REC)"}),
        ("Continental", {"fr": "Continental (UA)", "en": "Continental (AU)"}),
        ("Global", {"fr": "International", "en": "International"}),
        ("Other", {"fr": "Autre", "en": "Other"}),
    ]
    scope_labels = {k: t(lang, v["fr"], v["en"]) for k, v in scope_opts_raw}
    scope_keys = [k for k, _ in scope_opts_raw]
    scope_options = [""] + scope_keys

    prev_scope = resp_get("scope", "")
    scope_idx = scope_options.index(prev_scope) if prev_scope in scope_options else 0

    chosen_scope = st.selectbox(
        t(lang, "Portée", "Scope"),
        options=scope_options,
        index=scope_idx,
        format_func=lambda k: (t(lang, "— Sélectionner —", "— Select —") if k == "" else scope_labels.get(k, k)),
        help=t(
            lang,
            "Indiquez le périmètre principal de votre réponse : national, régional (CER), continental (UA) ou international.",
            "Indicate the main scope of your response: national, regional (REC), continental (AU), or international."
        )
    )
    resp_set("scope", chosen_scope)

    if resp_get("scope") == "Other":
        st.text_input(t(lang, "Préciser", "Specify"), key="scope_other_input", value=resp_get("scope_other", ""))
        resp_set("scope_other", st.session_state.get("scope_other_input", "").strip())
    else:
        resp_set("scope_other", "")



def rubric_4(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "Rubrique 4 : Domaines prioritaires", "Section 4: Priority domains"))

    st.info(
        t(
            lang,
            "Veuillez d’abord choisir 5 à 10 domaines (pré-sélection). Ensuite, choisissez exactement 5 domaines dans ce sous-ensemble (TOP 5).\n\nConseil : choisissez les domaines où la demande politique est forte.",
            "First select 5 to 10 domains (pre-selection). Then choose exactly 5 domains within that subset (TOP 5).\n\nTip: choose domains where policy demand is strong."
        )
    )

    domains = domains_from_longlist(df_long, lang)
    if not domains:
        st.error(
            t(
                lang,
                "La liste des domaines n’est pas disponible (longlist introuvable ou vide).",
                "Domain list is not available (longlist not found or empty).",
            )
        )
        st.caption(
            t(
                lang,
                "Vérifiez que le dépôt contient : data/indicator_longlist.csv (prioritaire) ou data/longlist.xlsx (ou ces fichiers à la racine).",
                "Check that the repository contains: data/indicator_longlist.csv (preferred) or data/longlist.xlsx (or these files at repo root).",
            )
        )
        return

    code_to_label = {c: lbl for c, lbl in domains}

    # Build display labels without showing codes (codes are stored internally)
    labels = [code_to_label[c] for c, _ in domains]
    # Disambiguate duplicates if any (rare)
    seen = {}
    for i, (c, _) in enumerate(domains):
        lbl = code_to_label[c]
        seen[lbl] = seen.get(lbl, 0) + 1
    display_labels = []
    label_to_code = {}
    for c, _ in domains:
        lbl = code_to_label[c]
        disp = lbl if seen[lbl] == 1 else f"{lbl} ({c})"
        display_labels.append(disp)
        label_to_code[disp] = c

    st.markdown(
        t(
            lang,
            """
### Étape 1 : Pré-sélection
Sélectionnez **entre 5 et 10 domaines** (sans doublons).
            """,
            """
### Step 1: Pre-selection
Select **5 to 10 domains** (no duplicates).
            """
        )
    )

    pre_default_codes = resp_get("preselected_domains", [])
    pre_default_disp = []
    for c in pre_default_codes:
        lbl = code_to_label.get(c, "")
        if not lbl:
            continue
        disp = lbl if seen.get(lbl, 1) == 1 else f"{lbl} ({c})"
        if disp in label_to_code:
            pre_default_disp.append(disp)

    # Avoid "first click not kept" by initializing widget state once (no default on every rerun)
    if "r4_preselection_ms" not in st.session_state:
        st.session_state["r4_preselection_ms"] = pre_default_disp

    pre_disp = st.multiselect(
        t(lang, "Pré-sélection (5–10 domaines)", "Pre-selection (5–10 domains)"),
        options=display_labels,
        max_selections=10,
        key="r4_preselection_ms",
        help=t(lang, "Choisissez au maximum 10 domaines. Une fois 10 domaines sélectionnés, les nouveaux clics seront ignorés.", 
               "Select up to 10 domains. Once 10 domains are selected, additional clicks are ignored.")
    )

    pre_codes = [label_to_code[x] for x in pre_disp]
    resp_set("preselected_domains", pre_codes)

    st.divider()
    st.markdown(
        t(
            lang,
            """
### Étape 2 : Classement TOP 5
Classez exactement **5 domaines** parmi votre pré-sélection.
            """,
            """
### Step 2: Rank TOP 5
Rank exactly **5 domains** from your pre-selection.
            """
        )
    )

    if len(pre_codes) < 5:
        st.warning(t(lang, "Sélectionnez d’abord au moins 5 domaines dans la pré-sélection.",
                     "Please pre-select at least 5 domains first."))
        resp_set("top5_domains", [])
        return

    top5: List[str] = []

    # Ranking with 5 selectboxes (no prefill + no duplicates)
    chosen_prev: List[str] = []
    for i in range(5):
        key = f"top5_rank_{i+1}"

        # Options for this rank = preselection minus already chosen
        remaining = [c for c in pre_codes if c not in chosen_prev]
        options = [""] + remaining  # "" placeholder (no prefill)

        prev = resp_get(key, "")
        if prev and prev in remaining:
            idx = options.index(prev)
        else:
            idx = 0

        choice = st.selectbox(
            t(lang, f"Rang {i+1}", f"Rank {i+1}"),
            options=options,
            index=idx,
            format_func=lambda c: (t(lang, "— Sélectionner —", "— Select —") if c == "" else code_to_label.get(c, c)),
            help=t(
                lang,
                "Choisissez un domaine unique pour chaque rang. Les domaines déjà choisis ne sont plus proposés aux rangs suivants.",
                "Choose a unique domain for each rank. Already selected domains are removed from the next ranks.",
            ),
            key=key,
        )

        if choice != "":
            top5.append(choice)
            chosen_prev.append(choice)

    resp_set("top5_domains", top5)


    errs = validate_r4(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))



def rubric_5(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "Rubrique 5 : Statistiques prioritaires et notation", "Section 5: Priority indicators and scoring"))

    top5 = resp_get("top5_domains", [])
    if not top5 or len(top5) != 5:
        st.warning(t(lang, "Veuillez d’abord finaliser le TOP 5 des domaines (Rubrique 4).",
                     "Please complete TOP 5 domains first (Section 4)."))
        return

    # mapping for domain display
    dom_map = {c: lbl for c, lbl in domains_from_longlist(df_long, lang)}

    st.markdown(
        t(
            lang,
            """
### Étape A : Sélection des statistiques
Pour chaque domaine du TOP 5 : choisissez **1 à 3 statistiques**.
- Total attendu : **entre 5 et 15** statistiques.
- Une statistique ne doit pas apparaître dans deux domaines.

### Étape B : Notation multicritères
Pour chaque statistique sélectionnée, attribuez une note (0–3) sur :
- **Écart de données** : manque actuel / insuffisance
- **Demande politique** : intérêt politique / stratégique
- **Faisabilité** : capacité de production à 12–24 mois
            """,
            """
### Step A: Select indicators
For each TOP 5 domain: select **1 to 3 indicators**.
- Expected total: **5 to 15** indicators.
- The same indicator must not be selected under two domains.

### Step B: Multi-criteria scoring
For each selected indicator, provide a score (0–3) for:
- **Data gap**
- **Political demand**
- **Feasibility (12–24 months)**
            """
        )
    )

    selected_by_domain: Dict[str, List[str]] = resp_get("selected_by_domain", {})
    if not isinstance(selected_by_domain, dict):
        selected_by_domain = {}

    scoring: Dict[str, Dict[str, Any]] = resp_get("scoring", {})
    if not isinstance(scoring, dict):
        scoring = {}

    # Ensure dict keys exist
    for d in top5:
        if d not in selected_by_domain:
            selected_by_domain[d] = []

    # UI selection per domain (codes hidden)
    for d in top5:
        st.markdown(f"#### {dom_map.get(d, d)}")

        stats_opts = stats_for_domain(df_long, d, lang)
        stat_code_to_label = {c: lbl for c, lbl in stats_opts}

        # build display labels without showing stat codes
        labels = [stat_code_to_label[c] for c, _ in stats_opts]
        seen = {}
        for c, _ in stats_opts:
            lbl = stat_code_to_label[c]
            seen[lbl] = seen.get(lbl, 0) + 1
        display_labels = []
        label_to_code = {}
        for c, _ in stats_opts:
            lbl = stat_code_to_label[c]
            disp = lbl if seen[lbl] == 1 else f"{lbl} ({c})"
            display_labels.append(disp)
            label_to_code[disp] = c

        default_codes = selected_by_domain.get(d, [])
        default_disp = []
        for c in default_codes:
            lbl = stat_code_to_label.get(c, "")
            if not lbl:
                continue
            disp = lbl if seen.get(lbl, 1) == 1 else f"{lbl} ({c})"
            if disp in label_to_code:
                default_disp.append(disp)

        key_ms = f"stats_ms_{d}"

        # Init widget state once (avoid "first click" issues)
        if key_ms not in st.session_state:
            st.session_state[key_ms] = default_disp

        picked_disp = st.multiselect(
            t(lang, "Choisir 1 à 3 statistiques", "Select 1 to 3 indicators"),
            options=display_labels,
            max_selections=3,
            key=key_ms,
            help=t(lang, "Sélectionnez au minimum 1 et au maximum 3 statistiques pour ce domaine.",
                   "Select at least 1 and at most 3 indicators for this domain.")
        )

        picked_codes = [label_to_code[x] for x in picked_disp]
        selected_by_domain[d] = picked_codes


    # Uniqueness check
    flattened = []
    for d in top5:
        flattened.extend(selected_by_domain.get(d, []))
    duplicates = [x for x in set(flattened) if flattened.count(x) > 1]
    if duplicates:
        st.error(
            t(
                lang,
                "Une ou plusieurs statistiques sont sélectionnées dans plusieurs domaines. Veuillez corriger.",
                "One or more indicators are selected under multiple domains. Please correct."
            )
        )

    resp_set("selected_by_domain", selected_by_domain)
    resp_set("selected_stats", flattened)

    # Map codes to labels for display in scoring
    global_map = {}
    for d in top5:
        for c, lbl in stats_for_domain(df_long, d, lang):
            global_map[c] = lbl

    st.divider()
    st.markdown("### " + t(lang, "Notation multicritères", "Multi-criteria scoring"))

    for s in flattened:
        if s not in scoring:
            scoring[s] = {"gap": None, "demand": None, "feasibility": None}

        st.markdown(f"**{global_map.get(s, s)}**")

        c1, c2, c3 = st.columns(3)
        with c1:
            opts = [None, 1, 2, 3, 0]  # None = placeholder (no prefill). 0 = NSP
            prev = scoring[s].get("gap", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["gap"] = st.selectbox(
                t(lang, "Écart de données", "Data gap"),
                options=opts,
                index=idx,
                format_func=score_format(lang),
                help=t(
                    lang,
                    "Définition : Écart actuel entre les besoins et les données disponibles.",
                    "Definition: current gap between needs and available data.",
                ),
                key=f"sc_gap_{s}",
            )
        with c2:
            opts = [None, 1, 2, 3, 0]
            prev = scoring[s].get("demand", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["demand"] = st.selectbox(
                t(lang, "Demande politique", "Political demand"),
                options=opts,
                index=idx,
                format_func=score_format(lang),
                help=t(
                    lang,
                    "Définition : importance stratégique / demande des décideurs.",
                    "Definition: strategic importance / demand from decision-makers.",
                ),
                key=f"sc_dem_{s}",
            )
        with c3:
            opts = [None, 1, 2, 3, 0]
            prev = scoring[s].get("feasibility", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["feasibility"] = st.selectbox(
                t(lang, "Faisabilité 12–24 mois", "Feasibility 12–24 months"),
                options=opts,
                index=idx,
                format_func=score_format(lang),
                help=t(
                    lang,
                    "Définition : capacité réaliste à produire la statistique d’ici 12–24 mois.",
                    "Definition: realistic ability to produce within 12–24 months.",
                ),
                key=f"sc_fea_{s}",
            )

    resp_set("scoring", scoring)

    errs = validate_r5(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_6(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 6 : Perspective de genre", "Section 6: Gender perspective"))
    st.markdown(
        t(
            lang,
            "Indiquez si les statistiques prioritaires doivent intégrer ces dimensions (Oui/Non/Selon indicateur/NSP).",
            "Indicate whether priority indicators should integrate these dimensions (Yes/No/Indicator-specific/UK)."
        )
    )

    options = [
        (t(lang, "Oui", "Yes"), "YES"),
        (t(lang, "Non", "No"), "NO"),
        (t(lang, "Selon indicateur", "Indicator-specific"), "SPEC"),
        (UK_FR if lang == "fr" else UK_EN, "UK"),
    ]
    labels = [x[0] for x in options]
    code_map = {x[0]: x[1] for x in options}

    items_fr = [
        "Désagrégation par sexe",
        "Désagrégation par âge",
        "Milieu urbain / rural",
        "Handicap",
        "Quintile de richesse",
    ]
    items_en = [
        "Disaggregation by sex",
        "Disaggregation by age",
        "Urban / rural",
        "Disability",
        "Wealth quintile",
    ]
    items = items_fr if lang == "fr" else items_en

    tbl = resp_get("gender_table", {})
    if not isinstance(tbl, dict):
        tbl = {}

    for it in items:
        rev_map = {v: k for k, v in code_map.items()}
        prev_code = tbl.get(it, None)
        idx = labels.index(rev_map[prev_code]) if prev_code in rev_map else None
        chosen = st.radio(it, options=labels, index=idx, horizontal=True, key=f"gender_{it}")
        tbl[it] = code_map.get(chosen, None)

    resp_set("gender_table", tbl)

    errs = validate_r6(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_8(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 8 : Capacité et faisabilité (12–24 mois)", "Section 8: Capacity and feasibility (12–24 months)"))
    st.markdown(
        t(
            lang,
            "Évaluez le niveau de disponibilité et d’adéquation des moyens pour produire les statistiques prioritaires dans les 12–24 mois à venir.",
            "Assess the availability and adequacy of resources to produce priority statistics in the coming 12–24 months."
        )
    )

    scale = [
        (t(lang, "Élevé", "High"), "HIGH"),
        (t(lang, "Moyen", "Medium"), "MED"),
        (t(lang, "Faible", "Low"), "LOW"),
        (UK_FR if lang == "fr" else UK_EN, "UK"),
    ]
    labels = [x[0] for x in scale]
    code_map = {x[0]: x[1] for x in scale}

    st.caption(t(lang, "Échelle : Élevé = capacité suffisante et opérationnelle ; Moyen = partiellement disponible ; Faible = insuffisant ; NSP = ne sait pas.",
                   "Scale: High = sufficient and operational; Medium = partially available; Low = insufficient; DK = does not know."))

    items_fr = [
        "Compétences statistiques disponibles",
        "Accès aux données administratives",
        "Financement disponible",
        "Outils numériques (collecte, traitement, diffusion)",
        "Cadre juridique pour le partage de données",
        "Coordination interinstitutionnelle",
    ]
    items_en = [
        "Available statistical skills",
        "Access to administrative data",
        "Available funding",
        "Digital tools (collection, processing, dissemination)",
        "Legal framework for data sharing",
        "Inter-institutional coordination",
    ]
    items = items_fr if lang == "fr" else items_en

    helps_fr = [
        "Ressources humaines : disponibilité de statisticiens/analystes qualifiés et expérience pertinente.",
        "Accès aux données administratives : disponibilité, qualité, régularité et conditions d’accès pour usage statistique.",
        "Financement : budget disponible et soutenable pour la production, y compris opérations de collecte/traitement.",
        "Outils numériques : disponibilité et adéquation des outils pour collecte, traitement, stockage, diffusion, interopérabilité (logiciels, matériel, connectivité, sécurité).",
        "Cadre juridique : existence et applicabilité des textes/accords permettant le partage de données à des fins statistiques (lois, décrets, protocoles, MoU, clauses de confidentialité).",
        "Coordination : mécanismes de coordination interinstitutionnelle (comités, conventions, échanges réguliers, standards communs).",
    ]
    helps_en = [
        "Human resources: availability of qualified statisticians/analysts and relevant experience.",
        "Access to administrative data: availability, quality, timeliness and conditions of access for statistical use.",
        "Funding: available and sustainable budget for production, including collection/processing operations.",
        "Digital tools: availability and adequacy of tools for collection, processing, storage, dissemination, interoperability (software, hardware, connectivity, security).",
        "Legal framework: existence and enforceability of texts/agreements enabling data sharing for statistical purposes (laws, decrees, protocols, MoUs, confidentiality clauses).",
        "Coordination: inter-institutional coordination mechanisms (committees, agreements, regular exchanges, shared standards).",
    ]
    helps = helps_fr if lang == "fr" else helps_en

    tbl = resp_get("capacity_table", {})
    if not isinstance(tbl, dict):
        tbl = {}

    rev_map = {v: k for k, v in code_map.items()}
    for it, hp in zip(items, helps):
        prev_code = tbl.get(it, None)
        idx = labels.index(rev_map[prev_code]) if prev_code in rev_map else None
        chosen = st.radio(it, options=labels, index=idx, horizontal=True, key=f"cap_{it}", help=hp)
        tbl[it] = code_map.get(chosen, None)

    resp_set("capacity_table", tbl)

    errs = validate_r8(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_9(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 9 : Harmonisation et qualité", "Section 9: Harmonization and quality"))
    st.markdown(
        t(
            lang,
            "Indiquez les exigences clés attendues en matière d’harmonisation et d’assurance qualité.",
            "Indicate key expectations regarding harmonization and quality assurance."
        )
    )

    opts_fr = [
        "Normes internationales (ONU, FMI, UA, etc.)",
        "Méthodologies harmonisées au niveau continental",
        "Calendrier de diffusion et révisions documentées",
        "Accès aux métadonnées",
        "Autre",
    ]
    opts_en = [
        "International standards (UN, IMF, AU, etc.)",
        "Continental harmonized methodologies",
        "Release calendar and documented revisions",
        "Access to metadata",
        "Other",
    ]
    opts = opts_fr if lang == "fr" else opts_en
    default = resp_get("quality_expectations", [])
    sel = st.multiselect(t(lang, "Sélectionnez", "Select"), options=opts, default=default, key="r9_multiselect")
    resp_set("quality_expectations", sel)
    if ("Autre" in sel) or ("Other" in sel):
        st.text_input(t(lang, "Préciser (Autre)", "Specify (Other)"),
                      key="q9_other_input", value=resp_get("quality_other", ""))
        resp_set("quality_other", st.session_state.get("q9_other_input", "").strip())
    else:
        resp_set("quality_other", "")


def rubric_10(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 10 : Diffusion", "Section 10: Dissemination"))
    st.markdown(
        t(
            lang,
            "Indiquez les canaux de diffusion jugés les plus utiles pour les statistiques prioritaires.",
            "Indicate the dissemination channels you find most useful for priority statistics."
        )
    )
    opts_fr = [
        "Portail web / tableaux de bord",
        "Communiqués / notes de conjoncture",
        "Microdonnées anonymisées (accès sécurisé)",
        "API / Open data",
        "Ateliers et webinaires",
        "Autre",
    ]
    opts_en = [
        "Web portal / dashboards",
        "Press releases / bulletins",
        "Anonymized microdata (secure access)",
        "API / Open data",
        "Workshops and webinars",
        "Other",
    ]
    opts = opts_fr if lang == "fr" else opts_en
    # Éviter les problèmes de clic (init du widget une seule fois)
    if "r10_multiselect" not in st.session_state:
        st.session_state["r10_multiselect"] = resp_get("dissemination_channels", [])
    sel = st.multiselect(
        t(lang, "Sélectionnez", "Select"),
        options=opts,
        key="r10_multiselect",
        help=t(lang, "Choisissez les canaux de diffusion les plus utiles.", "Select the most useful dissemination channels.")
    )
    resp_set("dissemination_channels", sel)
    if ("Autre" in sel) or ("Other" in sel):
        st.text_input(t(lang, "Préciser (Autre)", "Specify (Other)"),
                      key="q10_other_input", value=resp_get("dissemination_other", ""))
        resp_set("dissemination_other", st.session_state.get("q10_other_input", "").strip())
    else:
        resp_set("dissemination_other", "")


def rubric_12(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 12 : Questions ouvertes", "Section 12: Open questions"))
    st.markdown(
        t(
            lang,
            "Ces questions sont **optionnelles**, mais une alerte apparaîtra si vous laissez le champ vide.",
            "These questions are **optional**, but you will see a warning if left empty."
        )
    )

    q1 = st.text_area(
        t(lang, "1) Commentaires / recommandations clés", "1) Key comments / recommendations"),
        value=resp_get("open_q1", ""),
        height=120,
        key="open_q1_input"
    )
    resp_set("open_q1", q1.strip())

    q2 = st.text_area(
        t(
            lang,
            "2) Un ou des indicateur(s) statistique(s) socio-économique(s) essentiel(s) manquant(s) et justification(s)",
            "2) One or more missing essential socio-economic statistical indicator(s) and justification(s)"
        ),
        value=resp_get("open_q2", ""),
        height=120,
        key="open_q2_input"
    )
    resp_set("open_q2", q2.strip())

    q3 = st.text_area(
        t(lang, "3) Besoins de soutien (technique, financier, etc.)", "3) Support needs (technical, financial, etc.)"),
        value=resp_get("open_q3", ""),
        height=120,
        key="open_q3_input"
    )
    resp_set("open_q3", q3.strip())

    if not resp_get("open_q1", ""):
        st.warning(t(lang, "Alerte : la question 1 est vide (vous pouvez tout de même continuer).",
                     "Warning: question 1 is empty (you can still proceed)."))
    if not resp_get("open_q2", ""):
        st.warning(t(lang, "Alerte : la question 2 est vide (vous pouvez tout de même continuer).",
                     "Warning: question 2 is empty (you can still proceed)."))
    if not resp_get("open_q3", ""):
        st.warning(t(lang, "Alerte : la question 3 est vide (vous pouvez tout de même continuer).",
                     "Warning: question 3 is empty (you can still proceed)."))


def rubric_send(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "ENVOYER le questionnaire", "SUBMIT questionnaire"))

    errors = validate_all(lang)
    if errors:
        st.error(t(lang, "Le questionnaire contient des erreurs bloquantes :", "There are blocking errors:"))
        st.write("\n".join([f"- {e}" for e in errors]))
        st.info(t(lang, "Retournez aux rubriques concernées via la navigation.", "Go back to the relevant sections using navigation."))
        return

    # Optional warnings
    if not resp_get("open_q1", "") or not resp_get("open_q2", "") or not resp_get("open_q3", ""):
        st.warning(t(lang, "Certaines questions ouvertes sont vides (optionnel).", "Some open questions are empty (optional)."))

    st.info(t(lang, "Tout est prêt. Cliquez sur **ENVOYER** pour soumettre votre questionnaire.",
              "Everything is ready. Click **SUBMIT** to send your questionnaire."))

    # Empêcher les envois multiples (par email + par session)
    email = (resp_get("email", "") or "").strip()
    already_in_db = db_email_exists(email) if email else False
    already_in_session = bool(st.session_state.get("submitted_once", False))

    if already_in_db:
        st.error(t(lang, "Ce questionnaire a déjà été envoyé avec cet email. Un seul envoi est autorisé.",
                   "This questionnaire has already been submitted with this email. Only one submission is allowed."))

    if already_in_session and not already_in_db:
        st.info(t(lang, "Ce navigateur a déjà effectué un envoi. Pour un nouvel envoi, utilisez un autre email / session.",
                  "This browser session already submitted once. For a new submission, use another email / session."))

    disable_submit = already_in_db or already_in_session

    if st.button(t(lang, "✅ ENVOYER et enregistrer", "✅ SUBMIT and save"), disabled=disable_submit):
        submission_id = str(uuid.uuid4())
        payload = st.session_state.responses.copy()
        payload["submission_id"] = submission_id
        payload["submitted_at_utc"] = now_utc_iso()

        # Save locally (SQLite)
        db_save_submission(submission_id, lang, email, payload)

        # Optional storage integrations
        gs_ok, gs_msg = google_sheets_append(payload)
        dbx_ok, dbx_msg = dropbox_upload_json(submission_id, payload)

        st.success(t(lang, "Merci ! Votre questionnaire a été enregistré.", "Thank you! Your submission has been saved."))

        # Block multiple sends for this session
        st.session_state.submitted_once = True
        st.caption(f"ID : {submission_id}")

        # (No respondent-side download or storage details)
        st.info(t(lang, "Envoi terminé. Vous pouvez fermer cette page.", "Submission complete. You can close this page."))

        # Reset session marker (optional)
        st.session_state.submission_id = submission_id


# =========================
# Admin dashboard
# =========================

def admin_login(lang: str) -> None:
    st.subheader(t(lang, "Administration", "Administration"))
    pw = st.text_input(t(lang, "Mot de passe admin", "Admin password"), type="password")
    if st.button(t(lang, "Se connecter", "Login")):
        expected = st.secrets.get("ADMIN_PASSWORD", None)
        if expected and pw == expected:
            st.session_state.admin_authed = True
            st.success(t(lang, "Connexion réussie.", "Logged in."))
            st.rerun()
        else:
            st.error(t(lang, "Mot de passe incorrect ou secret ADMIN_PASSWORD manquant.", "Incorrect password or missing ADMIN_PASSWORD secret."))



def admin_dashboard(lang: str) -> None:
    st.subheader(t(lang, "Tableau de bord admin", "Admin dashboard"))

    # Load data from SQLite
    df = db_read_submissions(limit=20000)
    st.metric(t(lang, "Nombre de réponses", "Number of responses"), len(df))

    if df.empty:
        st.info(t(lang, "Aucune réponse pour le moment.", "No responses yet."))
        return

    # Parse payloads
    payloads = []
    for _, r in df.iterrows():
        try:
            payloads.append(json.loads(r["payload_json"]))
        except Exception:
            payloads.append({})

    # Flat view for quick export
    flat = pd.DataFrame([flatten_payload(p) for p in payloads])
    flat.insert(0, "submission_id", df["submission_id"].values)
    flat.insert(1, "submitted_at_utc", df["submitted_at_utc"].values)

    tab_quick, tab_super = st.tabs([
        t(lang, "Vue rapide", "Quick view"),
        t(lang, "Super admin", "Super admin")
    ])

    with tab_quick:
        st.dataframe(flat, use_container_width=True)

        # Export Excel (flat + raw)
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine="openpyxl") as writer:
            flat.to_excel(writer, sheet_name="submissions", index=False)
            df.to_excel(writer, sheet_name="raw_json", index=False)
        out.seek(0)

        st.download_button(
            t(lang, "Exporter en Excel", "Export to Excel"),
            data=out.getvalue(),
            file_name="consultation_stat_niang_export.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        # Download DB file
        if os.path.exists(DB_PATH):
            with open(DB_PATH, "rb") as f:
                st.download_button(
                    t(lang, "Télécharger la base SQLite (.db)", "Download SQLite database (.db)"),
                    data=f.read(),
                    file_name="responses.db",
                    mime="application/octet-stream",
                )

            # CSV export of the SQLite table (for easier sharing/analysis)
            csv_db = db_dump_csv_bytes(limit=2000000)
            st.download_button(
                t(lang, "Télécharger la base SQLite convertie en CSV", "Download SQLite exported as CSV"),
                data=csv_db,
                file_name="responses.csv",
                mime="text/csv",
            )

            # Full consolidated base (ZIP) for admin
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
                zf.writestr("responses.csv", csv_db)
                zf.writestr("consultation_stat_niang_export.xlsx", out.getvalue())
                zf.writestr("submissions_flat.csv", flat.to_csv(index=False))
                zf.writestr("responses_raw_json.csv", df.to_csv(index=False))
            zip_buf.seek(0)
            st.download_button(
                t(lang, "Télécharger la base consolidée complète (ZIP)", "Download full consolidated base (ZIP)"),
                data=zip_buf.getvalue(),
                file_name="consultation_stat_niang_full_base.zip",
                mime="application/zip",
            )

        # Upload exports to Dropbox (optional)
        st.markdown("#### Dropbox")
        if st.button(t(lang, "Envoyer l’Excel sur Dropbox", "Upload Excel to Dropbox"), key="dbx_excel_btn"):
            try:
                token = st.secrets["DROPBOX_ACCESS_TOKEN"]
                folder = st.secrets.get("DROPBOX_FOLDER", "/consultation_stat_niang")
                folder = folder if folder.startswith("/") else "/" + folder
                path = f"{folder}/exports/export_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.xlsx"
                dbx = dropbox.Dropbox(token)
                dbx.files_upload(out.getvalue(), path, mode=dropbox.files.WriteMode.overwrite)
                st.success(t(lang, "Excel envoyé sur Dropbox.", "Excel uploaded to Dropbox."))
            except Exception as e:
                st.error(f"Dropbox : {e}")

        if st.button(t(lang, "Envoyer la base .db sur Dropbox", "Upload database to Dropbox"), key="dbx_db_btn"):
            try:
                token = st.secrets["DROPBOX_ACCESS_TOKEN"]
                folder = st.secrets.get("DROPBOX_FOLDER", "/consultation_stat_niang")
                folder = folder if folder.startswith("/") else "/" + folder
                path = f"{folder}/exports/responses_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.db"
                dbx = dropbox.Dropbox(token)
                with open(DB_PATH, "rb") as fdb:
                    dbx.files_upload(fdb.read(), path, mode=dropbox.files.WriteMode.overwrite)
                st.success(t(lang, "Base envoyée sur Dropbox.", "Database uploaded to Dropbox."))
            except Exception as e:
                st.error(f"Dropbox : {e}")

        st.info(
            t(
                lang,
                "Astuce : utilisez ?admin=1 dans l’URL pour afficher l’espace admin.",
                "Tip: use ?admin=1 in the URL to access the admin space."
            )
        )

    with tab_super:
        # --- Build a richer dataset for analysis
        df_super = pd.DataFrame(payloads)
        df_super["submission_id"] = df["submission_id"].values
        df_super["submitted_at_utc"] = pd.to_datetime(df["submitted_at_utc"], errors="coerce", utc=True)

        # Filters
        st.markdown("### " + t(lang, "Filtres", "Filters"))

        colf1, colf2, colf3 = st.columns(3)
        with colf1:
            countries = sorted([c for c in df_super.get("pays", pd.Series(dtype=str)).dropna().unique().tolist() if str(c).strip() != ""])
            sel_countries = st.multiselect(t(lang, "Pays", "Country"), options=countries, default=[], key="f_country")
        with colf2:
            actors = sorted([a for a in df_super.get("type_acteur", pd.Series(dtype=str)).dropna().unique().tolist() if str(a).strip() != ""])
            sel_actors = st.multiselect(t(lang, "Type d’acteur", "Stakeholder type"), options=actors, default=[], key="f_actor")
        with colf3:
            # Period filter
            min_dt = df_super["submitted_at_utc"].min()
            max_dt = df_super["submitted_at_utc"].max()
            if pd.isna(min_dt) or pd.isna(max_dt):
                min_dt = pd.Timestamp.utcnow() - pd.Timedelta(days=30)
                max_dt = pd.Timestamp.utcnow()
            date_range = st.date_input(
                t(lang, "Période", "Period"),
                value=(min_dt.date(), max_dt.date()),
                key="f_period"
            )

        filtered = df_super.copy()
        if sel_countries:
            filtered = filtered[filtered["pays"].isin(sel_countries)]
        if sel_actors:
            filtered = filtered[filtered["type_acteur"].isin(sel_actors)]
        if isinstance(date_range, (list, tuple)) and len(date_range) == 2:
            # Streamlit date_input returns datetime.date; our column may be tz-aware (UTC).
            col = pd.to_datetime(filtered["submitted_at_utc"], utc=True, errors="coerce")

            start_d = pd.Timestamp(date_range[0])
            end_d = pd.Timestamp(date_range[1]) + pd.Timedelta(days=1) - pd.Timedelta(seconds=1)

            if start_d.tz is None:
                start_d = start_d.tz_localize("UTC")
            else:
                start_d = start_d.tz_convert("UTC")

            if end_d.tz is None:
                end_d = end_d.tz_localize("UTC")
            else:
                end_d = end_d.tz_convert("UTC")

            filtered = filtered[(col >= start_d) & (col <= end_d)]

        st.caption(t(lang, f"Réponses filtrées : {len(filtered)}", f"Filtered responses: {len(filtered)}"))

        if filtered.empty:
            st.warning(t(lang, "Aucune réponse dans ce filtre.", "No responses match these filters."))
            return

        # Longlist for labels (domain/stat)
        df_long = load_longlist()
        dom_lbl = domain_label_map(df_long, lang)
        stat_lbl = stat_label_map(df_long, lang)

        # --- Build aggregated prioritization table
        rows = []
        for _, p in filtered.iterrows():
            top5 = p.get("top5_domains", []) or []
            sel_by_dom = p.get("selected_by_domain", {}) or {}
            scoring = p.get("scoring", {}) or {}
            sid = p.get("submission_id", "")
            for d, stats in (sel_by_dom.items() if isinstance(sel_by_dom, dict) else []):
                if not isinstance(stats, list):
                    continue
                for s in stats:
                    sc = scoring.get(s, {})
                    gap = int(sc.get("gap", 0) or 0)
                    dem = int(sc.get("demand", 0) or 0)
                    fea = int(sc.get("feasibility", 0) or 0)
                    overall = (gap + dem + fea) / 3.0
                    rows.append({
                        "submission_id": sid,
                        "pays": p.get("pays", ""),
                        "type_acteur": p.get("type_acteur", ""),
                        "domain_code": d,
                        "domain_label": dom_lbl.get(d, d),
                        "stat_code": s,
                        "stat_label": stat_lbl.get(s, s),
                        "gap": gap, "demand": dem, "feasibility": fea, "overall": overall
                    })

        df_rows = pd.DataFrame(rows)
        if df_rows.empty:
            st.warning(t(lang, "Aucune statistique notée dans ces réponses.", "No scored indicators in these responses."))
            return

        # Aggregation
        by_stat = df_rows.groupby(["domain_code", "domain_label", "stat_code", "stat_label"], as_index=False).agg(
            n=("submission_id", "nunique"),
            mean_gap=("gap", "mean"),
            mean_demand=("demand", "mean"),
            mean_feasibility=("feasibility", "mean"),
            mean_overall=("overall", "mean"),
        ).sort_values(["domain_code", "mean_overall", "n"], ascending=[True, False, False])

        by_domain = df_rows.groupby(["domain_code", "domain_label"], as_index=False).agg(
            n_stats=("stat_code", "count"),
            n_submissions=("submission_id", "nunique"),
            mean_overall=("overall", "mean"),
        ).sort_values(["mean_overall", "n_submissions"], ascending=[False, False])

        st.markdown("### " + t(lang, "Tableau de priorisation agrégé", "Aggregated prioritization table"))
        st.dataframe(by_stat, use_container_width=True, height=420)

        # Export aggregated Excel
        out2 = io.BytesIO()
        with pd.ExcelWriter(out2, engine="openpyxl") as writer:
            by_domain.to_excel(writer, sheet_name="by_domain", index=False)
            by_stat.to_excel(writer, sheet_name="by_statistic", index=False)
            df_rows.to_excel(writer, sheet_name="scored_rows", index=False)
        out2.seek(0)

        st.download_button(
            t(lang, "Télécharger l’agrégé (Excel)", "Download aggregated (Excel)"),
            data=out2.getvalue(),
            file_name="prioritization_aggregated.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        colpub1, colpub2 = st.columns(2)
        with colpub1:
            if st.button(t(lang, "Publier l’agrégé sur Google Sheets", "Publish aggregates to Google Sheets"), key="btn_publish_gsheets"):
                ok1, msg1 = google_sheets_write_df(by_domain, "agg_by_domain")
                ok2, msg2 = google_sheets_write_df(by_stat, "agg_by_statistic")
                if ok1 and ok2:
                    st.success(t(lang, "Agrégats publiés sur Google Sheets (onglets : agg_by_domain, agg_by_statistic).",
                                 "Aggregates published to Google Sheets (tabs: agg_by_domain, agg_by_statistic)."))
                else:
                    st.error(f"Google Sheets : {msg1} / {msg2}")
        with colpub2:
            if st.button(t(lang, "Envoyer l’agrégé sur Dropbox", "Upload aggregates to Dropbox"), key="btn_publish_dropbox_agg"):
                ok, msg = dropbox_upload_bytes(out2.getvalue(), "prioritization_aggregated.xlsx", subfolder="aggregates")
                if ok:
                    st.success(t(lang, "Agrégé envoyé sur Dropbox.", "Aggregates uploaded to Dropbox."))
                else:
                    st.error(f"Dropbox : {msg}")

        # Rich Word report
        st.markdown("### " + t(lang, "Rapport Word (publication)", "Word report (publication)"))
        st.caption(t(lang, "Génère un rapport enrichi avec graphiques et annexes.", "Generates an enriched report with charts and annexes."))

        if st.button(t(lang, "Générer le rapport Word", "Generate Word report"), key="btn_word_publication"):
            try:
                doc_bytes = build_publication_report_docx(lang, filtered, by_domain, by_stat, df_rows)
                st.download_button(
                    t(lang, "Télécharger le rapport (.docx)", "Download report (.docx)"),
                    data=doc_bytes,
                    file_name="rapport_publication_priorisation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                st.success(t(lang, "Rapport généré.", "Report generated."))

                # Auto-upload to Dropbox (if configured)
                ok_dbx, msg_dbx = dropbox_upload_bytes(doc_bytes, "rapport_publication_priorisation.docx", subfolder="reports")
                if ok_dbx:
                    st.info(t(lang, "Rapport aussi envoyé sur Dropbox.", "Report also uploaded to Dropbox."))
                else:
                    # Not an error if Dropbox is not configured
                    if "manquant" in str(msg_dbx).lower() or "non configur" in str(msg_dbx).lower():
                        st.caption(t(lang, "Dropbox non configuré : le rapport n’a pas été envoyé.", "Dropbox not configured: report not uploaded."))
                    else:
                        st.warning(f"Dropbox : {msg_dbx}")
            except Exception as e:
                st.error(f"Word : {e}")


# =========================
# Main
# =========================

def main() -> None:
    st.set_page_config(page_title=APP_TITLE_FR, layout="wide")
    init_session()

    # Language toggle
    st.sidebar.title("🌐")
    lang = st.sidebar.selectbox(
        "Langue / Language",
        options=["fr", "en"],
        index=0 if st.session_state.lang == "fr" else 1
    )
    st.session_state.lang = lang

    # Admin access
    qp = get_query_params()
    is_admin = "admin" in qp and qp["admin"] and qp["admin"][0] in ["1", "true", "yes"]

    df_long = load_longlist()

    # Header
    st.title(t(lang, APP_TITLE_FR, APP_TITLE_EN))
    st.caption(t(lang, "Application unifiée (FR/EN) – codes masqués – contrôles qualité intégrés.",
                 "Unified app (FR/EN) – hidden codes – built-in quality controls."))

    # Sidebar navigation
    steps = get_steps(lang)
    render_sidebar(lang, steps)

    # Admin view
    if is_admin:
        if not st.session_state.admin_authed:
            admin_login(lang)
            return
        admin_dashboard(lang)
        return

    # Normal flow
    step_key = steps[st.session_state.nav_idx][0]

    if step_key == "R1":
        rubric_1(lang)
    elif step_key == "R2":
        rubric_2(lang)
    elif step_key == "R3":
        rubric_3(lang)
    elif step_key == "R4":
        rubric_4(lang, df_long)
    elif step_key == "R5":
        rubric_5(lang, df_long)
    elif step_key == "R6":
        rubric_6(lang)
    elif step_key == "R8":
        rubric_8(lang)
    elif step_key == "R9":
        rubric_9(lang)
    elif step_key == "R10":
        rubric_10(lang)
    elif step_key == "R12":
        rubric_12(lang)
    elif step_key == "SEND":
        rubric_send(lang, df_long)

    st.divider()
    nav_buttons(lang, steps, df_long)


if __name__ == "__main__":
    main()
