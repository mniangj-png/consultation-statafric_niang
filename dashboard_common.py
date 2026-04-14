from __future__ import annotations

import io
import json
import os
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import pandas as pd
import requests
from docx import Document
from docx.shared import Inches

RESPONSE_ROOT = os.getenv("RESPONSE_PATH_ROOT", "data/validation_doc")
DEFAULT_OWNER = os.getenv("GITHUB_OWNER", "mniangj-png")
DEFAULT_REPO = os.getenv("GITHUB_REPO", "consultation-statafric_niang")
DEFAULT_BRANCH = os.getenv("GITHUB_BRANCH", "main")

RESPONSE_LABELS_FR = {
    "go": "Validé",
    "go_with_reservations": "Validé sous réserve",
    "no_go": "Non-validé",
    "no_opinion": "Sans avis",
    "yes": "Oui",
    "mostly_yes": "Plutôt oui",
    "mostly_no": "Plutôt non",
    "no": "Non",
    "yes_limited_adjustments": "Oui, sous réserve d’ajustements limités",
    "no_substantial_revision": "Non, une révision plus substantielle est nécessaire",
    "discuss_in_workshop": "À discuter en atelier",
    "nso": "Institut national de statistique (INS)",
    "rec": "Communauté économique régionale (CER)",
    "director_general": "Directeur général",
    "statistician_general": "Statisticien général",
    "deputy_director_general": "Directeur général adjoint",
    "director_statistics": "Directeur des statistiques",
    "head_department": "Chef de département / unité",
    "programme_manager": "Responsable / coordinateur de programme",
    "technical_expert": "Expert technique",
    "other": "Autre",
}

STRATEGIC_ROWS_FR = {
    "strategic_prioritization_criteria": "Critères de priorisation retenus",
    "strategic_scoring_logic": "Logique de notation multicritère (scoring)",
    "strategic_core_extensions": "Distinction noyau / extensions",
    "strategic_gender_integration": "Intégration transversale du genre",
    "strategic_min_disaggregations": "Désagrégations minimales proposées",
    "strategic_data_sources": "Sources de données et dispositifs de production",
    "strategic_governance_roles": "Gouvernance et répartition des rôles",
    "strategic_roadmap_update": "Feuille de route de mise en œuvre et mécanisme de mise à jour",
}

DOMAIN_ROWS_FR = {
    "domain_d01": "D01 Croissance économique, transformation structurelle et commerce",
    "domain_d02": "D02 Emploi, travail décent et protection sociale",
    "domain_d03": "D03 Agriculture durable, sécurité alimentaire et nutrition",
    "domain_d04": "D04 Infrastructures, industrialisation et innovation",
    "domain_d05": "D05 Inclusion, pauvreté et inégalités",
    "domain_d06": "D06 Éducation, compétences et capital humain",
    "domain_d07": "D07 Santé, bien-être et accès universel",
    "domain_d08": "D08 Égalité des genres et autonomisation",
    "domain_d09": "D09 Environnement, résilience climatique et villes durables",
    "domain_d10": "D10 Gouvernance, paix et institutions",
    "domain_d11": "D11 Économie bleue et gestion des océans",
    "domain_d12": "D12 Partenariats et financement du développement",
}

STRATEGIC_KEYS = list(STRATEGIC_ROWS_FR.keys())
DOMAIN_KEYS = list(DOMAIN_ROWS_FR.keys())


@dataclass
class GitHubConfig:
    owner: str = DEFAULT_OWNER
    repo: str = DEFAULT_REPO
    branch: str = DEFAULT_BRANCH
    token: str = ""


def sanitize_token(token: str | None) -> str:
    token = (token or "").strip().strip('"').strip("'")
    if token.lower().startswith("bearer "):
        token = token[7:].strip()
    return token.replace("\r", "").replace("\n", "")


def get_github_config_from_env() -> GitHubConfig:
    return GitHubConfig(
        owner=os.getenv("GITHUB_OWNER", DEFAULT_OWNER),
        repo=os.getenv("GITHUB_REPO", DEFAULT_REPO),
        branch=os.getenv("GITHUB_BRANCH", DEFAULT_BRANCH),
        token=sanitize_token(os.getenv("GITHUB_TOKEN", "")),
    )


def get_github_config_from_streamlit(st_module) -> GitHubConfig:
    cfg = get_github_config_from_env()
    try:
        secrets = st_module.secrets
        if "github" in secrets:
            cfg.owner = secrets["github"].get("owner", cfg.owner)
            cfg.repo = secrets["github"].get("repo", cfg.repo)
            cfg.branch = secrets["github"].get("branch", cfg.branch)
            cfg.token = sanitize_token(secrets["github"].get("token", cfg.token))
        cfg.owner = secrets.get("GITHUB_OWNER", cfg.owner)
        cfg.repo = secrets.get("GITHUB_REPO", cfg.repo)
        cfg.branch = secrets.get("GITHUB_BRANCH", cfg.branch)
        cfg.token = sanitize_token(secrets.get("GITHUB_TOKEN", cfg.token))
    except Exception:
        pass
    return cfg


def github_headers(cfg: GitHubConfig) -> Dict[str, str]:
    headers = {
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    }
    if cfg.token:
        headers["Authorization"] = f"Bearer {cfg.token}"
    return headers


def github_get_json(url: str, cfg: GitHubConfig, params: Optional[Dict] = None) -> Dict | List:
    r = requests.get(url, headers=github_headers(cfg), params=params, timeout=60)
    r.raise_for_status()
    return r.json()


def resolve_branch(cfg: GitHubConfig) -> str:
    tried = []
    for branch in [cfg.branch, "main", "master"]:
        if not branch or branch in tried:
            continue
        tried.append(branch)
        try:
            url = f"https://api.github.com/repos/{cfg.owner}/{cfg.repo}/git/trees/{branch}"
            github_get_json(url, cfg)
            return branch
        except Exception:
            continue
    raise RuntimeError(f"Impossible d’accéder au dépôt GitHub. Branches testées : {', '.join(tried)}")


def list_repo_json_paths(cfg: GitHubConfig, root_path: str) -> Tuple[str, List[str]]:
    branch = resolve_branch(cfg)
    url = f"https://api.github.com/repos/{cfg.owner}/{cfg.repo}/git/trees/{branch}"
    tree = github_get_json(url, cfg, params={"recursive": "1"})
    paths = []
    prefix = root_path.rstrip("/") + "/"
    for item in tree.get("tree", []):
        if item.get("type") == "blob" and item.get("path", "").startswith(prefix) and item.get("path", "").endswith(".json"):
            paths.append(item["path"])
    return branch, sorted(paths)


def fetch_repo_file_text(cfg: GitHubConfig, path: str, branch: str) -> str:
    url = f"https://raw.githubusercontent.com/{cfg.owner}/{cfg.repo}/{branch}/{path}"
    headers = {}
    if cfg.token:
        headers["Authorization"] = f"Bearer {cfg.token}"
    r = requests.get(url, headers=headers, timeout=60)
    r.raise_for_status()
    return r.text


def load_json_records_from_repo(cfg: GitHubConfig, subfolder: str) -> Tuple[str, List[Dict]]:
    root = f"{RESPONSE_ROOT}/{subfolder}".rstrip("/")
    branch, paths = list_repo_json_paths(cfg, root)
    records: List[Dict] = []
    for path in paths:
        try:
            txt = fetch_repo_file_text(cfg, path, branch)
            item = json.loads(txt)
            if isinstance(item, dict):
                item["_source_path"] = path
                records.append(item)
        except Exception:
            continue
    return branch, records


def parse_iso_date(value: str | None) -> Optional[pd.Timestamp]:
    if not value:
        return None
    try:
        return pd.to_datetime(value, utc=True)
    except Exception:
        return None


def label_value(value):
    if isinstance(value, list):
        return ", ".join(map(str, value))
    return RESPONSE_LABELS_FR.get(value, value)


def flatten_record(rec: Dict) -> Dict:
    out = {}
    for k, v in rec.items():
        out[k] = label_value(v)
    return out


def records_to_dataframe(records: Sequence[Dict]) -> pd.DataFrame:
    rows = [flatten_record(r) for r in records]
    df = pd.DataFrame(rows)
    if df.empty:
        return df
    for dt_col in ["saved_at", "submitted_at", "expires_at"]:
        if dt_col in df.columns:
            df[dt_col] = pd.to_datetime(df[dt_col], errors="coerce", utc=True)
    for col in ["country_or_rec", "institution_type", "language", "status", "overall_validation", "operational_usability", "final_institutional_position"]:
        if col in df.columns:
            df[col] = df[col].fillna("")
    return df


def dataframe_to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8-sig")


def _excel_safe_scalar(value):
    if isinstance(value, pd.Timestamp):
        if pd.isna(value):
            return None
        if value.tzinfo is not None:
            value = value.tz_convert(None)
        # Write as text to avoid all Excel timezone issues
        return value.to_pydatetime().strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, datetime):
        if value.tzinfo is not None:
            value = value.astimezone(timezone.utc).replace(tzinfo=None)
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, pd.Timedelta):
        return str(value)
    if isinstance(value, dict):
        return json.dumps({str(k): _excel_safe_scalar(v) for k, v in value.items()}, ensure_ascii=False)
    if isinstance(value, (list, tuple, set)):
        return json.dumps([_excel_safe_scalar(v) for v in value], ensure_ascii=False)
    try:
        import numpy as np
        if isinstance(value, (np.integer, np.floating, np.bool_)):
            return value.item()
        if isinstance(value, np.ndarray):
            return json.dumps(value.tolist(), ensure_ascii=False)
    except Exception:
        pass
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass
    return value


def _prepare_dataframe_for_excel(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    if out.empty:
        return out
    # Force object dtype first so pandas does not preserve timezone-aware
    # datetime dtypes through Series.map/list assignment.
    out = out.astype(object)
    for col in out.columns:
        safe_values = []
        for value in out[col].tolist():
            safe_values.append(_excel_safe_scalar(value))
        out[col] = pd.Series(safe_values, dtype="object")
    return out


def dataframe_to_xlsx_bytes(sheets: Dict[str, pd.DataFrame]) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for name, df in sheets.items():
            safe_name = name[:31] if name else "Sheet1"
            safe_df = _prepare_dataframe_for_excel(df)
            safe_df.to_excel(writer, index=False, sheet_name=safe_name)
    output.seek(0)
    return output.getvalue()


def _json_safe(value):
    if isinstance(value, (pd.Timestamp, datetime)):
        try:
            return value.isoformat()
        except Exception:
            return str(value)
    if pd.isna(value):
        return None
    if isinstance(value, dict):
        return {str(k): _json_safe(v) for k, v in value.items()}
    if isinstance(value, (list, tuple, set)):
        return [_json_safe(v) for v in value]
    try:
        import numpy as np
        if isinstance(value, (np.integer, np.floating, np.bool_)):
            return value.item()
    except Exception:
        pass
    return value


def records_to_json_bytes(records: Sequence[Dict]) -> bytes:
    safe = [_json_safe(r) for r in list(records)]
    return json.dumps(safe, ensure_ascii=False, indent=2).encode("utf-8")


def response_count_table(df: pd.DataFrame, columns: Sequence[str], mapping: Dict[str, str]) -> pd.DataFrame:
    rows = []
    for col in columns:
        if col not in df.columns:
            continue
        s = df[col].fillna("")
        counts = s.value_counts(dropna=False)
        row = {"élément": mapping.get(col, col), "total réponses": int((s != "").sum())}
        for code, label in [("Validé", "Validé"), ("Validé sous réserve", "Validé sous réserve"), ("Non-validé", "Non-validé"), ("Sans avis", "Sans avis")]:
            row[label] = int(counts.get(code, 0))
        rows.append(row)
    out = pd.DataFrame(rows)
    if not out.empty:
        out["Taux de réserve (%)"] = ((out["Validé sous réserve"] / out["total réponses"]).fillna(0) * 100).round(1)
        out["Taux de non-validation (%)"] = ((out["Non-validé"] / out["total réponses"]).fillna(0) * 100).round(1)
    return out


def extract_justifications(df: pd.DataFrame, prefixes: Sequence[str], mapping: Dict[str, str]) -> pd.DataFrame:
    rows = []
    for prefix in prefixes:
        why_col = f"{prefix}_why"
        if prefix not in df.columns or why_col not in df.columns:
            continue
        mask = df[prefix].isin(["Validé sous réserve", "Non-validé"]) & df[why_col].fillna("").str.strip().ne("")
        subset = df.loc[mask, ["submission_id", "institution_acronym", "country_or_rec", prefix, why_col]].copy()
        if subset.empty:
            continue
        subset.insert(0, "élément", mapping.get(prefix, prefix))
        subset = subset.rename(columns={prefix: "position", why_col: "justification"})
        rows.append(subset)
    if not rows:
        return pd.DataFrame(columns=["élément", "submission_id", "institution_acronym", "country_or_rec", "position", "justification"])
    return pd.concat(rows, ignore_index=True)


def build_analysis_sheets(df: pd.DataFrame) -> Dict[str, pd.DataFrame]:
    if df.empty:
        empty = pd.DataFrame({"message": ["Aucune donnée disponible"]})
        return {"accueil": empty}

    summary = pd.DataFrame(
        {
            "indicateur": [
                "Nombre total d’enregistrements",
                "Nombre de soumissions finales",
                "Nombre de brouillons",
                "Nombre de pays / CER distincts",
                "Nombre d’institutions distinctes",
            ],
            "valeur": [
                len(df),
                int((df.get("status", "") == "submitted").sum()) if "status" in df.columns else 0,
                int((df.get("status", "") == "draft").sum()) if "status" in df.columns else 0,
                df.get("country_or_rec", pd.Series(dtype=str)).replace("", pd.NA).dropna().nunique(),
                df.get("institution_acronym", pd.Series(dtype=str)).replace("", pd.NA).dropna().nunique(),
            ],
        }
    )

    by_lang = pd.DataFrame(df.get("language", pd.Series(dtype=str)).value_counts(dropna=False)).reset_index()
    by_lang.columns = ["langue", "nombre"]

    by_inst = pd.DataFrame(df.get("institution_type", pd.Series(dtype=str)).value_counts(dropna=False)).reset_index()
    by_inst.columns = ["type_institution", "nombre"]

    overall = pd.DataFrame(df.get("overall_validation", pd.Series(dtype=str)).value_counts(dropna=False)).reset_index()
    overall.columns = ["validation_globale", "nombre"]

    final_pos = pd.DataFrame(df.get("final_institutional_position", pd.Series(dtype=str)).value_counts(dropna=False)).reset_index()
    final_pos.columns = ["position_finale", "nombre"]

    strategic = response_count_table(df, STRATEGIC_KEYS, STRATEGIC_ROWS_FR)
    domains = response_count_table(df, DOMAIN_KEYS, DOMAIN_ROWS_FR)
    justifs = extract_justifications(df, list(STRATEGIC_ROWS_FR) + list(DOMAIN_ROWS_FR) + ["overall_validation", "operational_usability", "final_institutional_position"], {
        **STRATEGIC_ROWS_FR,
        **DOMAIN_ROWS_FR,
        "overall_validation": "Validation globale du document",
        "operational_usability": "Document suffisamment opérationnel",
        "final_institutional_position": "Position finale de l’institution",
    })

    metadata = pd.DataFrame(
        {
            "paramètre": ["horodatage_génération_utc", "colonnes_source"],
            "valeur": [datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"), ", ".join(df.columns)],
        }
    )

    return {
        "synthese": summary,
        "soumissions_brutes": df.copy(),
        "langues": by_lang,
        "types_institution": by_inst,
        "validation_globale": overall,
        "position_finale": final_pos,
        "choix_methodo": strategic,
        "domaines": domains,
        "justifications": justifs,
        "metadata": metadata,
    }


def build_report_docx_bytes(df: pd.DataFrame) -> bytes:
    doc = Document()
    doc.add_heading("Rapport synthétique de validation", level=1)
    doc.add_paragraph(
        "Ce rapport est généré automatiquement à partir des réponses enregistrées dans le questionnaire de validation du document sur les statistiques socio-économiques prioritaires en Afrique."
    )
    doc.add_paragraph(f"Date de génération : {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}")

    if df.empty:
        doc.add_paragraph("Aucune donnée n’est disponible pour produire une synthèse.")
    else:
        total = len(df)
        submitted = int((df.get("status", "") == "submitted").sum()) if "status" in df.columns else 0
        drafts = int((df.get("status", "") == "draft").sum()) if "status" in df.columns else 0
        countries = df.get("country_or_rec", pd.Series(dtype=str)).replace("", pd.NA).dropna().nunique()

        doc.add_heading("1. Vue d’ensemble", level=2)
        p = doc.add_paragraph()
        p.add_run("Nombre total d’enregistrements : ").bold = True
        p.add_run(str(total))
        p = doc.add_paragraph()
        p.add_run("Soumissions finales : ").bold = True
        p.add_run(str(submitted))
        p = doc.add_paragraph()
        p.add_run("Brouillons : ").bold = True
        p.add_run(str(drafts))
        p = doc.add_paragraph()
        p.add_run("Pays / CER distincts : ").bold = True
        p.add_run(str(countries))

        doc.add_heading("2. Validation globale", level=2)
        overall = df.get("overall_validation", pd.Series(dtype=str)).value_counts(dropna=False)
        for label, count in overall.items():
            doc.add_paragraph(f"- {label or 'Non renseigné'} : {int(count)}")

        doc.add_heading("3. Position finale des institutions", level=2)
        final_pos = df.get("final_institutional_position", pd.Series(dtype=str)).value_counts(dropna=False)
        for label, count in final_pos.items():
            doc.add_paragraph(f"- {label or 'Non renseigné'} : {int(count)}")

        doc.add_heading("4. Points méthodologiques les plus réservés", level=2)
        strategic = response_count_table(df, STRATEGIC_KEYS, STRATEGIC_ROWS_FR)
        if not strategic.empty:
            top_reserved = strategic.sort_values(["Taux de non-validation (%)", "Taux de réserve (%)"], ascending=False).head(5)
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            hdr[0].text = "Élément"
            hdr[1].text = "Validé"
            hdr[2].text = "Sous réserve"
            hdr[3].text = "Non-validé"
            hdr[4].text = "Sans avis"
            for _, row in top_reserved.iterrows():
                cells = table.add_row().cells
                cells[0].text = str(row["élément"])
                cells[1].text = str(int(row["Validé"]))
                cells[2].text = str(int(row["Validé sous réserve"]))
                cells[3].text = str(int(row["Non-validé"]))
                cells[4].text = str(int(row["Sans avis"]))
        else:
            doc.add_paragraph("Aucune information disponible.")

        doc.add_heading("5. Domaines les plus réservés", level=2)
        domains = response_count_table(df, DOMAIN_KEYS, DOMAIN_ROWS_FR)
        if not domains.empty:
            top_domains = domains.sort_values(["Taux de non-validation (%)", "Taux de réserve (%)"], ascending=False).head(6)
            for _, row in top_domains.iterrows():
                doc.add_paragraph(
                    f"- {row['élément']} : {int(row['Validé sous réserve'])} sous réserve, {int(row['Non-validé'])} non-validés."
                )
        else:
            doc.add_paragraph("Aucune information disponible.")

        doc.add_heading("6. Extraits de justifications", level=2)
        justifs = extract_justifications(df, list(STRATEGIC_ROWS_FR) + list(DOMAIN_ROWS_FR) + ["overall_validation", "operational_usability", "final_institutional_position"], {
            **STRATEGIC_ROWS_FR,
            **DOMAIN_ROWS_FR,
            "overall_validation": "Validation globale du document",
            "operational_usability": "Document suffisamment opérationnel",
            "final_institutional_position": "Position finale de l’institution",
        })
        if not justifs.empty:
            for _, row in justifs.head(20).iterrows():
                doc.add_paragraph(
                    f"- {row['élément']} | {row['institution_acronym']} | {row['country_or_rec']} | {row['position']} : {row['justification']}"
                )
        else:
            doc.add_paragraph("Aucune justification textuelle disponible.")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def apply_filters(
    df: pd.DataFrame,
    statuses: Optional[Sequence[str]] = None,
    languages: Optional[Sequence[str]] = None,
    institution_types: Optional[Sequence[str]] = None,
    countries: Optional[Sequence[str]] = None,
    date_from: Optional[pd.Timestamp] = None,
    date_to: Optional[pd.Timestamp] = None,
) -> pd.DataFrame:
    out = df.copy()
    if out.empty:
        return out
    if statuses and "status" in out.columns:
        out = out[out["status"].isin(statuses)]
    if languages and "language" in out.columns:
        out = out[out["language"].isin(languages)]
    if institution_types and "institution_type" in out.columns:
        out = out[out["institution_type"].isin(institution_types)]
    if countries and "country_or_rec" in out.columns:
        out = out[out["country_or_rec"].isin(countries)]
    if date_from is not None and "submitted_at" in out.columns:
        out = out[out["submitted_at"] >= pd.Timestamp(date_from).tz_localize("UTC") if pd.Timestamp(date_from).tzinfo is None else out["submitted_at"] >= pd.Timestamp(date_from)]
    if date_to is not None and "submitted_at" in out.columns:
        end_ts = pd.Timestamp(date_to)
        if end_ts.tzinfo is None:
            # make the upper bound inclusive for the whole selected day
            end_ts = end_ts.tz_localize("UTC") + pd.Timedelta(days=1) - pd.Timedelta(microseconds=1)
        out = out[out["submitted_at"] <= end_ts]
    return out
