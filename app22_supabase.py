import streamlit as st
import pandas as pd
from datetime import datetime

# --- CONFIGURATION DE LA PAGE ---
st.set_page_config(page_title="Consultation STATAFRIC", layout="centered")

# On établit la connexion. Le système ira chercher les infos dans les "Secrets" de Streamlit Cloud


import streamlit as st
import pandas as pd
from datetime import datetime
import io
import json
import hashlib
import hmac
import secrets
import os
import re
import sqlite3
import zipfile
import uuid
import time
from datetime import datetime, timezone
from typing import Any, Dict, List, Tuple, Optional

import numpy as np
import pandas as pd
import streamlit as st

# Optional deps (Google Sheets / Dropbox)
try:
    import gspread  # type: ignore
    from google.oauth2.service_account import Credentials  # type: ignore
except Exception:
    gspread = None
    Credentials = None

try:
    import dropbox  # type: ignore
except Exception:
    dropbox = None


# Supabase (Postgres + Storage)
try:
    from supabase import create_client  # type: ignore
except Exception:
    create_client = None



# =========================
# Configuration
# =========================

APP_TITLE_FR = "Questionnaire de consultation pour l'identification des statistiques prioritaires"
APP_TITLE_EN = "Consultation questionnaire for identifying priority statistics"

DB_PATH = "responses.db"
LONG_LIST_CSV = os.path.join("data", "indicator_longlist.csv")
LONG_LIST_XLSX = os.path.join("data", "longlist.xlsx")


COUNTRY_XLSX = os.path.join("data", "COUNTRY_ISO3_with_EN.xlsx")
UK_FR = "NSP (Ne sais pas)"
UK_EN = "DNK (Do not know)"

# Version du scoring (pour compatibilité ascendante)
# v1 : ancien critère "gap" (écart) ; v2 : disponibilité inversée (Bonne=1) ; v3 : disponibilité directe (Bonne=3)
SCORING_VERSION = 3


# Scores affichés (notation multicritères)
# Barèmes de notation (0–3) par critère (scoring rationalisé)
# Remarque : 0 = NSP / DNK (ne sait pas) ; 1–3 = intensité croissante selon le libellé.
SCORE_SCALES = {
    "demand": {
        "fr": {0: "NSP", 1: "Faible", 2: "Moyenne", 3: "Élevée"},
        "en": {0: "DNK", 1: "Low", 2: "Medium", 3: "High"},
    },
    "availability": {
        "fr": {0: "NSP", 1: "Faible ou inexistante", 2: "Partielle", 3: "Bonne"},
        "en": {0: "DNK", 1: "Low or none", 2: "Partial", 3: "Good"},
    },
    "feasibility": {
        "fr": {0: "NSP", 1: "Difficile", 2: "Modérée", 3: "Facile"},
        "en": {0: "DNK", 1: "Difficult", 2: "Moderate", 3: "Easy"},
    },
}

def score_format(lang: str, criterion: str):
    """Formatter for score selectboxes (criterion-aware).

    We include a None option (placeholder) so we don't prefill answers.
    """
    placeholder_fr = "— Sélectionner —"
    placeholder_en = "— Select —"
    scale = SCORE_SCALES.get(criterion, SCORE_SCALES["demand"])
    mapping = scale["fr"] if lang == "fr" else scale["en"]

    def _fmt(v):
        if v is None or v == "":
            return placeholder_fr if lang == "fr" else placeholder_en
        try:
            iv = int(v)
        except Exception:
            return str(v)
        return mapping.get(iv, str(v))

    return _fmt


ROLE_OPTIONS_FR = [
    "DG/DGA/SG",
    "Directeur",
    "Conseiller",
    "Chef de division",
    "Chef de bureau",
    "Autre",
]
ROLE_OPTIONS_EN = [
    "DG/DGA/SG",
    "Director",
    "Advisor",
    "Head of division",
    "Head of office",
    "Other",
]


# =========================
# Helpers : i18n and state
# =========================

def t(lang: str, fr: str, en: str) -> str:
    return fr if lang == "fr" else en


def now_utc_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def get_query_params() -> Dict[str, List[str]]:
    """Compatibility across Streamlit versions."""
    try:
        # Streamlit >= 1.30
        qp = st.query_params  # type: ignore
        return {k: list(v) if isinstance(v, (list, tuple)) else [str(v)] for k, v in qp.items()}
    except Exception:
        try:
            return st.experimental_get_query_params()
        except Exception:
            return {}


def set_query_params(params: Dict[str, Any]) -> None:
    try:
        st.query_params.update(params)  # type: ignore
    except Exception:
        try:
            st.experimental_set_query_params(**params)
        except Exception:
            pass


def init_session() -> None:
    if "lang" not in st.session_state:
        st.session_state.lang = "fr"
    if "nav_idx" not in st.session_state:
        st.session_state.nav_idx = 0
    if "responses" not in st.session_state:
        st.session_state["responses"] = {}
    elif not isinstance(st.session_state.get("responses"), dict):
        st.session_state["responses"] = {}
    if "submission_id" not in st.session_state:
        st.session_state.submission_id = None
    if "admin_authed" not in st.session_state:
        st.session_state.admin_authed = False
    if "admin_role" not in st.session_state:
        st.session_state.admin_role = None  # "admin" | "superadmin"
    if "draft_id" not in st.session_state:
        st.session_state.draft_id = None

    if "draft_exists" not in st.session_state:
        st.session_state.draft_exists = False
    if "draft_resume_notice_shown" not in st.session_state:
        st.session_state.draft_resume_notice_shown = False
    if "draft_restored" not in st.session_state:
        st.session_state.draft_restored = False
    if "last_draft_save_ts" not in st.session_state:
        st.session_state.last_draft_save_ts = 0.0
    if "r12_substep" not in st.session_state:
        st.session_state["r12_substep"] = 0  # 0..2=open questions, 3=confirmation


def ensure_responses() -> None:
    """Garantit l’existence de st.session_state['responses'] (dict)."""
    if "responses" not in st.session_state or not isinstance(st.session_state.get("responses"), dict):
        st.session_state["responses"] = {}


def resp_get(key: str, default=None):
    ensure_responses()
    return st.session_state["responses"].get(key, default)


def resp_set(key: str, value) -> None:
    ensure_responses()
    st.session_state["responses"][key] = value



def normalize_availability(v_raw: Any, scoring_version: Any) -> int:
    """Normalise la disponibilité sur l'échelle 'Bonne=3' (SCORING_VERSION=3).

    - v3+ : on conserve la valeur telle quelle (0–3).
    - v1/v2 ou absence de version : on inverse (1<->3) car l'ancien codage correspondait à un "écart" / ou à une disponibilité inversée.
    """
    try:
        iv = int(v_raw)
    except Exception:
        return 0
    if iv == 0:
        return 0

    try:
        ver = int(scoring_version)
    except Exception:
        ver = 0

    if ver >= SCORING_VERSION:
        return iv

    # Inversion pour les versions antérieures : 1<->3, 2 inchangé
    if iv in (1, 2, 3):
        return 4 - iv
    return iv

def ensure_draft_id() -> Optional[str]:
    """Ensure a stable draft id exists (used for mobile resume)."""
    if st.session_state.get("draft_id"):
        return st.session_state.draft_id
    email = (resp_get("email", "") or "").strip()
    if not email:
        return None
    draft_id = str(uuid.uuid4())
    st.session_state.draft_id = draft_id
    # Keep any existing query params (admin, etc.)
    try:
        qp = get_query_params()
        qp["rid"] = [draft_id]
        set_query_params({k: v[0] if len(v) == 1 else v for k, v in qp.items()})
    except Exception:
        pass
    return draft_id


def autosave_draft(force: bool = False) -> Tuple[bool, str]:
    """Persist current responses to DB to mitigate mobile refresh/resets."""
    draft_id = st.session_state.get("draft_id")
    email = (resp_get("email", "") or "").strip()
    if not draft_id or not email:
        return False, "no_draft_id_or_email"
    now_ts = time.time()
    last_ts = float(st.session_state.get("last_draft_save_ts", 0.0) or 0.0)
    if (not force) and (now_ts - last_ts < 2.0):
        return True, "skipped_rate_limit"
    payload = {
        "responses": st.session_state.responses,
        "nav_idx": int(st.session_state.get("nav_idx", 0)),
        "lang": st.session_state.get("lang", "fr"),
    }
    try:
        db_save_draft(draft_id, email, payload)
        st.session_state.last_draft_save_ts = now_ts
        return True, "saved"
    except Exception as e:
        return False, str(e)


def maybe_restore_draft() -> None:
    """Restore responses from DB if URL contains rid and session is empty."""
    if st.session_state.get("draft_restored"):
        return
    st.session_state.draft_restored = True

    qp = get_query_params()
    rid = None
    if "rid" in qp and qp["rid"]:
        rid = qp["rid"][0]
    if not rid:
        return

    # Do not restore while in admin mode
    if "admin" in qp and qp["admin"] and qp["admin"][0] in ["1", "true", "yes"]:
        return

    payload = db_load_draft(rid)
    st.session_state.draft_exists = bool(payload)

    # Restore only if session is empty (avoid overriding ongoing input)
    if st.session_state.get("responses") and len(st.session_state.responses) > 0:
        st.session_state.draft_id = rid
        return


    if not payload:
        st.session_state.draft_id = rid
        return

    responses = payload.get("responses", {})
    if isinstance(responses, dict):
        st.session_state.responses = responses

    try:
        st.session_state.nav_idx = int(payload.get("nav_idx", 0))
    except Exception:
        st.session_state.nav_idx = 0

    lang = payload.get("lang", None)
    if lang in ["fr", "en"]:
        st.session_state.lang = lang

    st.session_state.draft_id = rid




# =========================
# Data : longlist loader
# =========================

@st.cache_data(show_spinner=False)
def load_longlist() -> pd.DataFrame:
    """
    Charge la longlist (statistiques par domaine) depuis :
    - CSV : data/indicator_longlist.csv (prioritaire)
    - XLSX : data/longlist.xlsx (fallback)
    Tolère aussi les fichiers placés à la racine du dépôt.

    Si aucun fichier n'est trouvé, l'application démarre quand même,
    mais les listes déroulantes de la Rubrique 4/5 seront vides.
    """
    csv_candidates = [
        LONG_LIST_CSV,
        "indicator_longlist.csv",
        os.path.join(".", "indicator_longlist.csv"),
        os.path.join(".", "data", "indicator_longlist.csv"),
    ]
    xlsx_candidates = [
        LONG_LIST_XLSX,
        "longlist.xlsx",
        os.path.join(".", "longlist.xlsx"),
        os.path.join(".", "data", "longlist.xlsx"),
    ]

    # 1) CSV (prioritaire si la traduction EN est suffisamment complète)
    df_csv = None
    df_csv_path = None
    for p in csv_candidates:
        if os.path.exists(p):
            df_csv = pd.read_csv(p, dtype=str).fillna("")
            df_csv_path = p
            break

    if df_csv is not None:
        # Sanity check : si beaucoup de libellés EN sont vides, on préfère l'XLSX (souvent plus à jour)
        if "stat_label_en" in df_csv.columns:
            miss_ratio = (df_csv["stat_label_en"].astype(str).str.strip() == "").mean()
        else:
            miss_ratio = 1.0

        if miss_ratio <= 0.20:
            df_csv.attrs["source_path"] = df_csv_path
            return df_csv

    # 2) XLSX (format utilisateur) (format utilisateur)
    for p in xlsx_candidates:
        if os.path.exists(p):
            df = pd.read_excel(p, dtype=str).fillna("")
            df.attrs["source_path"] = p

            # Colonnes attendues (minimum) : Domain_code, Domain_label_fr, Stat_label_fr
            if set(["Domain_code", "Domain_label_fr", "Stat_label_fr"]).issubset(df.columns):
                out = pd.DataFrame()
                out["domain_code"] = df["Domain_code"].astype(str).str.strip()

                # Labels FR (on retire le préfixe code "D01|...")
                out["domain_label_fr"] = df["Domain_label_fr"].astype(str).str.split("|", n=1).str[-1].str.strip()
                out["stat_code"] = df["Stat_label_fr"].astype(str).str.split("|", n=1).str[0].str.strip()
                out["stat_label_fr"] = df["Stat_label_fr"].astype(str).str.split("|", n=1).str[-1].str.strip()

                # Labels EN si disponibles, sinon fallback FR
                if "Domain_label_en" in df.columns:
                    out["domain_label_en"] = df["Domain_label_en"].astype(str).str.split("|", n=1).str[-1].str.strip()
                else:
                    out["domain_label_en"] = out["domain_label_fr"]

                if "Stat_label_en" in df.columns:
                    out["stat_label_en"] = df["Stat_label_en"].astype(str).str.split("|", n=1).str[-1].str.strip()
                else:
                    out["stat_label_en"] = out["stat_label_fr"]

                out.attrs["source_path"] = p
                return out[[
                    "domain_code",
                    "domain_label_fr",
                    "domain_label_en",
                    "stat_code",
                    "stat_label_fr",
                    "stat_label_en",
                ]]


    # Fallback final : si un CSV a été trouvé (même avec traduction EN incomplète), on le renvoie
    if df_csv is not None:
        df_csv.attrs["source_path"] = df_csv_path or ""
        return df_csv
# Aucun fichier trouvé : dataframe vide
    empty = pd.DataFrame(columns=[
        "domain_code",
        "domain_label_fr",
        "domain_label_en",
        "stat_code",
        "stat_label_fr",
        "stat_label_en",
    ])
    empty.attrs["source_path"] = ""
    return empty



# =========================
# Data : countries loader
# =========================

@st.cache_data(show_spinner=False)
def load_countries() -> pd.DataFrame:
    """
    Charge la liste des pays (ISO3 + noms FR/EN) depuis :
    - data/COUNTRY_ISO3_with_EN.xlsx (prioritaire)
    Tolère aussi le fichier placé à la racine du dépôt.

    Colonnes attendues (au minimum) :
    - COUNTRY_ISO3
    - COUNTRY_NAME_FR
    - COUNTRY_NAME_EN
    """
    candidates = [
        COUNTRY_XLSX,
        "COUNTRY_ISO3_with_EN.xlsx",
        os.path.join(".", "COUNTRY_ISO3_with_EN.xlsx"),
        os.path.join(".", "data", "COUNTRY_ISO3_with_EN.xlsx"),
        # Fallbacks (ancien nom éventuel)
        "COUNTRY_ISO3.xlsx",
        os.path.join(".", "COUNTRY_ISO3.xlsx"),
        os.path.join(".", "data", "COUNTRY_ISO3.xlsx"),
    ]
    for p in candidates:
        if os.path.exists(p):
            try:
                df = pd.read_excel(p, dtype=str).fillna("")
                # Normalisation des noms de colonnes
                df.columns = [str(c).strip() for c in df.columns]
                if "COUNTRY_ISO3" not in df.columns:
                    continue
                if "COUNTRY_NAME_FR" not in df.columns:
                    df["COUNTRY_NAME_FR"] = ""
                if "COUNTRY_NAME_EN" not in df.columns:
                    df["COUNTRY_NAME_EN"] = ""
                df["COUNTRY_ISO3"] = df["COUNTRY_ISO3"].astype(str).str.strip().str.upper()
                df["COUNTRY_NAME_FR"] = df["COUNTRY_NAME_FR"].astype(str).str.strip()
                df["COUNTRY_NAME_EN"] = df["COUNTRY_NAME_EN"].astype(str).str.strip()
                df = df[df["COUNTRY_ISO3"] != ""].copy()
                df.attrs["source_path"] = p
                return df
            except Exception:
                continue
    empty = pd.DataFrame(columns=["COUNTRY_ISO3", "COUNTRY_NAME_FR", "COUNTRY_NAME_EN"])
    empty.attrs["source_path"] = ""
    return empty


def country_maps(df_c: pd.DataFrame) -> Tuple[List[str], Dict[str, str], Dict[str, str]]:
    """Retourne (iso3_list, iso3_to_fr, iso3_to_en)."""
    if df_c is None or df_c.empty:
        return [], {}, {}
    iso3_to_fr: Dict[str, str] = {}
    iso3_to_en: Dict[str, str] = {}
    for _, r in df_c.iterrows():
        iso3 = str(r.get("COUNTRY_ISO3", "")).strip().upper()
        if not iso3:
            continue
        iso3_to_fr[iso3] = str(r.get("COUNTRY_NAME_FR", "")).strip()
        iso3_to_en[iso3] = str(r.get("COUNTRY_NAME_EN", "")).strip()
    iso3_list = sorted(set(list(iso3_to_fr.keys()) + list(iso3_to_en.keys())))
    return iso3_list, iso3_to_fr, iso3_to_en


def country_label(iso3: str, lang: str, iso3_to_fr: Dict[str, str], iso3_to_en: Dict[str, str]) -> str:
    if not iso3:
        return ""
    if lang == "en":
        return (iso3_to_en.get(iso3) or iso3_to_fr.get(iso3) or iso3).strip()
    return (iso3_to_fr.get(iso3) or iso3_to_en.get(iso3) or iso3).strip()


def domains_from_longlist(df_long: pd.DataFrame, lang: str) -> List[Tuple[str, str]]:
    if df_long.empty:
        return []
    col = "domain_label_fr" if lang == "fr" else "domain_label_en"
    tmp = df_long[["domain_code", col]].drop_duplicates().sort_values(["domain_code", col])
    return [(r["domain_code"], r[col]) for _, r in tmp.iterrows()]


def stats_for_domain(df_long: pd.DataFrame, domain_code: str, lang: str) -> List[Tuple[str, str]]:
    if df_long.empty or not domain_code:
        return []
    col = "stat_label_fr" if lang == "fr" else "stat_label_en"
    tmp = df_long[df_long["domain_code"] == domain_code][["stat_code", col]].drop_duplicates()
    tmp = tmp.sort_values(["stat_code", col])
    return [(r["stat_code"], (str(r[col]).strip() if str(r[col]).strip() else r["stat_code"])) for _, r in tmp.iterrows()]


# =========================
# Stockage : Supabase (Postgres) + JSON (optionnel)
# =========================


def domain_label_map(df_long: pd.DataFrame, lang: str) -> Dict[str, str]:
    """Map domain_code -> label in selected language."""
    if df_long is None or df_long.empty:
        return {}
    col = "domain_label_fr" if lang == "fr" else "domain_label_en"
    m = {}
    for _, r in df_long.drop_duplicates("domain_code").iterrows():
        code = str(r["domain_code"])
        lbl = str(r.get(col, "")).strip()
        if not lbl:
            lbl = str(r.get("domain_label_fr", "")).strip() or code
        m[code] = lbl
    return m

def stat_label_map(df_long: pd.DataFrame, lang: str) -> Dict[str, str]:
    """Map stat_code -> label in selected language."""
    if df_long is None or df_long.empty:
        return {}
    col = "stat_label_fr" if lang == "fr" else "stat_label_en"
    m = {}
    for _, r in df_long.drop_duplicates("stat_code").iterrows():
        code = str(r["stat_code"])
        lbl = str(r.get(col, "")).strip()
        if not lbl:
            lbl = str(r.get("stat_label_fr", "")).strip() or code
        m[code] = lbl
    return m

def build_publication_report_docx(lang: str, filtered_payloads: pd.DataFrame, by_domain: pd.DataFrame, by_stat: pd.DataFrame, scored_rows: pd.DataFrame) -> bytes:
    """
    Génère un rapport Word 'publication' enrichi :
    - profil des répondants
    - domaines TOP 5 (fréquences)
    - tableau agrégé des statistiques et scores moyens
    - graphiques (bar charts)
    - annexes
    """
    from docx import Document
    from docx.shared import Inches
    import matplotlib.pyplot as plt

    doc = Document()
    title = t(lang, "Rapport de synthèse – Consultation sur les statistiques prioritaires", "Summary report – Consultation on priority statistics")
    doc.add_heading(title, level=0)
    doc.add_paragraph(t(lang, f"Date : {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", f"Date: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"))
    doc.add_paragraph("")

    # Sample profile
    doc.add_heading(t(lang, "Profil des répondants", "Respondent profile"), level=1)
    n = len(filtered_payloads)
    doc.add_paragraph(t(lang, f"Nombre de réponses analysées : {n}", f"Number of responses analyzed: {n}"))

    # Countries
    if "pays" in filtered_payloads.columns:
        vc = filtered_payloads["pays"].fillna("").astype(str)
        vc = vc[vc.str.strip() != ""].value_counts().head(10)
        if len(vc) > 0:
            doc.add_paragraph(t(lang, "Top pays (10 premiers) :", "Top countries (top 10):"))
            for k, v in vc.items():
                doc.add_paragraph(f"- {k} : {v}", style=None)

    # Actor types
    if "type_acteur" in filtered_payloads.columns:
        vc = filtered_payloads["type_acteur"].fillna("").astype(str)
        vc = vc[vc.str.strip() != ""].value_counts()
        if len(vc) > 0:
            doc.add_paragraph(t(lang, "Répartition par type d’acteur :", "Distribution by stakeholder type:"))
            for k, v in vc.items():
                doc.add_paragraph(f"- {k} : {v}", style=None)

    # Domain aggregation
    doc.add_heading(t(lang, "Domaines prioritaires (scores moyens)", "Priority domains (mean scores)"), level=1)
    top_dom = by_domain.head(15).copy()
    # Table
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = t(lang, "Domaine", "Domain")
    hdr[1].text = t(lang, "Nb. soumissions", "Submissions")
    hdr[2].text = t(lang, "Nb. stats notées", "Scored indicators")
    hdr[3].text = t(lang, "Score moyen", "Mean score")
    for _, r in top_dom.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["domain_label"])
        row[1].text = str(int(r["n_submissions"]))
        row[2].text = str(int(r["n_stats"]))
        row[3].text = f"{float(r['mean_overall']):.2f}"

    # Chart domain
    try:
        fig = plt.figure()
        plt.bar(top_dom["domain_label"], top_dom["mean_overall"])
        plt.xticks(rotation=75, ha="right")
        plt.ylabel(t(lang, "Score moyen", "Mean score"))
        plt.tight_layout()
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format="png", dpi=150)
        plt.close(fig)
        img_stream.seek(0)
        doc.add_paragraph(t(lang, "Graphique : score moyen par domaine (top 15).", "Chart: mean score by domain (top 15)."))
        doc.add_picture(img_stream, width=Inches(6.5))
    except Exception:
        pass

    # Statistic aggregation
    doc.add_heading(t(lang, "Statistiques prioritaires (scores moyens)", "Priority indicators (mean scores)"), level=1)
    top_stat = by_stat.sort_values(["mean_overall", "n"], ascending=[False, False]).head(30).copy()
    table2 = doc.add_table(rows=1, cols=6)
    h = table2.rows[0].cells
    h[0].text = t(lang, "Domaine", "Domain")
    h[1].text = t(lang, "Statistique", "Indicator")
    h[2].text = t(lang, "N", "N")
    h[3].text = t(lang, "Demande", "Demand")
    h[4].text = t(lang, "Disponibilité", "Availability")
    h[5].text = t(lang, "Faisabilité", "Feasibility")
    for _, r in top_stat.iterrows():
        row = table2.add_row().cells
        row[0].text = str(r["domain_label"])
        row[1].text = str(r["stat_label"])
        row[2].text = str(int(r["n"]))
        row[3].text = f"{float(r['mean_demand']):.2f}"
        row[4].text = f"{float(r['mean_availability']):.2f}"
        row[5].text = f"{float(r['mean_feasibility']):.2f}"


    # Chart top overall indicators
    try:
        fig = plt.figure()
        plt.barh(top_stat["stat_label"].iloc[::-1], top_stat["mean_overall"].iloc[::-1])
        plt.xlabel(t(lang, "Score moyen", "Mean score"))
        plt.tight_layout()
        img2 = io.BytesIO()
        plt.savefig(img2, format="png", dpi=150)
        plt.close(fig)
        img2.seek(0)
        doc.add_paragraph(t(lang, "Graphique : score moyen par statistique (top 30).", "Chart: mean score by indicator (top 30)."))
        doc.add_picture(img2, width=Inches(6.5))
    except Exception:
        pass

    # Interpretation auto
    doc.add_heading(t(lang, "Interprétations automatiques", "Automatic interpretations"), level=1)
    # Simple rules
    best_dom = top_dom.iloc[0]
    doc.add_paragraph(
        t(
            lang,
            f"Le domaine le mieux noté est « {best_dom['domain_label']} » avec un score moyen de {best_dom['mean_overall']:.2f} (sur 3).",
            f"The highest-rated domain is “{best_dom['domain_label']}” with a mean score of {best_dom['mean_overall']:.2f} (out of 3)."
        )
    )
    best_stat = top_stat.iloc[0]
    doc.add_paragraph(
        t(
            lang,
            f"La statistique la mieux notée est « {best_stat['stat_label']} » (domaine : {best_stat['domain_label']}) avec un score moyen de {best_stat['mean_overall']:.2f}.",
            f"The highest-rated indicator is “{best_stat['stat_label']}” (domain: {best_stat['domain_label']}) with a mean score of {best_stat['mean_overall']:.2f}."
        )
    )

    # Annexes
    doc.add_heading(t(lang, "Annexes", "Annexes"), level=1)
    doc.add_paragraph(t(lang, "A1. Tableau détaillé (statistiques agrégées) – extrait", "A1. Detailed table (aggregated indicators) – excerpt"))
    annex = by_stat.head(50).copy()
    tab3 = doc.add_table(rows=1, cols=5)
    hh = tab3.rows[0].cells
    hh[0].text = t(lang, "Domaine", "Domain")
    hh[1].text = t(lang, "Statistique", "Indicator")
    hh[2].text = t(lang, "N", "N")
    hh[3].text = t(lang, "Score moyen", "Mean score")
    hh[4].text = t(lang, "Détail", "Detail")
    for _, r in annex.iterrows():
        rr = tab3.add_row().cells
        rr[0].text = str(r["domain_label"])
        rr[1].text = str(r["stat_label"])
        rr[2].text = str(int(r["n"]))
        rr[3].text = f"{float(r['mean_overall']):.2f}"
        if lang == "fr":
            rr[4].text = f"Demande={float(r['mean_demand']):.2f}, Disponibilité={float(r['mean_availability']):.2f}, Faisabilité={float(r['mean_feasibility']):.2f}"
        else:
            rr[4].text = f"Demand={float(r['mean_demand']):.2f}, Availability={float(r['mean_availability']):.2f}, Feasibility={float(r['mean_feasibility']):.2f}"


    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()





# =========================
# Stockage : Supabase (Postgres) + JSON (optionnel : Supabase Storage)
# =========================

SUPABASE_SUBMISSIONS_TABLE = os.environ.get("SUPABASE_SUBMISSIONS_TABLE", "submissions")
SUPABASE_DRAFTS_TABLE = os.environ.get("SUPABASE_DRAFTS_TABLE", "drafts")
SUPABASE_CONFIG_TABLE = os.environ.get("SUPABASE_CONFIG_TABLE", "app_config")
SUPABASE_STORAGE_BUCKET = os.environ.get("SUPABASE_STORAGE_BUCKET", "submissions-json")


def _sb_get(name: str) -> Optional[str]:
    """Get value from Streamlit secrets or env, without depending on _get_secret_or_env ordering."""
    # Streamlit secrets
    try:
        secrets_obj = getattr(st, "secrets", None)
        if secrets_obj is not None:
            try:
                if hasattr(secrets_obj, "__contains__") and name in secrets_obj:
                    v = secrets_obj[name]
                    if v not in (None, ""):
                        return str(v)
            except Exception:
                pass
            try:
                d = secrets_obj.to_dict() if hasattr(secrets_obj, "to_dict") else dict(secrets_obj)
            except Exception:
                d = {}
            if isinstance(d, dict):
                if name in d and d.get(name) not in (None, ""):
                    return str(d.get(name))
                for _k, _v in d.items():
                    if isinstance(_v, dict) and name in _v and _v.get(name) not in (None, ""):
                        return str(_v.get(name))
    except Exception:
        pass

    # Env
    v = os.environ.get(name, None)
    if v not in (None, ""):
        return str(v)
    return None


@st.cache_resource(show_spinner=False)
def _sb_client():
    """Supabase client cached per-process."""
    if create_client is None:
        return None
    url = _sb_get("SUPABASE_URL")
    # Prefer service role for server-side apps (kept as secret on Render).
    key = _sb_get("SUPABASE_SERVICE_ROLE_KEY") or _sb_get("SUPABASE_SERVICE_KEY") or _sb_get("SUPABASE_ANON_KEY")
    if not url or not key:
        return None
    try:
        return create_client(url, key)
    except Exception:
        return None


@st.cache_resource(show_spinner=False)
def _sb_runtime_state() -> Dict[str, Any]:
    return {"last_cleanup_ts": 0.0}


def _sb_maybe_cleanup() -> None:
    """Best-effort deletion of rows older than retention window.

    Note : this is an in-app fallback. For production, prefer a DB cron / scheduler.
    """
    try:
        retention_days = int(_sb_get("DATA_RETENTION_DAYS") or "60")
    except Exception:
        retention_days = 60

    # throttle to at most once per 12h per instance
    state = _sb_runtime_state()
    now_ts = time.time()
    if now_ts - float(state.get("last_cleanup_ts", 0.0)) < 12 * 3600:
        return

    sb = _sb_client()
    if sb is None:
        return

    # Compute threshold ISO
    try:
        threshold = (datetime.utcnow() - pd.Timedelta(days=retention_days)).replace(tzinfo=timezone.utc).isoformat()
    except Exception:
        threshold = datetime.utcnow().replace(tzinfo=timezone.utc).isoformat()

    # Best-effort deletes (ignore failures)
    try:
        sb.table(SUPABASE_SUBMISSIONS_TABLE).delete().lt("submitted_at_utc", threshold).execute()
    except Exception:
        pass
    try:
        sb.table(SUPABASE_DRAFTS_TABLE).delete().lt("updated_at_utc", threshold).execute()
    except Exception:
        pass


    # Nettoyage best-effort des fichiers JSON dans Supabase Storage (si bucket présent)
    try:
        bucket = _sb_get("SUPABASE_STORAGE_BUCKET") or SUPABASE_STORAGE_BUCKET
        if bucket:
            # Liste des objets dans le dossier 'submissions'
            res = sb.storage.from_(bucket).list(
                "submissions",
                {"limit": 1000, "offset": 0, "sortBy": {"column": "created_at", "order": "asc"}},
            )
            items = getattr(res, "data", res)
            to_remove = []
            if isinstance(items, list):
                for it in items:
                    if not isinstance(it, dict):
                        continue
                    name = it.get("name")
                    created_at = it.get("created_at") or it.get("updated_at")
                    # Comparaison lexicographique OK pour ISO8601 (timestamptz)
                    if name and created_at and str(created_at) < str(threshold):
                        to_remove.append(f"submissions/{name}")
            if to_remove:
                sb.storage.from_(bucket).remove(to_remove)
    except Exception:
        pass

    state["last_cleanup_ts"] = now_ts


def _sb_upload_submission_json(submission_id: str, payload: Dict[str, Any]) -> Tuple[bool, str]:
    """Upload one JSON file to Supabase Storage (optional)."""
    sb = _sb_client()
    if sb is None:
        return (False, "Supabase non configuré")
    bucket = _sb_get("SUPABASE_STORAGE_BUCKET") or SUPABASE_STORAGE_BUCKET
    if not bucket:
        return (False, "Bucket non configuré")
    try:
        path = f"submissions/{submission_id}.json"
        content = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
        # Python storage API : storage.from_(bucket).upload(path=..., file=...)
        sb.storage.from_(bucket).upload(
            path=path,
            file=content,
            file_options={"content-type": "application/json", "upsert": "true"},
        )
        return (True, f"JSON envoyé : {bucket}/{path}")
    except Exception as e:
        return (False, f"Upload JSON échoué : {e}")

def db_init() -> None:
    """Initialise le backend Supabase.

    Les tables doivent être créées côté Supabase (voir procédure).
    """
    _sb_maybe_cleanup()
    return

# =========================
# Admin auth helpers
# =========================

PBKDF2_ITERS = 200_000


def db_get_config(k: str) -> Optional[str]:
    sb = _sb_client()
    if sb is None:
        return None
    try:
        r = sb.table(SUPABASE_CONFIG_TABLE).select("v").eq("k", k).limit(1).execute()
        data = getattr(r, "data", None) or []
        if isinstance(data, list) and data:
            return data[0].get("v")
    except Exception:
        return None
    return None


def db_set_config(k: str, v: str) -> None:
    sb = _sb_client()
    if sb is None:
        return
    try:
        sb.table(SUPABASE_CONFIG_TABLE).upsert(
            {"k": k, "v": v, "updated_at_utc": now_utc_iso()},
            on_conflict="k",
        ).execute()
    except Exception:
        return


def db_delete_config(k: str) -> None:
    sb = _sb_client()
    if sb is None:
        return
    try:
        sb.table(SUPABASE_CONFIG_TABLE).delete().eq("k", k).execute()
    except Exception:
        return

def _pbkdf2_sha256_hex(password: str, salt: bytes, iterations: int = PBKDF2_ITERS) -> str:
    dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, iterations)
    return dk.hex()


def _safe_eq(a: str, b: str) -> bool:
    try:
        return hmac.compare_digest(a or "", b or "")
    except Exception:
        return (a or "") == (b or "")


def _get_secret_or_env(name: str) -> Tuple[Optional[str], Optional[str]]:
    """Return (value, source) where source is 'secrets' | 'env' | None.

    Robust to different Streamlit versions and secrets layouts:
    - Direct key : ADMIN_PASSWORD = "..."
    - Nested table : [general] ADMIN_PASSWORD = "..."
    """
    val: Optional[str] = None
    src: Optional[str] = None

    # 1) Streamlit secrets
    try:
        secrets_obj = getattr(st, "secrets", None)
        if secrets_obj is not None:
            # Direct access (preferred)
            try:
                if hasattr(secrets_obj, "__contains__") and name in secrets_obj:
                    v = secrets_obj[name]
                    if v not in (None, ""):
                        val = str(v)
                        src = "secrets"
            except Exception:
                pass

            # Fallback : convert to dict then search (supports nested sections)
            if not val:
                try:
                    d = secrets_obj.to_dict() if hasattr(secrets_obj, "to_dict") else dict(secrets_obj)
                except Exception:
                    d = {}

                if isinstance(d, dict):
                    if name in d and d.get(name) not in (None, ""):
                        val = str(d.get(name))
                        src = "secrets"
                    else:
                        # Search nested dicts
                        for _k, _v in d.items():
                            if isinstance(_v, dict) and name in _v and _v.get(name) not in (None, ""):
                                val = str(_v.get(name))
                                src = "secrets"
                                break
    except Exception:
        # Leave val/src as None
        pass

    # 2) Environment variable (only if not found in secrets)
    if not val:
        env_val = os.environ.get(name, None)
        if env_val not in (None, ""):
            val = str(env_val)
            src = "env"

    return (val, src)


def get_admin_auth_source() -> Tuple[str, str]:
    """Human-readable indicator of current admin password source."""
    h = db_get_config("ADMIN_PASSWORD_HASH")
    s = db_get_config("ADMIN_PASSWORD_SALT")
    if h and s:
        return ("db", t(st.session_state.get("lang", "fr"), "base (haché)", "database (hashed)"))
    v, src = _get_secret_or_env("ADMIN_PASSWORD")
    if v and src:
        return (src, t(st.session_state.get("lang", "fr"), src, src))
    return ("none", t(st.session_state.get("lang", "fr"), "non configuré", "not configured"))


def verify_admin_password(pw: str) -> bool:
    pw = pw or ""
    h = db_get_config("ADMIN_PASSWORD_HASH")
    s = db_get_config("ADMIN_PASSWORD_SALT")
    it = db_get_config("ADMIN_PASSWORD_ITERS")
    if h and s:
        try:
            salt = bytes.fromhex(s)
            iterations = int(it) if it else PBKDF2_ITERS
            calc = _pbkdf2_sha256_hex(pw, salt, iterations)
            return _safe_eq(calc, h)
        except Exception:
            return False

    expected, _src = _get_secret_or_env("ADMIN_PASSWORD")
    if expected:
        return _safe_eq(pw, str(expected))
    return False


def verify_superadmin_password(pw: str) -> bool:
    pw = pw or ""
    expected, _src = _get_secret_or_env("SUPERADMIN_PASSWORD")
    if expected:
        return _safe_eq(pw, str(expected))
    return False


def set_admin_password(new_pw: str) -> None:
    """Set (hashed) admin password override in DB."""
    new_pw = (new_pw or "").strip()
    if not new_pw:
        raise ValueError("empty password")
    salt = secrets.token_bytes(16)
    h = _pbkdf2_sha256_hex(new_pw, salt, PBKDF2_ITERS)
    db_set_config("ADMIN_PASSWORD_HASH", h)
    db_set_config("ADMIN_PASSWORD_SALT", salt.hex())
    db_set_config("ADMIN_PASSWORD_ITERS", str(PBKDF2_ITERS))


def reset_admin_password_to_secrets_env() -> None:
    """Remove DB override so app falls back to secrets/env."""
    db_delete_config("ADMIN_PASSWORD_HASH")
    db_delete_config("ADMIN_PASSWORD_SALT")
    db_delete_config("ADMIN_PASSWORD_ITERS")


def db_email_exists(email: str) -> bool:
    email = (email or "").strip().lower()
    if not email:
        return False
    sb = _sb_client()
    if sb is None:
        return False
    try:
        r = sb.table(SUPABASE_SUBMISSIONS_TABLE).select("submission_id").eq("email", email).limit(1).execute()
        data = getattr(r, "data", None) or []
        return bool(data)
    except Exception:
        return False


def db_save_submission(submission_id: str, lang: str, email: str, payload: Dict[str, Any]) -> Tuple[bool, str, bool, str]:
    """Enregistre la soumission dans Supabase Postgres (jsonb) et, si possible, dépose un fichier JSON dans Supabase Storage."""
    email = (email or "").strip().lower()
    sb = _sb_client()
    if sb is None:
        return (False, "Supabase non configuré (SUPABASE_URL / SUPABASE_SERVICE_ROLE_KEY).", False, "JSON non envoyé")
    row = {
        "submission_id": submission_id,
        "submitted_at_utc": now_utc_iso(),
        "lang": lang,
        "email": email,
        "payload": payload,  # jsonb
    }
    try:
        sb.table(SUPABASE_SUBMISSIONS_TABLE).insert(row).execute()
        ok_db = True
        msg_db = "Enregistré dans Supabase."
    except Exception as e:
        ok_db = False
        msg_db = f"Échec d'enregistrement Supabase : {e}"
        return (ok_db, msg_db, False, "JSON non envoyé")

    json_ok, json_msg = _sb_upload_submission_json(submission_id, payload)
    return (ok_db, msg_db, json_ok, json_msg)


def db_save_draft(draft_id: str, email: str, payload: Dict[str, Any]) -> None:
    draft_id = (draft_id or "").strip()
    email = (email or "").strip().lower()
    if not draft_id or not email:
        return
    sb = _sb_client()
    if sb is None:
        return
    try:
        sb.table(SUPABASE_DRAFTS_TABLE).upsert(
            {"draft_id": draft_id, "updated_at_utc": now_utc_iso(), "email": email, "payload": payload},
            on_conflict="draft_id",
        ).execute()
    except Exception:
        return


def db_load_draft(draft_id: str) -> Optional[Dict[str, Any]]:
    draft_id = (draft_id or "").strip()
    if not draft_id:
        return None
    sb = _sb_client()
    if sb is None:
        return None
    try:
        r = sb.table(SUPABASE_DRAFTS_TABLE).select("payload").eq("draft_id", draft_id).limit(1).execute()
        data = getattr(r, "data", None) or []
        if isinstance(data, list) and data:
            return data[0].get("payload") or None
    except Exception:
        return None
    return None


def db_delete_draft(draft_id: str) -> None:
    draft_id = (draft_id or "").strip()
    if not draft_id:
        return
    sb = _sb_client()
    if sb is None:
        return
    try:
        sb.table(SUPABASE_DRAFTS_TABLE).delete().eq("draft_id", draft_id).execute()
    except Exception:
        return


def db_read_submissions(limit: int = 2000) -> pd.DataFrame:
    sb = _sb_client()
    if sb is None:
        return pd.DataFrame(columns=["submission_id", "submitted_at_utc", "lang", "email", "payload_json"])
    try:
        r = (
            sb.table(SUPABASE_SUBMISSIONS_TABLE)
            .select("submission_id,submitted_at_utc,lang,email,payload")
            .order("submitted_at_utc", desc=True)
            .limit(limit)
            .execute()
        )
        data = getattr(r, "data", None) or []
        rows = []
        for row in (data if isinstance(data, list) else []):
            payload = row.get("payload", {})
            rows.append(
                {
                    "submission_id": row.get("submission_id"),
                    "submitted_at_utc": row.get("submitted_at_utc"),
                    "lang": row.get("lang"),
                    "email": row.get("email"),
                    "payload_json": json.dumps(payload, ensure_ascii=False),
                }
            )
        return pd.DataFrame(rows, columns=["submission_id", "submitted_at_utc", "lang", "email", "payload_json"])
    except Exception:
        return pd.DataFrame(columns=["submission_id", "submitted_at_utc", "lang", "email", "payload_json"])


def db_dump_csv_bytes(limit: int = 2000000) -> bytes:
    """Export des soumissions (Supabase) en CSV (octets)."""
    df = db_read_submissions(limit=limit)
    return df.to_csv(index=False).encode("utf-8-sig")

def flatten_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Create a 'flat' row for exports / Google Sheets (comprehensive).
    - Keeps keys stable across FR/EN by mapping table items to canonical ids.
    - Serializes list/dict fields into '; ' / JSON strings as needed.
    """
    def _join_list(v: Any) -> str:
        if isinstance(v, list):
            return "; ".join([str(x) for x in v if x is not None and str(x).strip() != ""])
        return ""

    def _json(v: Any) -> str:
        try:
            return json.dumps(v, ensure_ascii=False)
        except Exception:
            return ""

    # Canonical mappings for table questions (FR/EN)
    GENDER_ITEM_MAP = {
        "Désagrégation par sexe": "sex",
        "Disaggregation by sex": "sex",
        "Sexe": "sex",
        "Sex": "sex",
        "Désagrégation par âge": "age",
        "Disaggregation by age": "age",
        "Âge": "age",
        "Age": "age",
        "Milieu urbain / rural": "urban_rural",
        "Urban / rural": "urban_rural",
        "Milieu urbain/rural": "urban_rural",
        "Urban/rural residence": "urban_rural",
        "Handicap": "disability",
        "Disability": "disability",
        "Quintile de richesse": "wealth_quintile",
        "Wealth quintile": "wealth_quintile",
        "Violences basées sur le genre (VBG)": "gbv",
        "Gender-based violence (GBV)": "gbv",
        "Temps domestique non rémunéré": "unpaid_domestic",
        "Unpaid domestic work": "unpaid_domestic",
    }
    CAPACITY_ITEM_MAP = {
        "Compétences statistiques disponibles": "skills_hr",
        "Available statistical skills": "skills_hr",
        "Accès aux données administratives": "access_admin_data",
        "Access to administrative data": "access_admin_data",
        "Financement disponible": "funding",
        "Available funding": "funding",
        "Financement": "funding",
        "Funding": "funding",
        "Outils numériques (collecte, traitement, diffusion)": "digital_tools",
        "Digital tools (collection, processing, dissemination)": "digital_tools",
        "Outils numériques": "digital_tools",
        "Digital tools": "digital_tools",
        "Cadre juridique pour le partage de données": "legal_framework",
        "Legal framework for data sharing": "legal_framework",
        "Cadre juridique": "legal_framework",
        "Legal framework": "legal_framework",
        "Coordination interinstitutionnelle": "institutional_coordination",
        "Inter-institutional coordination": "institutional_coordination",
        "Coordination institutionnelle": "institutional_coordination",
        "Institutional coordination": "institutional_coordination",
    }

    def _extract_table(table_obj: Any, mapping: Dict[str, str], prefix: str) -> Dict[str, Any]:
        out_tbl: Dict[str, Any] = {}
        # Ensure stable columns even when a respondent skips the section
        canons = sorted(set(mapping.values()))
        for canon in canons:
            out_tbl[f"{prefix}_{canon}"] = ""
            out_tbl[f"{prefix}_{canon}_spec"] = ""
        if not isinstance(table_obj, dict):
            return out_tbl
        for label, canon in mapping.items():
            cell = table_obj.get(label, None)
            if isinstance(cell, dict):
                out_tbl[f"{prefix}_{canon}"] = cell.get("code", "")
                out_tbl[f"{prefix}_{canon}_spec"] = cell.get("spec", "")
            elif isinstance(cell, str):
                out_tbl[f"{prefix}_{canon}"] = cell
        return out_tbl

    out: Dict[str, Any] = {}

    # Identification (Rubrique 2)
    out["organisation"] = payload.get("organisation", "")
    out["pays"] = payload.get("pays", "")
    out["pays_name_fr"] = payload.get("pays_name_fr", "")
    out["pays_name_en"] = payload.get("pays_name_en", "")
    out["type_acteur"] = payload.get("type_acteur", "")
    out["fonction"] = payload.get("fonction", "")
    out["email"] = payload.get("email", "")
    out["lang"] = payload.get("lang", "")

    # Rubrique 3 : portée
    out["scope"] = payload.get("scope", "")
    out["scope_other"] = payload.get("scope_other", "")

    # Rubrique 4 : domaines
    pre = payload.get("preselection_domains", [])
    out["preselection_domains"] = _join_list(pre)
    out["nb_preselection_domains"] = len(pre) if isinstance(pre, list) else 0

    top5 = payload.get("top5_domains", [])
    for i in range(5):
        out[f"top_domain_{i+1}"] = top5[i] if i < len(top5) else ""

    # Rubrique 5 : stats et notation
    selected_stats = payload.get("selected_stats", [])
    out["nb_stats"] = len(selected_stats) if isinstance(selected_stats, list) else 0
    out["stats_list"] = _join_list(selected_stats)
    out["selected_by_domain_json"] = _json(payload.get("selected_by_domain", {}))
    out["scoring_json"] = _json(payload.get("scoring", {}))
    out["scoring_version"] = payload.get("scoring_version", "")

    # Rubrique 6 : perspective de genre (table)
    out.update(_extract_table(payload.get("gender_table", {}), GENDER_ITEM_MAP, "gender"))

    # Rubrique 8 : capacité & faisabilité (table)
    out.update(_extract_table(payload.get("capacity_table", {}), CAPACITY_ITEM_MAP, "capacity"))

    # Rubrique 9 : harmonisation & qualité
    out["quality_expectations"] = _join_list(payload.get("quality_expectations", []))
    out["quality_other"] = payload.get("quality_other", "")

    # Rubrique 10 : diffusion
    out["dissemination_channels"] = _join_list(payload.get("dissemination_channels", []))
    out["dissemination_other"] = payload.get("dissemination_other", "")

    # Rubrique 12 : questions ouvertes
    out["comment_1"] = payload.get("open_q1", "")
    out["missing_indicators"] = payload.get("open_q2", "")
    out["support_needs"] = payload.get("open_q3", "")

    return out


def google_sheets_append(payload: Dict[str, Any]) -> Tuple[bool, str]:
    """
    Append a row into a Google Sheet if configured.
    Requires secrets:
      GOOGLE_SHEET_ID
      GOOGLE_SERVICE_ACCOUNT (dict)
    """
    if gspread is None or Credentials is None:
        return False, "Bibliothèques Google Sheets non disponibles (gspread/google-auth)."
    try:
        sheet_id = st.secrets.get("GOOGLE_SHEET_ID", None)
        sa_info = st.secrets.get("GOOGLE_SERVICE_ACCOUNT", None)
        if not sheet_id or not sa_info:
            return False, "Google Sheets non configuré (secrets manquants)."
        scopes = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
        creds = Credentials.from_service_account_info(sa_info, scopes=scopes)
        gc = gspread.authorize(creds)
        sh = gc.open_by_key(sheet_id)
        try:
            ws = sh.worksheet("responses")
        except Exception:
            ws = sh.add_worksheet(title="responses", rows=2000, cols=80)

        row = flatten_payload(payload)
        # Ensure header
        existing = ws.get_all_values()
        if not existing:
            ws.append_row(list(row.keys()), value_input_option="RAW")
        else:
            header = existing[0]
            # Add any missing columns at end
            for k in row.keys():
                if k not in header:
                    header.append(k)
            # If header grew, update first row
            ws.update("A1", [header])
        # Align order with header
        header = ws.row_values(1)
        values = [row.get(h, "") for h in header]
        ws.append_row(values, value_input_option="RAW")
        return True, "OK"
    except Exception as e:
        return False, f"Erreur Google Sheets : {e}"


def google_sheets_write_df(df: pd.DataFrame, worksheet_title: str, sheet_id: Optional[str] = None) -> Tuple[bool, str]:
    """
    Write a dataframe to a worksheet (overwrite), if configured.
    Uses secrets:
      GOOGLE_SHEET_ID (or provided sheet_id)
      GOOGLE_SERVICE_ACCOUNT (dict)
    """
    if gspread is None or Credentials is None:
        return False, "Bibliothèques Google Sheets non disponibles (gspread/google-auth)."
    try:
        sid = sheet_id or st.secrets.get("GOOGLE_SHEET_ID", None)
        sa_info = st.secrets.get("GOOGLE_SERVICE_ACCOUNT", None)
        if not sid or not sa_info:
            return False, "Google Sheets non configuré (secrets manquants)."

        scopes = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
        creds = Credentials.from_service_account_info(sa_info, scopes=scopes)
        gc = gspread.authorize(creds)
        sh = gc.open_by_key(sid)

        try:
            ws = sh.worksheet(worksheet_title)
        except Exception:
            ws = sh.add_worksheet(title=worksheet_title, rows=max(200, len(df) + 20), cols=max(20, len(df.columns) + 10))

        # Prepare values (header + rows), keep everything as string to avoid type surprises
        values = [list(df.columns)]
        if not df.empty:
            values += df.astype(object).where(pd.notnull(df), "").astype(str).values.tolist()

        ws.clear()
        ws.update("A1", values, value_input_option="RAW")
        return True, "OK"
    except Exception as e:
        return False, f"Erreur Google Sheets : {e}"


def dropbox_upload_bytes(content: bytes, filename: str, subfolder: str = "exports") -> Tuple[bool, str]:
    """
    Upload arbitrary bytes to Dropbox, if configured.
    Requires secret: DROPBOX_ACCESS_TOKEN
    Optional: DROPBOX_FOLDER (default /consultation_stat_niang)
    """
    if dropbox is None:
        return False, "Bibliothèque Dropbox non disponible."
    try:
        token = st.secrets.get("DROPBOX_ACCESS_TOKEN", None)
        if not token:
            return False, "Dropbox non configuré (DROPBOX_ACCESS_TOKEN manquant)."
        folder = st.secrets.get("DROPBOX_FOLDER", "/consultation_stat_niang")
        folder = folder if folder.startswith("/") else "/" + folder
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        path = f"{folder}/{subfolder}/{ts}_{filename}"
        dbx = dropbox.Dropbox(token)
        dbx.files_upload(content, path, mode=dropbox.files.WriteMode.overwrite)
        return True, "OK"
    except Exception as e:
        return False, f"Erreur Dropbox : {e}"

def dropbox_upload_json(submission_id: str, payload: Dict[str, Any]) -> Tuple[bool, str]:
    """
    Upload the JSON submission to Dropbox if configured.
    Requires secret: DROPBOX_ACCESS_TOKEN
    Optional: DROPBOX_FOLDER (default /consultation_stat_niang)
    """
    if dropbox is None:
        return False, "Bibliothèque Dropbox non disponible."
    try:
        token = st.secrets.get("DROPBOX_ACCESS_TOKEN", None)
        if not token:
            return False, "Dropbox non configuré (DROPBOX_ACCESS_TOKEN manquant)."
        folder = st.secrets.get("DROPBOX_FOLDER", "/consultation_stat_niang")
        folder = folder if folder.startswith("/") else "/" + folder
        path = f"{folder}/submissions/{submission_id}.json"
        dbx = dropbox.Dropbox(token)
        content = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
        dbx.files_upload(content, path, mode=dropbox.files.WriteMode.overwrite)
        return True, "OK"
    except Exception as e:
        return False, f"Erreur Dropbox : {e}"


# =========================
# Validation logic (quality controls)
# =========================

EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


def validate_r2(lang: str) -> List[str]:
    errs: List[str] = []
    organisation = resp_get("organisation", "").strip()
    pays = resp_get("pays", "").strip()
    type_acteur = resp_get("type_acteur", "").strip()
    fonction = resp_get("fonction", "").strip()
    fonction_autre = resp_get("fonction_autre", "").strip()
    email = resp_get("email", "").strip()

    if not organisation:
        errs.append(t(lang, "Organisation : champ obligatoire.", "Organization: required field."))
    # Contrôle qualité : éviter les sigles seuls
    elif len(organisation) < 12:
        errs.append(t(lang, "Organisation : indiquez le libellé complet (au moins 12 caractères) et non le sigle.", "Organization: please provide the full name (at least 12 characters), not only an acronym."))
    if not pays:
        errs.append(t(lang, "Pays de résidence : champ obligatoire.", "Country of Residence: required field."))
    if not type_acteur:
        errs.append(t(lang, "Type d’acteur : champ obligatoire.", "Stakeholder type: required field."))
    if not fonction:
        errs.append(t(lang, "Fonction : champ obligatoire.", "Role/Function: required field."))
    if fonction == "Autre" or fonction == "Other":
        if not fonction_autre:
            errs.append(t(lang, "Fonction (Autre) : précisez.", "Role (Other): please specify."))
    if not email:
        errs.append(t(lang, "Email : champ obligatoire.", "Email: required field."))
    elif not EMAIL_RE.match(email):
        errs.append(t(lang, "Email : format invalide.", "Email: invalid format."))
    return errs



def validate_r3(lang: str) -> List[str]:
    errs: List[str] = []
    scope = (resp_get("scope", "") or "").strip()
    if not scope:
        errs.append(t(lang, "Rubrique 3 : veuillez sélectionner une portée.", "Section 3: please select a scope."))
        return errs
    snds = (resp_get("snds_status", "") or "").strip()
    if not snds:
        errs.append(
            t(
                lang,
                "Rubrique 3 : veuillez indiquer le statut de la SNDS / plan national statistique.",
                "Section 3: please indicate the status of the NSDS / national statistical plan.",
            )
        )

    if scope == "Other":
        other = (resp_get("scope_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 3 : précisez l’option « Autre ».", "Section 3: please specify the \"Other\" option."))
    return errs

def validate_r4(lang: str) -> List[str]:
    errs: List[str] = []
    pre = resp_get("preselected_domains", [])
    top5 = resp_get("top5_domains", [])
    if not isinstance(pre, list):
        pre = []
    if not isinstance(top5, list):
        top5 = []
    if len(pre) < 5 or len(pre) > 10:
        errs.append(t(lang, "Rubrique 4 : pré-sélectionnez entre 5 et 10 domaines.", "Section 4: pre-select 5 to 10 domains."))
    if len(set(pre)) != len(pre):
        errs.append(t(lang, "Rubrique 4 : la pré-sélection contient des doublons.", "Section 4: duplicates found in pre-selection."))
    if len(top5) != 5:
        errs.append(t(lang, "Rubrique 4 : le TOP 5 doit contenir exactement 5 domaines.", "Section 4: TOP 5 must contain exactly 5 domains."))
    else:
        if len(set(top5)) != 5:
            errs.append(t(lang, "Rubrique 4 : le TOP 5 contient des doublons.", "Section 4: TOP 5 contains duplicates."))
        missing = [d for d in top5 if d not in pre]
        if missing:
            errs.append(t(lang, "Rubrique 4 : chaque domaine du TOP 5 doit provenir de la pré-sélection.", "Section 4: TOP 5 must be selected from pre-selection."))
    return errs


def validate_r5(lang: str) -> List[str]:
    errs: List[str] = []
    top5 = resp_get("top5_domains", [])
    selected_by_domain: Dict[str, List[str]] = resp_get("selected_by_domain", {})
    if not isinstance(selected_by_domain, dict):
        selected_by_domain = {}

    all_stats: List[str] = []
    for d in top5:
        stats = selected_by_domain.get(d, [])
        if not isinstance(stats, list):
            stats = []
        if len(stats) < 1:
            errs.append(t(lang, f"Rubrique 5 : choisissez au moins 1 statistique pour {d}.",
                          f"Section 5: select at least 1 indicator for {d}."))
        if len(stats) > 3:
            errs.append(t(lang, f"Rubrique 5 : maximum 3 statistiques pour {d}.",
                          f"Section 5: maximum 3 indicators for {d}."))
        all_stats.extend(stats)

    if len(all_stats) < 5 or len(all_stats) > 15:
        errs.append(t(lang, "Rubrique 5 : le total des statistiques doit être entre 5 et 15.",
                      "Section 5: total number of indicators must be between 5 and 15."))

    if len(set(all_stats)) != len(all_stats):
        errs.append(t(lang, "Rubrique 5 : une même statistique ne doit pas être sélectionnée plusieurs fois.",
                      "Section 5: the same indicator must not be selected more than once."))

    # scoring
    scoring: Dict[str, Dict[str, Any]] = resp_get("scoring", {})
    if not isinstance(scoring, dict):
        scoring = {}

    for s in all_stats:
        if s not in scoring:
            errs.append(t(lang, f"Rubrique 5 : vous devez noter la statistique {s}.",
                          f"Section 5: you must score indicator {s}."))
            continue
        for k in ["demand", "availability", "feasibility"]:
            sc_row = scoring.get(s, {}) or {}
            k_lbl = {
                "demand": t(lang, "demande", "demand"),
                "availability": t(lang, "disponibilité", "availability"),
                "feasibility": t(lang, "faisabilité", "feasibility"),
            }.get(k, k)

            # Backward compatibility: legacy key "gap" -> "availability"
            if k == "availability":
                v_raw = sc_row.get("availability", sc_row.get("gap", None))
            else:
                v_raw = sc_row.get(k, None)

            if v_raw is None or str(v_raw).strip() == "":
                errs.append(t(lang, f"Rubrique 5 : la note '{k_lbl}' manque pour {s}.",
                              f"Section 5: missing score '{k_lbl}' for {s}."))
            else:
                try:
                    v = int(v_raw)
                    if v < 0 or v > 3:
                        errs.append(t(lang, f"Rubrique 5 : note invalide pour {s} ({k_lbl}).",
                                      f"Section 5: invalid score for {s} ({k_lbl})."))
                except Exception:
                    errs.append(t(lang, f"Rubrique 5 : note invalide pour {s} ({k_lbl}).",
                                  f"Section 5: invalid score for {s} ({k_lbl})."))
    return errs


def validate_r6(lang: str) -> List[str]:
    errs: List[str] = []
    tbl = resp_get("gender_table", {})
    if not isinstance(tbl, dict) or not tbl:
        errs.append(t(lang, "Rubrique 6 : veuillez renseigner le tableau.", "Section 6: please complete the table."))
        return errs
    for k, v in tbl.items():
        if not v:
            errs.append(t(lang, f"Rubrique 6 : ligne non renseignée : {k}.", f"Section 6: missing answer for: {k}."))
    return errs


def validate_r8(lang: str) -> List[str]:
    errs: List[str] = []

    # Tableau capacité / faisabilité
    tbl = resp_get("capacity_table", {})
    if not isinstance(tbl, dict) or not tbl:
        errs.append(t(lang, "Rubrique 8 : veuillez renseigner le tableau.", "Section 8: please complete the table."))
        return errs
    for k, v in tbl.items():
        if not v:
            errs.append(t(lang, f"Rubrique 8 : ligne non renseignée : {k}.", f"Section 8: missing answer for: {k}."))

    return errs


def validate_r9(lang: str) -> List[str]:
    errs: List[str] = []
    sel = resp_get("quality_expectations", [])
    if not isinstance(sel, list) or len([x for x in sel if str(x).strip() != ""]) == 0:
        errs.append(t(lang, "Rubrique 9 : veuillez sélectionner au moins une option.", "Section 9: please select at least one option."))
        return errs
    if ("Autre" in sel) or ("Other" in sel):
        other = (resp_get("quality_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 9 : précisez l’option « Autre ».", "Section 9: please specify the \"Other\" option."))
    return errs


def validate_r10(lang: str) -> List[str]:
    errs: List[str] = []
    sel = resp_get("dissemination_channels", [])
    if not isinstance(sel, list) or len([x for x in sel if str(x).strip() != ""]) == 0:
        errs.append(t(lang, "Rubrique 10 : veuillez sélectionner au moins une option.", "Section 10: please select at least one option."))
        return errs
    if ("Autre" in sel) or ("Other" in sel):
        other = (resp_get("dissemination_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 10 : précisez l’option « Autre ».", "Section 10: please specify the \"Other\" option."))
    return errs


def validate_r11(lang: str) -> List[str]:
    errs: List[str] = []
    sel = resp_get("data_sources", [])
    if not isinstance(sel, list):
        sel = []

    sel_clean = [str(x).strip() for x in sel if str(x).strip()]
    if len(sel_clean) < 2:
        errs.append(t(lang, "Rubrique 11 : sélectionnez au moins 2 sources.", "Section 11: please select at least 2 sources."))
        return errs
    if len(sel_clean) > 4:
        errs.append(t(lang, "Rubrique 11 : sélectionnez au maximum 4 sources.", "Section 11: please select at most 4 sources."))

    if ("Autres" in sel_clean) or ("Other" in sel_clean):
        other = (resp_get("data_sources_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 11 : précisez l’option « Autres ».", "Section 11: please specify the 'Other' option."))

    return errs


def validate_r12(lang: str) -> List[str]:
    errs: List[str] = []
    sub = int(st.session_state.get("r12_substep", 0) or 0)
    if sub < 3:
        errs.append(
            t(
                lang,
                "Rubrique 12 : veuillez traiter les questions ouvertes une à une (bouton « Question suivante ») jusqu’à la Confirmation.",
                "Section 12: please go through the open questions one by one (use the “Next question” button) until Confirmation.",
            )
        )

    cc = (resp_get("consulted_colleagues", "") or "").strip()
    if cc not in ("YES", "NO"):
        errs.append(
            t(
                lang,
                "Rubrique 12 : veuillez indiquer si vous avez consulté d’autres collègues (Oui/Non).",
                "Section 12: please indicate whether you consulted other colleagues (Yes/No).",
            )
        )
    return errs

def validate_all(lang: str) -> List[str]:
    errs = []
    errs.extend(validate_r2(lang))
    errs.extend(validate_r3(lang))
    errs.extend(validate_r4(lang))
    errs.extend(validate_r5(lang))
    errs.extend(validate_r6(lang))
    errs.extend(validate_r8(lang))
    errs.extend(validate_r9(lang))
    errs.extend(validate_r10(lang))
    errs.extend(validate_r11(lang))
    errs.extend(validate_r12(lang))
    # Open questions (text fields) remain optional; warnings are shown in Section 12 / Submit.
    return errs


# =========================
# Navigation
# =========================

def get_steps(lang: str) -> List[Tuple[str, str]]:
    # Rubric 7 added, plus final SEND tab
    return [
        ("R1", t(lang, "Rubrique 1 : Instructions", "Section 1: Instructions")),
        ("R2", t(lang, "Rubrique 2 : Identification du répondant", "Section 2: Respondent identification")),
        ("R3", t(lang, "Rubrique 3 : Portée de la réponse", "Section 3: Scope of response")),
        ("R4", t(lang, "Rubrique 4 : Domaines prioritaires", "Section 4: Priority domains")),
        ("R5", t(lang, "Rubrique 5 : Statistiques prioritaires et notation", "Section 5: Priority indicators and scoring")),
        ("R6", t(lang, "Rubrique 6 : Dimension genre", "Section 6: Gender dimension")),
        ("R7", t(lang, "Rubrique 7 : Priorités genre", "Section 7: Gender priorities")),
        ("R8", t(lang, "Rubrique 8 : Capacité et faisabilité (12–24 mois)", "Section 8: Capacity and feasibility (12–24 months)")),
        ("R9", t(lang, "Rubrique 9 : Harmonisation et qualité", "Section 9: Harmonization and quality")),
        ("R10", t(lang, "Rubrique 10 : Diffusion", "Section 10: Dissemination")),
        ("R11", t(lang, "Rubrique 11 : Sources de données pertinentes", "Section 11: Relevant data sources")),
        ("R12", t(lang, "Rubrique 12 : Questions ouvertes", "Section 12: Open questions")),
        ("SEND", t(lang, "ENVOYER", "SUBMIT")),
    ]


def render_sidebar(lang: str, steps: List[Tuple[str, str]]) -> None:
    st.sidebar.header(t(lang, "Navigation", "Navigation"))
    labels = [s[1] for s in steps]

    # Keep sidebar selection in sync with nav_idx
    st.session_state.nav_radio = int(st.session_state.nav_idx)
    chosen = st.sidebar.radio(
        t(lang, "Aller à", "Go to"),
        options=list(range(len(labels))),
        index=int(st.session_state.nav_idx),
        format_func=lambda i: labels[i],
        key="nav_radio"
    )

    # User clicked in sidebar
    if int(chosen) != int(st.session_state.nav_idx):
        st.session_state.nav_idx = int(chosen)

    st.sidebar.divider()
    st.sidebar.caption(
        t(
            lang,
            "Note : les contrôles qualité peuvent bloquer la progression si une contrainte n’est pas respectée.",
            "Note: quality checks may prevent moving forward when constraints are not met."
        )
    )

    st.sidebar.markdown("---")
    st.sidebar.caption(
        t(
            lang,
            "NSP : Ne sait pas (score 0). Utilisez NSP uniquement si l’information est indisponible.",
            "UK: Unknown (score 0). Use UK only when information is unavailable."
        )
    )
    st.sidebar.markdown("---")
    st.sidebar.subheader(t(lang, "Brouillon", "Draft"))
    if st.sidebar.button(t(lang, "Sauvegarder maintenant", "Save now")):
        ok, msg = autosave_draft(force=True)
        if ok:
            st.sidebar.success(t(lang, "Brouillon sauvegardé.", "Draft saved."))
        else:
            st.sidebar.error(t(lang, "Brouillon non sauvegardé.", "Draft not saved."))
    if st.session_state.get("draft_id"):
        st.sidebar.caption(
            t(
                lang,
                "Reprise : conservez l’URL de cette page (paramètre rid=...).",
                "Resume: keep this page URL (rid=... parameter)."
            )
        )



def nav_buttons(lang: str, steps: List[Tuple[str, str]], df_long: pd.DataFrame) -> None:
    """Bottom nav buttons, with blocking based on current step validations."""
    step_key = steps[st.session_state.nav_idx][0]
    errors: List[str] = []

    # Blocking rules per step
    if step_key == "R2":
        errors = validate_r2(lang)
    elif step_key == "R3":
        errors = validate_r3(lang)
    elif step_key == "R4":
        errors = validate_r4(lang)
    elif step_key == "R5":
        errors = validate_r5(lang)
    elif step_key == "R6":
        errors = validate_r6(lang)
    elif step_key == "R7":
        errors = validate_r7(lang)
    elif step_key == "R8":
        errors = validate_r8(lang)
    elif step_key == "R9":
        errors = validate_r9(lang)
    elif step_key == "R10":
        errors = validate_r10(lang)
    elif step_key == "R11":
        errors = validate_r11(lang)
    elif step_key == "R12":
        errors = validate_r12(lang)

    col1, col2, col3 = st.columns([1, 1, 3])
    with col1:
        prev_disabled = st.session_state.nav_idx <= 0
        if st.button(t(lang, "⬅ Précédent", "⬅ Previous"), disabled=prev_disabled):
            autosave_draft(force=True)
            st.session_state.nav_idx = max(0, st.session_state.nav_idx - 1)
            st.rerun()
    with col2:
        next_disabled = (st.session_state.nav_idx >= len(steps) - 1) or bool(errors)
        if st.button(t(lang, "Suivant ➡", "Next ➡"), disabled=next_disabled):
            autosave_draft(force=True)
            st.session_state.nav_idx = min(len(steps) - 1, st.session_state.nav_idx + 1)
            st.rerun()
    with col3:
        if errors:
            st.error("\n".join(errors))


# =========================
# UI : Rubrics
# =========================

def rubric_1(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 1 : Instructions", "Section 1: Instructions"))
    st.markdown(
        t(
            lang,
            """
### Objectif
Ce questionnaire vise à recueillir votre avis sur **les statistiques socio-économiques prioritaires** à produire et diffuser au niveau continental.

### Comment répondre
1. **Identifiez** votre organisation (Rubrique 2).
2. **Pré-sélectionnez 5 à 10 domaines** et classez un **TOP 5** (Rubrique 4).
3. Pour chaque domaine du TOP 5 : choisissez **1 à 3 statistiques** et attribuez des **notes** (Rubrique 5).
4. Complétez les rubriques transversales : **genre**, **capacité/faisabilité**, **etc.**.
5. **N'hésitez pas à consulter les infobulles ⍰ pour plus de précisions.**

            """,
            """
### Purpose
This questionnaire collects your views on **priority socio-economic statistics** to be produced and disseminated at continental level.

### How to answer
1. **Identify** your organization (Section 2).
2. **Pre-select 5–10 domains** and rank a **TOP 5** (Section 4).
3. For each TOP 5 domain: select **1–3 indicators** and provide **scores** (Section 5).
4. Complete cross-cutting sections: **gender**, **capacity/feasibility**,  **etc.**.
5. **Feel free to consult the ⍰ tooltips for more details.**

            """
        )
    )


def rubric_2(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 2 : Identification du répondant", "Section 2: Respondent identification"))
    st.info(
        t(
            lang,
            "Merci de renseigner ces informations. Elles servent uniquement à l’analyse et ne seront pas publiées nominativement.",
            "Please provide these details. They are used for analysis and will not be published in a personally identifiable way."
        )
    )

    resp_set("lang", lang)

    st.text_input(t(lang, "Nom de l'organisation", "Organization Name"), key="org_input", value=resp_get("organisation", ""))
    resp_set("organisation", st.session_state.get("org_input", "").strip())
    st.caption(t(lang, "Merci d’indiquer le libellé complet (évitez le sigle seul).", "Please provide the full organization name (avoid acronym only)."))
    col1, col2 = st.columns(2)
    with col1:
        # Pays : liste déroulante (ISO3 + libellés FR/EN)
        df_countries = load_countries()
        iso3_list, iso3_to_fr, iso3_to_en = country_maps(df_countries)

        prev_country = (resp_get("pays", "") or "").strip()
        prev_iso3 = (prev_country.split("|", 1)[0].strip().upper() if "|" in prev_country else prev_country.strip().upper())
        if prev_iso3 not in iso3_list:
            prev_iso3 = ""

        if not iso3_list:
            # Fallback si le fichier pays est introuvable
            st.text_input(t(lang, "Pays", "Country"), key="country_input", value=resp_get("pays", ""))
            resp_set("pays", st.session_state.get("country_input", "").strip())
            resp_set("pays_name_fr", "")
            resp_set("pays_name_en", "")
        else:
            options = [""] + sorted(iso3_list, key=lambda x: country_label(x, lang, iso3_to_fr, iso3_to_en).lower())

            chosen_iso3 = st.selectbox(
                t(lang, "Pays", "Country"),
                options=options,
                index=options.index(prev_iso3) if prev_iso3 in options else 0,
                format_func=lambda x: (
                    t(lang, "— Sélectionner —", "— Select —") if x == ""
                    else f"{country_label(x, lang, iso3_to_fr, iso3_to_en)} ({x})"
                ),
                help=t(lang, "Choisissez votre pays (liste ISO3).", "Select your country (ISO3 list)."),
                key="country_iso3_select",
            )
            resp_set("pays", chosen_iso3)
            # Libellés normalisés (utile pour les exports / analyses)
            if chosen_iso3:
                resp_set("pays_name_fr", iso3_to_fr.get(chosen_iso3, "").strip())
                resp_set("pays_name_en", (iso3_to_en.get(chosen_iso3, "") or iso3_to_fr.get(chosen_iso3, "")).strip())
            else:
                resp_set("pays_name_fr", "")
                resp_set("pays_name_en", "")

    with col2:
        st.text_input(
            t(lang, "Email", "Email"),
            key="email_input",
            value=resp_get("email", ""),
            help=t(
                lang,
                "Saisissez une adresse email valide (ex. nom@domaine.tld).",
                "Enter a valid email address (e.g., name@domain.tld).",
            ),
        )
        resp_set("email", st.session_state.get("email_input", "").strip())

    # Brouillon : crée un identifiant de reprise dès que l’email est renseigné
    ensure_draft_id()
    autosave_draft(force=False)

    # Afficher le message de reprise dès la première session après saisie de l’email
    _email_now = resp_get("email", "")
    if _email_now and ("@" in _email_now) and not st.session_state.get("draft_resume_notice_shown"):
        st.warning(
            t(
                lang,
                "La saisie est sauvegardée. En cas de suspension de moins de 48 heures, reprenez là où vous vous étiez arrêté en ré-ouvrant le lien contenant rid (à conserver / mettre en favori / retrouver dans l'historique).",
                "Your input is saved. If you pause for less than 48 hours, resume where you left off by reopening the link containing rid (keep it / bookmark it / find it in your browser history).",
            )
        )
        st.session_state["draft_resume_notice_shown"] = True
    type_options = [
        ("NSO", {"fr": "Institut national de statistique", "en": "National Statistical Office"}),
        ("Ministry", {"fr": "Ministère / Service statistique sectoriel", "en": "Ministry / Sector statistical unit"}),
        ("REC", {"fr": "Communauté économique régionale", "en": "Regional Economic Community"}),
        ("AU", {"fr": "Union Africaine (UA)", "en": "African Union (AU)"}),
        ("CivilSoc", {"fr": "Société civile", "en": "Civil society"}),
        ("DevPartner", {"fr": "Partenaire technique et financier", "en": "Development partner"}),
        ("Academia", {"fr": "Université / Recherche", "en": "Academia / Research"}),
        ("Other", {"fr": "Autre", "en": "Other"}),
    ]
    type_labels = [t(lang, x[1]["fr"], x[1]["en"]) for x in type_options]
    type_keys = [x[0] for x in type_options]

    # Type d’acteur : pas de pré-remplissage (placeholder)
    type_opts = [""] + type_keys
    prev_type = resp_get("type_acteur", "")
    idx = type_opts.index(prev_type) if prev_type in type_opts else 0

    chosen_type = st.selectbox(
        t(lang, "Type d’acteur", "Stakeholder type"),
        options=type_opts,
        index=idx,
        format_func=lambda k: (t(lang, "— Sélectionner —", "— Select —") if k == "" else type_labels[type_keys.index(k)]),
        help=t(lang, "Choisissez la catégorie correspondant le mieux à votre organisation.", 
               "Choose the category that best matches your organization.")
    )
    resp_set("type_acteur", chosen_type)
# Fonction dropdown : pas de pré-remplissage (placeholder)
    role_opts = ROLE_OPTIONS_FR if lang == "fr" else ROLE_OPTIONS_EN
    role_options = [""] + role_opts
    prev_role = resp_get("fonction", "")
    role_idx = role_options.index(prev_role) if prev_role in role_options else 0

    chosen_role = st.selectbox(
        t(lang, "Fonction", "Role/Function"),
        options=role_options,
        index=role_idx,
        format_func=lambda x: (t(lang, "— Sélectionner —", "— Select —") if x == "" else x),
        help=t(lang, "Indiquez votre fonction principale dans l’organisation.", "Indicate your main role in the organization."),
    )
    resp_set("fonction", chosen_role)
    if chosen_role in ["Autre", "Other"]:
        st.text_input(t(lang, "Préciser (fonction)", "Specify (role)"),
                      key="fonction_autre_input", value=resp_get("fonction_autre", ""))
        resp_set("fonction_autre", st.session_state.get("fonction_autre_input", "").strip())
    else:
        resp_set("fonction_autre", "")

    # Live errors
    errs = validate_r2(lang)
    if errs:
        st.warning(t(lang, "Veuillez corriger les éléments ci-dessous :", "Please fix the following:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_3(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 3 : Portée de la réponse", "Section 3: Scope of response"))
    st.markdown(
        t(
            lang,
            "Indiquez le périmètre principal de votre réponse. Cela aide à interpréter vos priorités.",
            "Indicate the main scope of your response. This helps interpret your priorities."
        )
    )

    scope_opts_raw = [
        ("National", {"fr": "National", "en": "National"}),
        ("Regional", {"fr": "Régional (CER)", "en": "Regional (REC)"}),
        ("Continental", {"fr": "Continental (UA)", "en": "Continental (AU)"}),
        ("Global", {"fr": "International", "en": "International"}),
        ("Other", {"fr": "Autre", "en": "Other"}),
    ]
    scope_labels = {k: t(lang, v["fr"], v["en"]) for k, v in scope_opts_raw}
    scope_keys = [k for k, _ in scope_opts_raw]
    scope_options = [""] + scope_keys

    prev_scope = resp_get("scope", "")
    scope_idx = scope_options.index(prev_scope) if prev_scope in scope_options else 0

    chosen_scope = st.selectbox(
        t(lang, "Portée", "Scope"),
        options=scope_options,
        index=scope_idx,
        format_func=lambda k: (t(lang, "— Sélectionner —", "— Select —") if k == "" else scope_labels.get(k, k)),
        help=t(
            lang,
            "Indiquez le périmètre principal de votre réponse : national, régional (CER), continental (UA) ou international.",
            "Indicate the main scope of your response: national, regional (REC), continental (AU), or international."
        )
    )
    resp_set("scope", chosen_scope)

    # SNDS / Plan statistique national (obligatoire)
    snds_opts = ["", "YES", "NO", "PREP", "IMPL_PREP", "NSP"]
    snds_labels = {
        "YES": t(lang, "Oui", "Yes"),
        "NO": t(lang, "Non", "No"),
        "PREP": t(lang, "En préparation", "In preparation"),
        "IMPL_PREP": t(lang, "En cours de mise en œuvre ET nouvelle en préparation", "Under implementation AND new one in preparation"),
        "NSP": t(lang, "NSP", "DK"),
    }
    prev_snds = (resp_get("snds_status", "") or "").strip()
    idx_snds = snds_opts.index(prev_snds) if prev_snds in snds_opts else 0
    chosen_snds = st.selectbox(
        t(
            lang,
            "Statut de la SNDS / plan national statistique en cours",
            "Status of the current NSDS / national statistical plan",
        ),
        options=snds_opts,
        index=idx_snds,
        format_func=lambda k: (
            t(lang, "— Sélectionner —", "— Select —") if k == "" else snds_labels.get(k, k)
        ),
        help=t(
            lang,
            "Indiquez si une stratégie / plan statistique national est en cours, non, en préparation, ou NSP.",
            "Indicate whether an NSDS / national statistical plan is current, not in place, under preparation, or DK.",
        ),
        key="snds_status_select",
    )
    resp_set("snds_status", chosen_snds)


    # Contrôle qualité (alerte) : cohérence acteur × portée
    _actor = (resp_get("type_acteur", "") or "").strip()
    _scope = (resp_get("scope", "") or "").strip()
    if _actor in ["NSO", "Ministry"] and _scope and _scope != "National":
        st.warning(
            t(
                lang,
                "Alerte : vous avez indiqué « Institut national de statistique » ou « Ministère », mais la portée n’est pas « National ». Merci de vérifier la cohérence.",
                "Warning: you selected “National Statistical Office” or “Ministry”, but the scope is not “National”. Please check consistency."
            )
        )

    if resp_get("scope") == "Other":
        st.text_input(t(lang, "Préciser", "Specify"), key="scope_other_input", value=resp_get("scope_other", ""))
        resp_set("scope_other", st.session_state.get("scope_other_input", "").strip())
    else:
        resp_set("scope_other", "")



def rubric_4(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "Rubrique 4 : Domaines prioritaires", "Section 4: Priority domains"))

    st.info(
        t(
            lang,
            "Veuillez d’abord choisir 5 à 10 domaines (pré-sélection). Ensuite, choisissez exactement 5 domaines dans ce sous-ensemble (TOP 5).\n\nConseil : choisissez les domaines où la demande politique est forte.",
            "First select 5 to 10 domains (pre-selection). Then choose exactly 5 domains within that subset (TOP 5).\n\nTip: choose domains where policy demand is strong."
        )
    )

    domains = domains_from_longlist(df_long, lang)
    if not domains:
        st.error(
            t(
                lang,
                "La liste des domaines n’est pas disponible (longlist introuvable ou vide).",
                "Domain list is not available (longlist not found or empty).",
            )
        )
        st.caption(
            t(
                lang,
                "Vérifiez que le dépôt contient : data/indicator_longlist.csv (prioritaire) ou data/longlist.xlsx (ou ces fichiers à la racine).",
                "Check that the repository contains: data/indicator_longlist.csv (preferred) or data/longlist.xlsx (or these files at repo root).",
            )
        )
        return

    code_to_label = {c: lbl for c, lbl in domains}

    # Build display labels without showing codes (codes are stored internally)
    labels = [code_to_label[c] for c, _ in domains]
    # Disambiguate duplicates if any (rare)
    seen = {}
    for i, (c, _) in enumerate(domains):
        lbl = code_to_label[c]
        seen[lbl] = seen.get(lbl, 0) + 1
    display_labels = []
    label_to_code = {}
    for c, _ in domains:
        lbl = code_to_label[c]
        disp = lbl if seen[lbl] == 1 else f"{lbl} ({c})"
        display_labels.append(disp)
        label_to_code[disp] = c

    st.markdown(
        t(
            lang,
            """
### Étape 1 : Pré-sélection
Sélectionnez **entre 5 et 10 domaines** (sans doublons).
            """,
            """
### Step 1: Pre-selection
Select **5 to 10 domains** (no duplicates).
            """
        )
    )

    pre_default_codes = resp_get("preselected_domains", [])
    pre_default_disp = []
    for c in pre_default_codes:
        lbl = code_to_label.get(c, "")
        if not lbl:
            continue
        disp = lbl if seen.get(lbl, 1) == 1 else f"{lbl} ({c})"
        if disp in label_to_code:
            pre_default_disp.append(disp)

    # Avoid "first click not kept" by initializing widget state once (no default on every rerun)
    if "r4_preselection_ms" not in st.session_state:
        st.session_state["r4_preselection_ms"] = pre_default_disp

    pre_disp = st.multiselect(
        t(lang, "Pré-sélection (5–10 domaines)", "Pre-selection (5–10 domains)"),
        options=display_labels,
        max_selections=10,
        key="r4_preselection_ms",
        help=t(lang, "Choisissez au maximum 10 domaines. Une fois 10 domaines sélectionnés, les nouveaux clics seront ignorés.", 
               "Select up to 10 domains. Once 10 domains are selected, additional clicks are ignored.")
    )

    pre_codes = [label_to_code[x] for x in pre_disp]
    resp_set("preselected_domains", pre_codes)

    st.divider()
    st.markdown(
        t(
            lang,
            """
### Étape 2 : Classement TOP 5
Classez exactement **5 domaines** parmi votre pré-sélection.
            """,
            """
### Step 2: Rank TOP 5
Rank exactly **5 domains** from your pre-selection.
            """
        )
    )

    if len(pre_codes) < 5:
        st.warning(t(lang, "Sélectionnez d’abord au moins 5 domaines dans la pré-sélection.",
                     "Please pre-select at least 5 domains first."))
        resp_set("top5_domains", [])
        return

    top5: List[str] = []

    # Ranking with 5 selectboxes (no prefill + no duplicates)
    chosen_prev: List[str] = []
    for i in range(5):
        key = f"top5_rank_{i+1}"

        # Options for this rank = preselection minus already chosen
        remaining = [c for c in pre_codes if c not in chosen_prev]
        options = [""] + remaining  # "" placeholder (no prefill)

        prev = resp_get(key, "")
        if prev and prev in remaining:
            idx = options.index(prev)
        else:
            idx = 0

        choice = st.selectbox(
            t(lang, f"Rang {i+1}", f"Rank {i+1}"),
            options=options,
            index=idx,
            format_func=lambda c: (t(lang, "— Sélectionner —", "— Select —") if c == "" else code_to_label.get(c, c)),
            help=t(
                lang,
                "Choisissez un domaine unique pour chaque rang. Les domaines déjà choisis ne sont plus proposés aux rangs suivants.",
                "Choose a unique domain for each rank. Already selected domains are removed from the next ranks.",
            ),
            key=key,
        )

        if choice != "":
            top5.append(choice)
            chosen_prev.append(choice)

    resp_set("top5_domains", top5)


    errs = validate_r4(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))



def rubric_5(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "Rubrique 5 : Statistiques prioritaires et notation", "Section 5: Priority indicators and scoring"))

    top5 = resp_get("top5_domains", [])
    if not top5 or len(top5) != 5:
        st.warning(t(lang, "Veuillez d’abord finaliser le TOP 5 des domaines (Rubrique 4).",
                     "Please complete TOP 5 domains first (Section 4)."))
        return

    # mapping for domain display
    dom_map = {c: lbl for c, lbl in domains_from_longlist(df_long, lang)}

    st.markdown(
        t(
            lang,
            """
### Étape A : Sélection des statistiques
Pour chaque domaine du TOP 5 : choisissez **1 à 3 statistiques**.
- Total attendu : **entre 5 et 15** statistiques.
- Une statistique ne doit pas apparaître dans deux domaines.

### Étape B : Notation multicritères (scoring rationalisé)
Pour chaque statistique sélectionnée, attribuez une note (0–3) sur :
- **Demande politique** : importance pour le pilotage / les priorités
- **Disponibilité actuelle** : niveau de production existant (faible=1, partielle=2, bonne=3)
- **Faisabilité (12–24 mois)** : capacité réaliste à produire ou améliorer d’ici 12–24 mois
            """,
            """
### Step A: Select indicators
For each TOP 5 domain: select **1 to 3 indicators**.
- Expected total: **5 to 15** indicators.
- The same indicator must not be selected under two domains.

### Step B: Multi-criteria scoring (streamlined)
For each selected indicator, provide a score (0–3) for:
- **Political demand**
- **Current availability** (low=1, partial=2, good=3)
- **Feasibility (12–24 months)**
            """
        )
    )

    selected_by_domain: Dict[str, List[str]] = resp_get("selected_by_domain", {})
    if not isinstance(selected_by_domain, dict):
        selected_by_domain = {}

    scoring: Dict[str, Dict[str, Any]] = resp_get("scoring", {})
    if not isinstance(scoring, dict):
        scoring = {}

    # Ensure dict keys exist
    for d in top5:
        if d not in selected_by_domain:
            selected_by_domain[d] = []

    # UI selection per domain (codes hidden)
    for d in top5:
        st.markdown(f"#### {dom_map.get(d, d)}")

        stats_opts = stats_for_domain(df_long, d, lang)
        stat_code_to_label = {c: lbl for c, lbl in stats_opts}

        # build display labels without showing stat codes
        labels = [stat_code_to_label[c] for c, _ in stats_opts]
        seen = {}
        for c, _ in stats_opts:
            lbl = stat_code_to_label[c]
            seen[lbl] = seen.get(lbl, 0) + 1
        display_labels = []
        label_to_code = {}
        for c, _ in stats_opts:
            lbl = stat_code_to_label[c]
            disp = lbl if seen[lbl] == 1 else f"{lbl} ({c})"
            display_labels.append(disp)
            label_to_code[disp] = c

        default_codes = selected_by_domain.get(d, [])
        default_disp = []
        for c in default_codes:
            lbl = stat_code_to_label.get(c, "")
            if not lbl:
                continue
            disp = lbl if seen.get(lbl, 1) == 1 else f"{lbl} ({c})"
            if disp in label_to_code:
                default_disp.append(disp)

        key_ms = f"stats_ms_{d}"

        # Init widget state once (avoid "first click" issues)
        if key_ms not in st.session_state:
            st.session_state[key_ms] = default_disp

        picked_disp = st.multiselect(
            t(lang, "Choisir 1 à 3 statistiques", "Select 1 to 3 indicators"),
            options=display_labels,
            max_selections=3,
            key=key_ms,
            help=t(lang, "Sélectionnez au minimum 1 et au maximum 3 statistiques pour ce domaine.",
                   "Select at least 1 and at most 3 indicators for this domain.")
        )

        picked_codes = [label_to_code[x] for x in picked_disp]
        selected_by_domain[d] = picked_codes


    # Uniqueness check
    flattened = []
    for d in top5:
        flattened.extend(selected_by_domain.get(d, []))
    duplicates = [x for x in set(flattened) if flattened.count(x) > 1]
    if duplicates:
        st.error(
            t(
                lang,
                "Une ou plusieurs statistiques sont sélectionnées dans plusieurs domaines. Veuillez corriger.",
                "One or more indicators are selected under multiple domains. Please correct."
            )
        )

    resp_set("selected_by_domain", selected_by_domain)
    resp_set("selected_stats", flattened)

    # Map codes to labels for display in scoring
    global_map = {}
    for d in top5:
        for c, lbl in stats_for_domain(df_long, d, lang):
            global_map[c] = lbl

    st.divider()
    st.markdown("### " + t(lang, "Notation multicritères", "Multi-criteria scoring"))

    for s in flattened:
        if s not in scoring or not isinstance(scoring.get(s), dict):
            scoring[s] = {}

        # Backward compatibility: legacy key "gap" (écart) -> "availability"
        # On normalise sur l'échelle v3 (Bonne=3)
        if "availability" not in scoring[s] and "gap" in scoring[s]:
            scoring[s]["availability"] = normalize_availability(scoring[s].get("gap", 0), 0)

        # Ensure keys exist
        for k in ["demand", "availability", "feasibility"]:
            if k not in scoring[s]:
                scoring[s][k] = None

        st.markdown(f"**{global_map.get(s, s)}**")

        c1, c2, c3 = st.columns(3)
        opts = [None, 1, 2, 3, 0]  # None = placeholder (no prefill). 0 = NSP / DK

        with c1:
            prev = scoring[s].get("demand", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["demand"] = st.selectbox(
                t(lang, "Demande politique", "Political demand"),
                options=opts,
                index=idx,
                format_func=score_format(lang, "demand"),
                help=t(
                    lang,
                    "Définition : importance de l’indicateur pour le pilotage des politiques publiques, la redevabilité et les priorités.",
                    "Definition: importance for steering public policies, accountability and priorities.",
                ),
                key=f"sc_dem_{s}",
            )

        with c2:
            prev = scoring[s].get("availability", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["availability"] = st.selectbox(
                t(lang, "Disponibilité actuelle", "Current availability"),
                options=opts,
                index=idx,
                format_func=score_format(lang, "availability"),
                help=t(
                    lang,
                    "Définition : l’indicateur est-il déjà produit régulièrement, avec une couverture et une qualité suffisantes, et sous une forme utilisable ? (Bonne = score plus élevé).",
                    "Definition: is the indicator already produced regularly with sufficient coverage and quality, in a usable form? (Good = higher score).",
                ),
                key=f"sc_avail_{s}",
            )

        with c3:
            prev = scoring[s].get("feasibility", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["feasibility"] = st.selectbox(
                t(lang, "Faisabilité (12–24 mois)", "Feasibility (12–24 months)"),
                options=opts,
                index=idx,
                format_func=score_format(lang, "feasibility"),
                help=t(
                    lang,
                    "Définition : capacité réaliste à améliorer ou produire l’indicateur d’ici 12–24 mois, compte tenu des sources, capacités et prérequis.",
                    "Definition: realistic ability to improve or produce the indicator within 12–24 months, considering sources, capacities and prerequisites.",
                ),
                key=f"sc_fea_{s}",
            )

    resp_set("scoring", scoring)

    errs = validate_r5(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_6(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 6 : Dimension genre", "Section 6: Gender dimension"))
    st.markdown(
        t(
            lang,
            "Indiquez si les statistiques prioritaires doivent intégrer ces dimensions (Oui/Non/Selon indicateur/NSP).",
            "Indicate whether priority indicators should integrate these dimensions (Yes/No/Indicator-specific/UK)."
        )
    )

    options = [
        (t(lang, "Oui", "Yes"), "YES"),
        (t(lang, "Non", "No"), "NO"),
        (t(lang, "Selon indicateur", "Indicator-specific"), "SPEC"),
        (UK_FR if lang == "fr" else UK_EN, "UK"),
    ]
    labels = [x[0] for x in options]
    code_map = {x[0]: x[1] for x in options}

    items_fr = [
        "Désagrégation par sexe",
        "Désagrégation par âge",
        "Milieu urbain / rural",
        "Handicap",
        "Quintile de richesse",
        "Violences basées sur le genre (VBG)",
        "Temps domestique non rémunéré",
    ]
    items_en = [
        "Disaggregation by sex",
        "Disaggregation by age",
        "Urban / rural",
        "Disability",
        "Wealth quintile",
        "Gender-based violence (GBV)",
        "Unpaid domestic work",
    ]
    items = items_fr if lang == "fr" else items_en

    tbl = resp_get("gender_table", {})
    if not isinstance(tbl, dict):
        tbl = {}

    for it in items:
        rev_map = {v: k for k, v in code_map.items()}
        prev_code = tbl.get(it, None)
        idx = labels.index(rev_map[prev_code]) if prev_code in rev_map else None
        chosen = st.radio(it, options=labels, index=idx, horizontal=True, key=f"gender_{it}")
        tbl[it] = code_map.get(chosen, None)

    resp_set("gender_table", tbl)

    errs = validate_r6(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def validate_r7(lang: str) -> List[str]:
    errs: List[str] = []
    p1 = (resp_get("gender_priority_1", "") or "").strip()
    p2 = (resp_get("gender_priority_2", "") or "").strip()
    p3 = (resp_get("gender_priority_3", "") or "").strip()
    other = (resp_get("gender_priority_other", "") or "").strip()

    if not p1:
        errs.append(
            t(
                lang,
                "Rubrique 7 : veuillez sélectionner au moins une priorité genre (Priorité 1).",
                "Section 7: please select at least one gender priority (Priority 1).",
            )
        )

    # No rank 3 without rank 2
    if p3 and not p2:
        errs.append(
            t(
                lang,
                "Rubrique 7 : veuillez renseigner la Priorité 2 avant la Priorité 3.",
                "Section 7: please fill Priority 2 before Priority 3.",
            )
        )

    # Uniqueness
    chosen = [x for x in [p1, p2, p3] if x]
    if len(set(chosen)) != len(chosen):
        errs.append(
            t(
                lang,
                "Rubrique 7 : les priorités genre doivent être différentes (pas de doublons).",
                "Section 7: gender priorities must be distinct (no duplicates).",
            )
        )

    # Other text required if OTHER selected
    if "OTHER" in chosen and not other:
        errs.append(
            t(
                lang,
                "Rubrique 7 : précisez l’option « Autre ».",
                "Section 7: please specify the 'Other' option.",
            )
        )

    return errs


def rubric_7(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 7 : Priorités genre", "Section 7: Gender priorities"))
    st.markdown(
        t(
            lang,
            "Sélectionnez de 1 à 3 priorités genre en commençant par la plus importante.",
            "Select 1 to 3 gender priorities, starting with the most important.",
        )
    )

    gp_opts = ["", "ECO", "SERV", "GBV", "PART_DEC", "CARE", "OTHER"]
    gp_labels = {
        "ECO": t(lang, "Autonomisation économique", "Economic empowerment"),
        "SERV": t(lang, "Accès aux services", "Access to services"),
        "GBV": t(lang, "Violences basées sur le genre (VBG)", "Gender based violence (GBV)"),
        "PART_DEC": t(lang, "Participation aux instances décisionnelles", "Participation in decision-making bodies"),
        "CARE": t(lang, "Temps domestique non rémunéré", "Unpaid domestic and care work"),
        "OTHER": t(lang, "Autre", "Other"),
        "": t(lang, "— Sélectionner —", "— Select —"),
    }

    # Rank 1 (required)
    prev1 = (resp_get("gender_priority_1", "") or "").strip()
    idx1 = gp_opts.index(prev1) if prev1 in gp_opts else 0
    p1 = st.selectbox(
        t(lang, "Vos trois (3) priorités genre – Priorité 1 (obligatoire)", "Your three (3) gender priorities – Priority 1 (required)"),
        options=gp_opts,
        index=idx1,
        format_func=lambda k: gp_labels.get(k, k),
        key="gender_priority_1_select",
    )
    resp_set("gender_priority_1", p1)
    # Backward compatibility (previous single-priority field)
    resp_set("gender_priority_main", p1)

    # Rank 2 (optional), exclude already chosen (except empty)
    opts2 = [""] + [k for k in gp_opts if k not in ["", p1]]
    prev2 = (resp_get("gender_priority_2", "") or "").strip()
    idx2 = opts2.index(prev2) if prev2 in opts2 else 0
    p2 = st.selectbox(
        t(lang, "Priorité 2 (optionnelle)", "Priority 2 (optional)"),
        options=opts2,
        index=idx2,
        format_func=lambda k: gp_labels.get(k, k),
        key="gender_priority_2_select",
    )
    resp_set("gender_priority_2", p2)

    # Rank 3 (optional), exclude already chosen (except empty)
    opts3 = [""] + [k for k in gp_opts if k not in ["", p1, p2]]
    prev3 = (resp_get("gender_priority_3", "") or "").strip()
    idx3 = opts3.index(prev3) if prev3 in opts3 else 0
    p3 = st.selectbox(
        t(lang, "Priorité 3 (optionnelle)", "Priority 3 (optional)"),
        options=opts3,
        index=idx3,
        format_func=lambda k: gp_labels.get(k, k),
        key="gender_priority_3_select",
    )
    resp_set("gender_priority_3", p3)

    chosen_any = [x for x in [p1, p2, p3] if x]
    if "OTHER" in chosen_any:
        other = st.text_input(
            t(lang, "Autre : préciser", "Other: please specify"),
            key="gender_priority_other_input",
            value=resp_get("gender_priority_other", ""),
        )
        resp_set("gender_priority_other", (other or "").strip())
    else:
        resp_set("gender_priority_other", "")

    errs = validate_r7(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))

def rubric_8(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 8 : Capacité et faisabilité (12–24 mois)", "Section 8: Capacity and feasibility (12–24 months)"))
    st.markdown(
        t(
            lang,
            "Évaluez le niveau de disponibilité et d’adéquation des moyens pour produire les statistiques prioritaires dans les 12–24 mois à venir.",
            "Assess the availability and adequacy of resources to produce priority statistics in the coming 12–24 months."
        )
    )

    scale = [
        (t(lang, "Élevé", "High"), "HIGH"),
        (t(lang, "Moyen", "Medium"), "MED"),
        (t(lang, "Faible", "Low"), "LOW"),
        (UK_FR if lang == "fr" else UK_EN, "UK"),
    ]
    labels = [x[0] for x in scale]
    code_map = {x[0]: x[1] for x in scale}

    st.caption(t(lang, "Échelle : Élevé = capacité suffisante et opérationnelle ; Moyen = partiellement disponible ; Faible = insuffisant ; NSP = ne sait pas.",
                   "Scale: High = sufficient and operational; Medium = partially available; Low = insufficient; DK = does not know."))

    items_fr = [
        "Compétences statistiques disponibles",
        "Accès aux données administratives",
        "Financement disponible",
        "Outils numériques (collecte, traitement, diffusion)",
        "Cadre juridique pour le partage de données",
        "Coordination interinstitutionnelle",
    ]
    items_en = [
        "Available statistical skills",
        "Access to administrative data",
        "Available funding",
        "Digital tools (collection, processing, dissemination)",
        "Legal framework for data sharing",
        "Inter-institutional coordination",
    ]
    items = items_fr if lang == "fr" else items_en

    helps_fr = [
        "Ressources humaines : disponibilité de statisticiens/analystes qualifiés et expérience pertinente.",
        "Accès aux données administratives : disponibilité, qualité, régularité et conditions d’accès pour usage statistique.",
        "Financement : budget disponible et soutenable pour la production, y compris opérations de collecte/traitement.",
        "Outils numériques : disponibilité et adéquation des outils pour collecte, traitement, stockage, diffusion, interopérabilité (logiciels, matériel, connectivité, sécurité).",
        "Cadre juridique : existence et applicabilité des textes/accords permettant le partage de données à des fins statistiques (lois, décrets, protocoles, MoU, clauses de confidentialité).",
        "Coordination : mécanismes de coordination interinstitutionnelle (comités, conventions, échanges réguliers, standards communs).",
    ]
    helps_en = [
        "Human resources: availability of qualified statisticians/analysts and relevant experience.",
        "Access to administrative data: availability, quality, timeliness and conditions of access for statistical use.",
        "Funding: available and sustainable budget for production, including collection/processing operations.",
        "Digital tools: availability and adequacy of tools for collection, processing, storage, dissemination, interoperability (software, hardware, connectivity, security).",
        "Legal framework: existence and enforceability of texts/agreements enabling data sharing for statistical purposes (laws, decrees, protocols, MoUs, confidentiality clauses).",
        "Coordination: inter-institutional coordination mechanisms (committees, agreements, regular exchanges, shared standards).",
    ]
    helps = helps_fr if lang == "fr" else helps_en

    tbl = resp_get("capacity_table", {})
    if not isinstance(tbl, dict):
        tbl = {}

    rev_map = {v: k for k, v in code_map.items()}
    for it, hp in zip(items, helps):
        prev_code = tbl.get(it, None)
        idx = labels.index(rev_map[prev_code]) if prev_code in rev_map else None
        chosen = st.radio(it, options=labels, index=idx, horizontal=True, key=f"cap_{it}", help=hp)
        tbl[it] = code_map.get(chosen, None)

    resp_set("capacity_table", tbl)

    errs = validate_r8(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_9(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 9 : Harmonisation et qualité", "Section 9: Harmonization and quality"))
    st.markdown(
        t(
            lang,
            "Indiquez 1 à 3 exigences attendues en matière d’harmonisation et d’assurance qualité.",
            "Indicate 1 to 3 expectations regarding harmonization and quality assurance."
        )
    )

    opts_fr = [
        "Manuels de normes et méthodes communes (par domaine) disponibles",
        "Cadre d’assurance qualité fonctionnel",
        "Procédures de validation et certification des données",
        "Mécanismes de cohérence des données nationales entre secteurs",
        "Renforcement des capacités techniques du SSN",
        "Renforcement du leadership de l’INS au sein du SSN",
        "Groupes techniques spécialisés (GTS/UA) opérationnels",
        "Autre (préciser) ",
     ]
    opts_en = [
        "Manuals on common standards and methods (by domain) available",
        "Functional quality assurance framework (quality toolkit) ",
        "Data validation and certification procedures (certified quality) ",
        "Toolkit / mechanisms for cross-sector consistency of national data",
        "Strengthening NSS technical capacity",
        "Strengthening NSO leadership within the NSS",
        "Specialized Technical Groups (STGs/AU) operational",
        "Other (specify) ",
    ]
    opts = opts_fr if lang == "fr" else opts_en
    # Stabilité mobile : initialiser le widget une seule fois
    if "r9_multiselect" not in st.session_state:
        st.session_state["r9_multiselect"] = resp_get("quality_expectations", [])
    sel = st.multiselect(t(lang, "Sélectionnez", "Select"), options=opts, key="r9_multiselect", max_selections=3)
    resp_set("quality_expectations", sel)
    if ("Autre" in sel) or ("Other" in sel):
        st.text_input(t(lang, "Préciser (Autre)", "Specify (Other)"),
                      key="q9_other_input", value=resp_get("quality_other", ""))
        resp_set("quality_other", st.session_state.get("q9_other_input", "").strip())
    else:
        resp_set("quality_other", "")

    errs = validate_r9(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))

    # Auto-save draft (mobile)
    autosave_draft(force=False)



def rubric_10(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 10 : Diffusion", "Section 10: Dissemination"))
    st.markdown(
        t(
            lang,
            "Indiquez 1 à 3 canaux de diffusion jugés les plus utiles pour les statistiques prioritaires.",
            "Indicate 1 to 3 dissemination channels you find most useful for priority statistics."
        )
    )
    opts_fr = [
        "Portail web / tableaux de bord",
        "Communiqués / notes de conjoncture",
        "Microdonnées anonymisées (accès sécurisé)",
        "API / Open data",
        "Ateliers et webinaires",
        "Autre",
    ]
    opts_en = [
        "Web portal / dashboards",
        "Press releases / bulletins",
        "Anonymized microdata (secure access)",
        "API / Open data",
        "Workshops and webinars",
        "Other",
    ]
    opts = opts_fr if lang == "fr" else opts_en
    # Éviter les problèmes de clic (init du widget une seule fois)
    if "r10_multiselect" not in st.session_state:
        st.session_state["r10_multiselect"] = resp_get("dissemination_channels", [])
    sel = st.multiselect(
        t(lang, "Sélectionnez", "Select"),
        options=opts,
        max_selections=3,
        key="r10_multiselect",
        help=t(lang, "Choisissez les canaux de diffusion les plus utiles.", "Select the most useful dissemination channels.")
    )
    resp_set("dissemination_channels", sel)
    if ("Autre" in sel) or ("Other" in sel):
        st.text_input(t(lang, "Préciser (Autre)", "Specify (Other)"),
                      key="q10_other_input", value=resp_get("dissemination_other", ""))
        resp_set("dissemination_other", st.session_state.get("q10_other_input", "").strip())
    else:
        resp_set("dissemination_other", "")

    errs = validate_r10(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))

    # Auto-save draft (mobile)
    autosave_draft(force=False)



def rubric_11(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 11 : Sources de données pertinentes", "Section 11: Relevant data sources"))
    st.markdown(
        t(
            lang,
            "Sélectionnez **2 à 4** sources de données les plus importantes pour produire les statistiques prioritaires.",
            "Select **2 to 4** of the most important data sources to produce the priority statistics.",
        )
    )

    opts_fr = [
        "Enquêtes ménages",
        "Enquêtes entreprises",
        "Recensements",
        "Données administratives",
        "Registres état-civil",
        "Données géospatiales",
        "Données privées",
        "Autres",
    ]
    opts_en = [
        "Household surveys",
        "Enterprise surveys",
        "Censuses",
        "Administrative data",
        "Civil registration and vital statistics (CRVS)",
        "Geospatial data",
        "Private data",
        "Other",
    ]

    options = opts_fr if lang == "fr" else opts_en

    prev = resp_get("data_sources", [])
    if not isinstance(prev, list):
        prev = []

    sel = st.multiselect(
        t(
            lang,
            "2 à 4 sources de données les plus pertinentes",
            "2 to 4 most relevant data sources",
        ),
        options=options,
        default=[x for x in prev if x in options],
        max_selections=4,
        help=t(
            lang,
            "Choisissez entre 2 et 4 options. Si vous choisissez Autres, précisez.",
            "Choose between 2 and 4 options. If you select Other, please specify.",
        ),
        key="data_sources_multiselect",
    )
    resp_set("data_sources", sel)

    other_label = "Autres" if lang == "fr" else "Other"
    if other_label in sel:
        other = st.text_input(
            t(lang, "Autres : préciser", "Other: please specify"),
            key="data_sources_other_input",
            value=resp_get("data_sources_other", ""),
        )
        resp_set("data_sources_other", (other or "").strip())
    else:
        resp_set("data_sources_other", "")

    errs = validate_r11(lang)
    if errs:
        for e in errs:
            st.error(e)

def rubric_12(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 12 : Questions ouvertes", "Section 12: Open questions"))
    st.markdown(
        t(
            lang,
            "Ces questions sont **optionnelles**. Vous pouvez les laisser vides. Toutefois, elles sont présentées **une à une** pour faciliter la saisie.",
            "These questions are **optional**. You may leave them blank. They are presented **one by one** to facilitate completion.",
        )
    )

    if "r12_substep" not in st.session_state:
        st.session_state["r12_substep"] = 0  # 0..2=open questions, 3=confirmation

    sub = int(st.session_state.get("r12_substep", 0) or 0)

    if sub == 0:
        st.markdown("#### " + t(lang, "Question 1 / 3", "Question 1 / 3"))
        q1 = st.text_area(
            t(lang, "1) Commentaires / recommandations clés", "1) Key comments / recommendations"),
            value=resp_get("open_q1", ""),
            height=160,
            key="open_q1_input",
        )
        resp_set("open_q1", (q1 or "").strip())
        if not resp_get("open_q1", ""):
            st.warning(t(lang, "Alerte : la question 1 est vide (vous pouvez tout de même continuer).",
                         "Warning: question 1 is empty (you can still proceed)."))

    elif sub == 1:
        st.markdown("#### " + t(lang, "Question 2 / 3", "Question 2 / 3"))
        q2 = st.text_area(
            t(
                lang,
                "2) Un ou des indicateur(s) statistique(s) socio-économique(s) essentiel(s) manquant(s) et justification(s)",
                "2) One or more missing essential socio-economic statistical indicator(s) and justification(s)",
            ),
            value=resp_get("open_q2", ""),
            height=160,
            key="open_q2_input",
        )
        resp_set("open_q2", (q2 or "").strip())
        if not resp_get("open_q2", ""):
            st.warning(t(lang, "Alerte : la question 2 est vide (vous pouvez tout de même continuer).",
                         "Warning: question 2 is empty (you can still proceed)."))

    elif sub == 2:
        st.markdown("#### " + t(lang, "Question 3 / 3", "Question 3 / 3"))
        q3 = st.text_area(
            t(lang, "3) Besoins de soutien (technique, financier, etc.)", "3) Support needs (technical, financial, etc.)"),
            value=resp_get("open_q3", ""),
            height=160,
            key="open_q3_input",
        )
        resp_set("open_q3", (q3 or "").strip())
        if not resp_get("open_q3", ""):
            st.warning(t(lang, "Alerte : la question 3 est vide (vous pouvez tout de même continuer).",
                         "Warning: question 3 is empty (you can still proceed)."))

    else:
        st.markdown("#### " + t(lang, "Confirmation", "Confirmation"))
        st.info(
            t(
                lang,
                "Dernière étape : merci d’indiquer si vous avez consulté d’autres collègues. Cette question est obligatoire.",
                "Final step: please indicate whether you consulted other colleagues. This question is mandatory.",
            )
        )

        cc_opts = ["", "YES", "NO"]
        cc_labels = {"YES": t(lang, "Oui", "Yes"), "NO": t(lang, "Non", "No")}
        prev_cc = (resp_get("consulted_colleagues", "") or "").strip()
        idx_cc = cc_opts.index(prev_cc) if prev_cc in cc_opts else 0
        chosen_cc = st.selectbox(
            t(
                lang,
                "Avez-vous consulté d’autres collègues pour remplir ce questionnaire ?",
                "Did you consult other colleagues to complete this questionnaire?",
            ),
            options=cc_opts,
            index=idx_cc,
            format_func=lambda k: (t(lang, "— Sélectionner —", "— Select —") if k == "" else cc_labels.get(k, k)),
            key="consulted_colleagues_select",
        )
        resp_set("consulted_colleagues", chosen_cc)

    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 3])
    with col1:
        prev_disabled = (sub <= 0)
        if st.button(t(lang, "⬅ Question précédente", "⬅ Previous question"), disabled=prev_disabled, key="r12_prev_btn"):
            st.session_state["r12_substep"] = max(0, sub - 1)
            st.rerun()
    with col2:
        if sub < 2:
            next_label = t(lang, "Question suivante ➡", "Next question ➡")
        elif sub == 2:
            next_label = t(lang, "Aller à la confirmation ➡", "Go to confirmation ➡")
        else:
            next_label = t(lang, "OK (rubrique terminée)", "OK (section completed)")
        next_disabled = (sub >= 3)
        if st.button(next_label, disabled=next_disabled, key="r12_next_btn"):
            st.session_state["r12_substep"] = min(3, sub + 1)
            st.rerun()
    with col3:
        st.caption(t(lang, "Progression : 1/3 → 2/3 → 3/3 → Confirmation.",
                     "Progress: 1/3 → 2/3 → 3/3 → Confirmation."))

    errs = validate_r12(lang)
    if errs:
        for e in errs:
            st.error(e)

def rubric_send(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "ENVOYER le questionnaire", "SUBMIT questionnaire"))

    errors = validate_all(lang)
    if errors:
        st.error(t(lang, "Le questionnaire contient des erreurs bloquantes :", "There are blocking errors:"))
        st.write("\n".join([f"- {e}" for e in errors]))
        st.info(t(lang, "Retournez aux rubriques concernées via la navigation.", "Go back to the relevant sections using navigation."))
        return

    # Optional warnings
    if not resp_get("open_q1", "") or not resp_get("open_q2", "") or not resp_get("open_q3", ""):
        st.warning(t(lang, "Certaines questions ouvertes sont vides (optionnel).", "Some open questions are empty (optional)."))

    st.info(
        t(
            lang,
            "Tout est prêt. Cliquez sur **ENVOYER** pour soumettre votre questionnaire.",
            "Everything is ready. Click **SUBMIT** to send your questionnaire.",
        )
    )

    # Empêcher les envois multiples (par email + par session)
    email = (resp_get("email", "") or "").strip()
    already_in_db = db_email_exists(email) if email else False
    already_in_session = bool(st.session_state.get("submitted_once", False))

    if already_in_db:
        st.error(
            t(
                lang,
                "Ce questionnaire a déjà été envoyé avec cet email. Un seul envoi est autorisé.",
                "This questionnaire has already been submitted with this email. Only one submission is allowed.",
            )
        )

    if already_in_session and not already_in_db:
        st.info(
            t(
                lang,
                "Ce navigateur a déjà effectué un envoi. Pour un nouvel envoi, utilisez un autre email / session.",
                "This browser session already submitted once. For a new submission, use another email / session.",
            )
        )

    disable_submit = already_in_db or already_in_session

    if st.button(t(lang, "✅ ENVOYER et enregistrer", "✅ SUBMIT and save"), disabled=disable_submit):
        submission_id = str(uuid.uuid4())
        payload = st.session_state.responses.copy()
        payload["submission_id"] = submission_id
        payload["submitted_at_utc"] = now_utc_iso()
        payload["scoring_version"] = SCORING_VERSION

        ok_db, msg_db, ok_json, msg_json = db_save_submission(submission_id, lang, email, payload)
        if not ok_db:
            st.error(msg_db)
            return

        st.success(t(lang, "Merci ! Votre questionnaire a été enregistré.", "Thank you! Your submission has been saved."))

        # Bloquer les envois multiples pour cette session
        st.session_state.submitted_once = True
        st.caption(f"ID : {submission_id}")

        # Information simple côté répondant (ne pas exposer les détails techniques)
        st.info(t(lang, "Envoi terminé. Vous pouvez fermer cette page.", "Submission complete. You can close this page."))

        st.session_state.submission_id = submission_id

# =========================
# Main
# =========================

def main() -> None:
    st.set_page_config(page_title=APP_TITLE_FR, layout="wide")
    init_session()
    maybe_restore_draft()

    # Language toggle
    st.sidebar.title("🌐")
    lang = st.sidebar.selectbox(
        "Langue / Language",
        options=["fr", "en"],
        index=0 if st.session_state.lang == "fr" else 1
    )
    st.session_state.lang = lang

    # Admin access
    qp = get_query_params()
    is_admin = "admin" in qp and qp["admin"] and qp["admin"][0] in ["1", "true", "yes"]

    df_long = load_longlist()

    # Header
    st.title(t(lang, APP_TITLE_FR, APP_TITLE_EN))
    st.caption(t(lang, "Application unifiée (FR/EN) – codes masqués – contrôles qualité intégrés.",
                 "Unified app (FR/EN) – hidden codes – built-in quality controls."))
    if st.session_state.get("draft_exists") and not st.session_state.get("draft_resume_notice_shown"):
        st.warning(
            t(
                lang,
                "La saisie est sauvegardée. En cas de suspension de moins de 48 heures, reprenez-la là où vous vous étiez arrêté en ré-ouvrant le lien contenant rid (à conserver / mettre en favori / retrouver dans l’historique).",
                "Your entry is saved. If you pause for less than 48 hours, resume where you left off by reopening the link containing rid (bookmark / save / find it in your browser history).",
            ),
            icon="💾",
        )
        st.session_state.draft_resume_notice_shown = True

    # Sidebar navigation
    steps = get_steps(lang)
    render_sidebar(lang, steps)

    # Admin view
    if is_admin:
        if not st.session_state.admin_authed:
            admin_login(lang)
            return
        admin_dashboard(lang)
        return

    # Normal flow
    step_key = steps[st.session_state.nav_idx][0]

    if step_key == "R1":
        rubric_1(lang)
    elif step_key == "R2":
        rubric_2(lang)
    elif step_key == "R3":
        rubric_3(lang)
    elif step_key == "R4":
        rubric_4(lang, df_long)
    elif step_key == "R5":
        rubric_5(lang, df_long)
    elif step_key == "R6":
        rubric_6(lang)
    elif step_key == "R7":
        rubric_7(lang)
    elif step_key == "R8":
        rubric_8(lang)
    elif step_key == "R9":
        rubric_9(lang)
    elif step_key == "R10":
        rubric_10(lang)
    elif step_key == "R11":
        rubric_11(lang)
    elif step_key == "R12":
        rubric_12(lang)
    elif step_key == "SEND":
        rubric_send(lang, df_long)

    st.divider()
    nav_buttons(lang, steps, df_long)


if __name__ == "__main__":
    main()
